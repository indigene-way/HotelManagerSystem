#! /usr/bin/python
# -*- coding: utf-8 -*-

import sys, os
import re 
import uuid
import time

import sqlite3
from PyQt5.QtSql import *

from PyQt5 import QtCore, QtGui, uic
from PyQt5.QtCore import *
from PyQt5.QtCore import QDate
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *

import matplotlib.pyplot as plt
# %matplotlib  inline
import numpy as np 

from docx import *

#============================================================================================================DATA BASE
conn = sqlite3.connect('dataBase.db')

query = conn.cursor()
				
#===========================================================================TABLE LOGIN
# query.execute("DROP TABLE Login")
try:
	query.execute("SELECT id FROM Login ORDER BY id DESC")
except:
	# query.execute("DROP TABLE Login")
	conn.execute("""CREATE TABLE Login (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					user_name VARCHAR(50) ,
					user_password VARCHAR(30))""")
					
	# print("DataBase  Login created succefully")
	query.execute("INSERT INTO Login (user_name, user_password) VALUES ('admin','admin')")
				
#===========================================================================HOTEL
#===========================================================================ROOMS
# query.execute("DROP TABLE Rooms")
try:
	query.execute("SELECT id FROM Rooms")
except:
	conn.execute("""CREATE TABLE Rooms (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					room_name VARCHAR(50) ,
					room_type VARCHAR(30) ,
					room_price INTEGER ,
					room_state INTEGER, 
					room_date_in DATE, 
					room_date_out DATE, 
					room_reserved DATE, 
					room_client VARCHAR(80) )""")
	# print("Rooms TABLE CREATED SUCCESSFULLY")

#===========================================================================CLIENTS
# query.execute("DROP TABLE Clients")
try:
	query.execute("SELECT id FROM Clients")
except:
	conn.execute("""CREATE TABLE Clients (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					client_name VARCHAR(50) ,
					client_sexe VARCHAR(30) ,
					client_national VARCHAR(50) , 
					client_civil VARCHAR(10), 
					client_birth DATE, 
					client_birth_place VARCHAR(50), 
					client_adress VARCHAR(255), 
					client_job VARCHAR(50),
					client_company VARCHAR(50), 
					client_kids INTEGER, 
					client_id VARCHAR(50),
					client_id_num VARCHAR(50),
					client_id_delivery VARHAR(30),
					client_id_at VARCHAR(50),
					client_id_by VARCHAR(50),
					client_enter VARHAR(30),
					client_sejour INTEGER, 
					client_date_in DATE, 
					client_date_out DATE, 
					client_pay_mthd VARCHAR(20),
					client_room VARCHAR(20),
					client_total INTEGER)""")
	# print("Clients TABLE CREATED SUCCESSFULLY")
#===========================================================================TABLE USER
try:
	query.execute("SELECT user FROM User ")
except:
	# query.execute("DROP TABLE User")
	conn.execute("""CREATE TABLE User (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					user VARCHAR(50))""")
					
	# print("DataBase  User created succefully")
	query.execute("INSERT INTO User(user) VALUES('')")
				
#===========================================================================TABLE SOCIETY
try:
	query.execute("SELECT id FROM Society ORDER BY id DESC")
except:
	# query.execute("DROP TABLE Society")
	conn.execute("""CREATE TABLE Society (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					society_name VARCHAR(50) ,
					society_number VARCHAR(30))""")
					
	# print("DataBase  Society created succefully")
	query.execute("INSERT INTO Society(society_name, society_number) VALUES ('AssyDevSociety','0549484715')")
#===========================================================================TABLE CATEGORY
try:
	query.execute("SELECT id FROM Categories ORDER BY id DESC")
	# print(str(query.fetchone()[0]).strip("(',')"))
except:
	conn.execute("""CREATE TABLE Categories (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					category_name VARCHAR(50) ,
					category_sub_1 VARCHAR(30) ,
					category_sub_2 VARCHAR(30) ,
					category_sub_3 VARCHAR(30) )""")
					
	# print("DataBase  Categories created succefully")
	query.execute("INSERT INTO Categories (category_name) VALUES\
					('Entrées'),\
					('Plats'),\
					('Sandwiches'),\
					('Pizzas'),\
					('Boissons Fraiches'),\
					('Boissons Chaudes'),\
					('Désserts'),\
					('Jus et Smoothies'),\
					('Autres'),\
					('Supplements');")

#===========================================================================TABLE Products

# query.execute("DROP TABLE Products")
try:
	query.execute("SELECT id FROM Products ORDER BY id DESC")
except:
	# query.execute("DROP TABLE Products")
	conn.execute("""CREATE TABLE Products (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					product_category_name VARCHAR(50),
					product_subCategory VARCHAR(50),
					product_name VARCHAR(50) ,
					product_price INTEGER ,
					product_qnt INTEGER ,
					product_comptable INTEGER ,
					product_stockable INTEGER ,
					product_stock INTEGER ,
					product_alert INTEGER ,
					product_total INTEGER,
					product_comptable_cst INTEGER)""")
	# print("PRODUCT TABLE CREATED SUCCESSFULLY")

#============================================================= REFERENCEMENT TICKETS

try:
	query.execute("SELECT id FROM Ref ORDER BY id DESC")
except:
	# query.execute("DROP TABLE Ref")
	conn.execute("""CREATE TABLE Ref (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE, 
				ref INTEGER )""")


	query.execute("INSERT INTO Ref (ref) VALUES (0)")
	query.execute("INSERT INTO Ref (ref) VALUES (0)")
	# print("TABLE Ref Created")

#===========================================================================REGISTER CATEGORY
try:
	query.execute("SELECT id FROM Register ORDER BY id DESC")
except:
	# query.execute("DROP TABLE Register")
	conn.execute("""CREATE TABLE Register (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					register_date DATE,
					register_sum_init INTEGER,
					register_recette_total INTEGER,
					register_depense_total INTEGER,
					register_ajout_total INTEGER )""")
					
#===========================================================================DEPENSE 
# query.execute("DROP TABLE Register_dep")
try:
	query.execute("SELECT id FROM Register_dep ORDER BY id DESC")
except:
	conn.execute("""CREATE TABLE Register_dep (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					dep_date DATE,
					dep_type VARCHAR(15),
					dep_description VARCHAR(255),
					dep_value INTEGER)""")
					
	# print("DataBase  Register_dep created succefully")
					
#===========================================================================AJOUT 
# query.execute("DROP TABLE Register_add")
try:
	query.execute("SELECT id FROM Register_add ORDER BY id DESC")
except:
	conn.execute("""CREATE TABLE Register_add (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					add_date DATE,
					add_depo VARCHAR(50),
					add_description VARCHAR(255),
					add_value INTEGER)""")
					
	# print("DataBase  Register_add created succefully")
				
#===========================================================================TABLE TABLE
# query.execute("DROP TABLE Tables")
try:
	query.execute("SELECT id FROM Tables ORDER BY id DESC")
except:
	conn.execute("""CREATE TABLE Tables (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					table_name VARCHAR(15),
					table_place VARCHAR(20),
					table_total INTEGER,
					table_state VARCHAR(10))""")
					
	# print("DataBase Tables created succefully")
	query.execute("INSERT INTO Tables (table_name) VALUES ('A Emporter')")
				
#===========================================================================TABLE_CONTENT TABLE
# query.execute("DROP TABLE Tables_content")
try :
	query.execute("SELECT id FROM Tables_content")
except:
	conn.execute("""CREATE TABLE Tables_content (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					table_name VARCHAR(15),
					table_prod VARCHAR(30),
					table_price INTEGER,
					table_qnt INTEGER,
					table_prodTotal INTEGER,
					table_date DATE)""")		
	# print("DataBase Tables_content created succefully")
	
#============================================================= MAC ADRESS
# query.execute("DROP TABLE MAC")
try:
	query.execute("SELECT id FROM MAC ORDER BY id DESC")
except:
	query.execute("""CREATE TABLE MAC (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE,
					Key VARCHAR(50))""")
	# print("MAC CREATED")


	query.execute("INSERT INTO MAC (Key) VALUES ('')")
	

#============================================================= TABLE PROVIDER

try:
	query.execute("SELECT id FROM Providers ORDER BY id DESC")
except:
	# query.execute("DROP TABLE Providers")

	conn.execute("""CREATE TABLE Providers (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					provider_name VARCHAR(50) ,
					provider_category VARCHAR(50),
					provider_phone VARCHAR(10) ,
					provider_date_add DATE ,
					provider_date_transact DATE,
					provider_product1 VARCHAR(50),
					provider_product2 VARCHAR(50),
					provider_product3 VARCHAR(50),
					provider_product4 VARCHAR(50),
					provider_product5 VARCHAR(50),
					provider_price1 INTEGER,
					provider_price2 INTEGER,
					provider_price3 INTEGER,
					provider_price4 INTEGER,
					provider_price5 INTEGER)""")

	# print("DataBase Providers created succefully")

#============================================================Provider HISTORY

try:
	query.execute("SELECT id FROM Providers_History ORDER BY id DESC")
except:
	# query.execute("DROP TABLE Providers_History")
	conn.execute("""CREATE TABLE Providers_History (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					transact_name VARCHAR(50) ,
					transact_category VARCHAR(50),
					transact_phone VARCHAR(10) ,
					transact_products VARCHAR(50),
					transact_quantite INTEGER,
					transact_total INTEGER,
					transact_date_transact DATE)""")

	# print("DataBase Providers_History created succefully")

#============================================================WORKERS CATEGORY

try:
	query.execute("SELECT id FROM Workers_category ORDER BY id DESC")
except :
	# query.execute("DROP TABLE Workers_category")
	conn.execute("""CREATE TABLE Workers_category (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					category_name VARCHAR(50))""")

	# print("DataBase workers created succefully")
	
	query.execute("INSERT INTO Workers_category (category_name) values ('Manager')")
	query.execute("INSERT INTO Workers_category (category_name) values ('Cuisinier')")
	query.execute("INSERT INTO Workers_category (category_name) values ('Serveur')")
	query.execute("INSERT INTO Workers_category (category_name) values ('Caissier')")
	query.execute("INSERT INTO Workers_category (category_name) values ('Agent Entretien')")
	query.execute("INSERT INTO Workers_category (category_name) values ('Agent Sécurité')")
	query.execute("INSERT INTO Workers_category (category_name) values ('')")
	query.execute("INSERT INTO Workers_category (category_name) values ('')")
	query.execute("INSERT INTO Workers_category (category_name) values ('')")
	query.execute("INSERT INTO Workers_category (category_name) values ('')")

	# print("DataBase INSERTED created succefully")
	
#============================================================WORKER

try:
	query.execute("SELECT id FROM Workers ORDER BY id DESC")
except :
	# query.execute("DROP TABLE Workers")
	conn.execute("""CREATE TABLE Workers (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					worker_name VARCHAR(50) ,
					worker_category VARCHAR(50),
					worker_number VARCHAR(50),
					worker_salary INTEGER ,
					worker_salary_type VARCHAR(50),
					worker_start DATE,
					worker_payday DATE,
					worker_advance INTEGER,
					worker_advance_date DATE)""")

	# print("DataBase workers created succefully")
	
#============================================================WORKER AVANCE

try:
	query.execute("SELECT id FROM Worker_avance ORDER BY id DESC")
except :
	# query.execute("DROP TABLE Workers")
	conn.execute("""CREATE TABLE Worker_avance (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					worker_name VARCHAR(50) ,
					worker_category VARCHAR(50),
					worker_start DATE,
					worker_advance INTEGER,
					worker_advance_date DATE)""")

	# print("DataBase Worker_avance created succefully")
	
#============================================================WORKER SALAIR

try:
	query.execute("SELECT id FROM Worker_salaire ORDER BY id DESC")
except :
	# query.execute("DROP TABLE Workers")
	conn.execute("""CREATE TABLE Worker_salaire (
					id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE ,
					worker_name VARCHAR(50) ,
					worker_category VARCHAR(50),
					worker_start DATE,
					worker_salary INTEGER,
					worker_salary_date DATE)""")

	# print("DataBase Worker_salaire created succefully")

conn.commit()				
#============================================================================================================END DATA BASE
#============================================================================================================HOTEL
#============================================================================================================ROOMS				
qtTableCreator= "DESIGN/DIALOGS/roomCreatorDialog.ui"
Ui_roomCreatorDialog, QtBaseClass = uic.loadUiType(qtTableCreator)

class RoomCreatorDialog(QDialog, Ui_roomCreatorDialog):#EDIT : MODIF Product Name,Price DIALOG
	
	def __init__(self):
		QDialog.__init__(self)
		Ui_roomCreatorDialog.__init__(self)
		self.setupUi(self)
		
		self.setWindowTitle("Edition des Chambres")
		
		self.roomData()
		
	def roomData(self):
		#SELECT Rooms INTO LISTS:
		query.execute("SELECT room_name FROM Rooms ORDER BY id ")
		self.roomNameList = list()
		self.roomNameList = query.fetchall()
		
		query.execute("SELECT room_type FROM Rooms ORDER BY id ")
		self.roomTypeList = list()
		self.roomTypeList = query.fetchall()
		
		
		query.execute("SELECT room_price FROM Rooms ORDER BY id ")
		self.roomPriceList = list()
		self.roomPriceList = query.fetchall()
		
		
		query.execute("SELECT room_state FROM Rooms ORDER BY id ")
		self.roomStateList = list()
		self.roomStateList = query.fetchall()
		
		
		query.execute("SELECT room_date_in FROM Rooms ORDER BY id ")
		self.roomDateInList = list()
		self.roomDateInList = query.fetchall()
		
		
		query.execute("SELECT room_date_out FROM Rooms ORDER BY id ")
		self.roomDateOutList = list()
		self.roomDateOutList = query.fetchall()
		
		
		query.execute("SELECT room_reserved FROM Rooms ORDER BY id ")
		self.roomReservedList = list()
		self.roomReservedList = query.fetchall()
		
		
		query.execute("SELECT room_client FROM Rooms ORDER BY id ")
		self.roomClientList = list()
		self.roomClientList = query.fetchall()
		
#============================================================================================================DESIGN PATTERN EXCEPT FACTORY		
#========== PRODUCT WIDGET UIC CONVERT & Load

qtProductWidget = "DESIGN/WIDGETS/roomWidget.ui" # Enter file here.

Ui_RoomWidget, QtBaseClass = uic.loadUiType(qtProductWidget)

#==========INIT CLASS

class RoomWidget(QWidget, Ui_RoomWidget):

	def __init__(self):
	
		QWidget.__init__(self)#QWidget init
		Ui_RoomWidget.__init__(self)#ui init
		self.setupUi(self)#UI init
		
		self.setWindowTitle("Reception et Chambre")
		
		self.day = time.strftime("%d")
		self.month = time.strftime("%m")
		self.year = time.strftime("%Y")
		
		self.Date = time.strftime("%Y-%m-%d")
		self.localDateTime = time.strftime("%d-%m-%Y      %H:%M:%S")
		self.localDate = time.strftime("%d-%m-%Y")
		self.referenceSelector = ReferenceSelector()
		self.society = self.referenceSelector.society 
		self.numero = self.referenceSelector.numero  
		
		#==========INIT CLASSES OF TYPE DIALOGS FROM MODULE "CreatorDialog"
		self.register = RegisterCreatorDialog()
		self.Edit = RoomCreatorDialog()
		
		#==========DELEGATES INSTANCES :
		self.signals()
		self.messageFactory = MessageFactory()
		
		self.roomDataManager()#DATAMANAGER		
		self.roomListSignal()#SIGNAL : QListWidget item clicked		
		self.EditSignals()		
		self.roomCount()
		self.freeRoomTester()
	
	def freeRoomTester(self):
		self.roomSelector("DISTINCT room_name","room_date_out ='"+self.Date+"'")
		for prod in self.selecterAll :
			self.messageFactory.raiseAlert("La chambre : "+str(prod).strip("(',')")+" a atteint la date limite d'occupation")
	
	def signals(self):
		self.roomEditer.clicked.connect(self.slots)
		
		self.reserve.clicked.connect(self.slots)
		self.libere.clicked.connect(self.slots)
		self.annul.clicked.connect(self.slots)
		self.printer.clicked.connect(self.slots)
		
		self.clientStay.valueChanged.connect(self.slots)
	
	def reserveButton(self):	
		if self.roomState.text() == "Libre" :
			#CLIENT 
			if self.clientName.text() != "" and  self.clientNational.text() != "" and  self.clientId.text() != "" and self.clientStay.value() != 0 :
				self.dateOutCalcul()
				#BINDING CLIENT DATA
				query.execute("INSERT INTO Clients (client_name,client_sexe,client_civil,client_national,client_birth,client_birth_place,\
				client_adress,client_job,client_company,client_kids,client_id,client_id_num,client_id_delivery,client_id_at,client_id_by,client_enter,\
				client_sejour,client_date_in,client_date_out,client_pay_mthd,client_room,client_total)\
				VALUES(\
				'"+self.clientName.text()+"','"+self.clientSexe.currentText()+"','"+self.clientCivil.currentText()+"','"+self.clientNational.text()+"',\
				'"+self.clientBirth.text()+"','"+self.clientBirthPlace.text()+"','"+self.clientAdress.text()+"','"+self.clientJob.text()+"',\
				'"+self.clientCompany.text()+"','"+str(self.clientKids.value())+"',\
				'"+self.clientIdPiece.text()+"','"+self.clientId.text()+"','"+self.clientIdDate.text()+"','"+self.clientIdAt.text()+"',\
				'"+self.clientIdBy.text()+"','"+self.clientEnter.text()+"',\
				'"+str(self.clientStay.value())+"','"+self.clientDateIn.text()+"','"+self.clientDateOut.text()+"','"+self.clientPayMthd.currentText()+"',\
				'"+self.roomNum.text()+"','"+str(self.clientStay.value() * int(self.roomPrice.text()))+"')")
				conn.commit()
				#BINDING ROOM DATA roomState
				query.execute("UPDATE Rooms SET room_state ='1',room_date_in='"+self.Date+"',room_date_out='"+self.clientDateOut.text()+"',\
				room_client='"+self.clientName.text()+"' WHERE room_name ='"+self.roomNum.text()+"'")
				conn.commit()
				
				self.Edit.roomData()
				self.roomCount()
				self.roomCleaner()
				
				self.messageFactory.raiseAdder("Client(e)")
			else:
				self.messageFactory.raiseCaseExcept("Toutle les cases")
		else:
			self.messageFactory.raiseCaseExcept("une chambre LIBRE")
	
	def libereButton(self):	
		if self.roomState.text() == "Occupée" :
			#CLIENT 
				#BINDING ROOM DATA
				query.execute("UPDATE Rooms SET room_state ='0',room_date_out='2999-01-01' WHERE room_name ='"+self.roomNum.text()+"'")
				conn.commit()
				
				self.Edit.roomData()
				self.roomCount()
				self.roomCleaner()
				
				self.messageFactory.raiseAdder("Libération de chambre")
		else:
			self.messageFactory.raiseCaseExcept("une chambre occupée dans la liste")

	def printButton(self):	
		if self.roomState.text() == "Occupée" :
				
				self.messageFactory.raiseSuccess("Fichier imprimé avec succés")
		else:
			self.messageFactory.raiseIndefinedExcept("Aucun client trouvé.\nVeillez séléctionner une chambre occupée.")
			# self.messageFactory.raiseIndefinedExcept("Oups quelque chose s'est mal passé")
		
	def slots(self):
		if self.sender() == self.roomEditer :
			self.Edit.show()
			
		if self.sender() == self.reserve :
			self.reserveButton()
			
		if self.sender() == self.libere :
			self.libereButton()
			
		if self.sender() == self.printer :
			self.printButton()
			
		if self.sender() == self.annul :
			if self.roomState.text() == "Libre" :
				self.roomCleaner()
			elif self.roomState.text() == "" :
				pass
			else:
				self.messageFactory.raiseIndefinedExcept("Cette réservation a déja été validée\nVous ne pouvez pas l'annuler")
		if self.sender() == self.clientStay :
			if self.roomState.text() == "Libre" :
				self.viewTotal.display(int(self.roomPrice.text()) * self.clientStay.value())
			

	def dateOutCalcul(self):
		self.mois = int(self.month)
		self.adding = int(self.day) + self.clientStay.value()
		
		if int(self.mois) == 1 and self.adding > 31 :
				self.mois = self.mois + 1
				self.adding = self.adding - 31
		elif int(self.mois) == 3 and self.adding > 31 :
				self.mois = self.mois + 1
				self.adding = self.adding - 31
		elif int(self.mois) == 4 and self.adding > 30 :
				self.mois = self.mois + 1
				self.adding = self.adding - 30
		elif int(self.mois) == 5 and self.adding > 31 :
				self.mois = self.mois + 1
				self.adding = self.adding - 31
		elif int(self.mois) == 6 and self.adding > 30 :
				self.mois = self.mois + 1
				self.adding = self.adding - 30
		elif int(self.mois) == 7 and self.adding > 31 :
				self.mois = self.mois + 1
				self.adding = self.adding - 31
		elif int(self.mois) == 8 and self.adding > 31 :
				self.mois = self.mois + 1
				self.adding = self.adding - 31
		elif int(self.mois) == 9 and self.adding > 30 :
				self.mois = self.mois + 1
				self.adding = self.adding - 30
		elif int(self.mois) == 10 and self.adding > 31 :
				self.mois = self.mois + 1
				self.adding = self.adding - 31
		elif int(self.mois) == 11 and self.adding > 30 :
				self.mois = self.mois + 1
				self.adding = self.adding - 30
		elif int(self.mois) == 12 and self.adding > 31 :	
				self.mois = 1
				self.adding = self.adding - 31
		
		elif int(self.mois) == 2 and self.year % 4 != 0 : 
				if self.adding > 28 :
					self.mois = self.mois + 1
					self.adding = self.adding - 28
		elif int(self.mois) == 2 and self.year % 4 == 0 : 
			if self.adding > 29 :
				self.mois = self.mois + 1
				self.adding = self.adding - 29
		
		if self.mois <= 9 :
			if self.adding < 10 :
				self.dateOut =str(self.year)+"-0"+str(self.mois)+"-0"+str(self.adding)
			else :
				self.dateOut =str(self.year)+"-0"+str(self.mois)+"-"+str(self.adding)
		else :
			if self.adding < 10 :
				self.dateOut =str(self.year)+"-"+str(self.mois)+"-0"+str(self.adding)
			else:
				self.dateOut =str(self.year)+"-0"+str(self.mois)+"-"+str(self.adding)
			
		self.clientDateOut.setText(str(self.dateOut))	

	def roomDataManager(self):#SELECT PROVIDERS LIST THEM; PRODUCTLIST,	QNTLIST, PRICELIST
			
		#SET PRODUCTS AND PRICES INTO LISTS FROM SELF
		
		query.execute("SELECT DISTINCT room_name FROM Rooms ORDER BY id ")
		self.roomCat = []
		self.roomCat = query.fetchall()	

	def roomCleaner(self):	
		self.roomIndicator.setText("Chambres")
		self.roomIndic.setText("0")
		#ROOMS :
		self.roomNum.setText("")
		self.roomType.setText("")
		self.roomPrice.setText("")
		self.roomState.setText("")
		
		#CLIENT :
		self.clientName.setText("")
		self.clientNational.setText("")
		self.clientId.setText("")
		self.clientIdDate.setText("")
		self.clientIdPiece.setText("")
		self.clientIdAt.setText("")
		self.clientIdBy.setText("")
		self.clientEnter.setText("")
		
		self.clientCivil.clear()
		self.clientCivil.addItem("Marié(e)")#civil combo	
		self.clientCivil.addItem("Célibataire")#civil combo	

		self.clientSexe.clear()
		self.clientSexe.addItem("Homme")	
		self.clientSexe.addItem("Femme")
		
		self.clientCompany.setText("")
		self.clientBirth.setText("")
		self.clientBirthPlace.setText("")
		self.clientAdress.setText("")
		self.clientJob.setText("")
		
		self.clientStay.setValue(0)#sejour spin
		self.clientKids.setValue(0)#kids spin

		self.clientPay.clear()
		self.clientPay.addItem("A la réservation")	
		self.clientPay.addItem("Au départ")

		self.clientPayMthd.clear()
		self.clientPayMthd.addItem("Cash")	
		self.clientPayMthd.addItem("Cheque")
		self.clientPayMthd.addItem("Carte de crédit")	

		self.clientDateIn.setText(self.Date)
		self.clientDateOut.setText("")
		
		self.viewTotal.display(0)

	def roomCount(self):	#INIT fiRSVieW
		#SET ProvidersLists AS QLISTWIDGET ITEMS :
		self.Edit.roomData()
		self.roomListWidget.clear()
		self.freeListWidget.clear()
		self.Edit.modifComboBox.clear()
		self.Edit.delComboBox.clear()
		self.Edit.modifComboBox.addItem("...")
		self.Edit.delComboBox.addItem("...")
		
		i=0
		for prov  in self.Edit.roomNameList :
			self.roomListWidget.addItem("Chambre : "+str(prov).strip("(',')"))
			
			self.roomNbr.setText(str(i+1))
			self.Edit.roomNbr.setText(str(i+1))
			self.Edit.modifComboBox.addItem(str(self.Edit.roomTypeList[i]).strip("(',')"))
			self.Edit.delComboBox.addItem(str(self.Edit.roomTypeList[i]).strip("(',')"))
			if str(self.Edit.roomStateList[i]).strip("(',')") == "0" :
				self.freeListWidget.addItem("Chambre "+str(prov).strip("(',')")+" : Libre")
			i+=1
		
	def clientSelector(self,select,ref):
		
		#SELECT Clients INTO LISTS:
		query.execute("SELECT * FROM Clients ORDER BY id")
		
		#SELECT Clients INTO LISTS:
		query.execute("SELECT "+str(select)+" FROM Clients WHERE "+str(ref)+" ORDER BY id DESC")
		self.selection = str(query.fetchone()).strip("(',')")
		
		#SELECT Clients INTO LISTS:
		query.execute("SELECT "+str(select)+" FROM Clients WHERE "+str(ref)+" ORDER BY id ASC")
		self.selectionAll = query.fetchall()		

	def roomSelector(self,select,ref):
		
		#SELECT Clients INTO LISTS:
		query.execute("SELECT * FROM Rooms ORDER BY id")
		
		#SELECT Clients INTO LISTS:
		query.execute("SELECT "+str(select)+" FROM Rooms WHERE "+str(ref)+" ORDER BY id DESC")
		self.selecter = str(query.fetchone()).strip("(',')")
		
		#SELECT Clients INTO LISTS:
		query.execute("SELECT "+str(select)+" FROM Rooms WHERE "+str(ref)+" ORDER BY id ASC")
		self.selecterAll = query.fetchall()
		
	def roomListSignal(self):
		
		self.roomListWidget.itemClicked.connect(self.listwidgetClickedSlot)

	def listwidgetClickedSlot(self, item):
		self.roomCleaner()
		i=0
		p=0
		for prov in self.Edit.roomNameList :
			if item.text() == "Chambre : "+str(prov).strip("(',')"):
		
				self.roomNum.setText(str(prov).strip("(',')"))
				self.roomType.setText(str(self.Edit.roomTypeList[i]).strip("(',')"))
				self.roomPrice.setText(str(self.Edit.roomPriceList[i]).strip("(',')"))
				if str(self.Edit.roomStateList[i]).strip("(',')") == '0' :
					self.roomState.setText("Libre")
					self.roomState.setStyleSheet("color :green;")
				else :
					self.roomState.setText("Occupée")
					self.roomState.setStyleSheet("color :red;")	
					
					self.clientSelector("client_name","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientName.setText(str(self.selection).strip("(',')"))
					
					self.clientSelector("client_national","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientNational.setText(str(self.selection).strip("(',')"))
					
					self.clientSelector("client_id","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientIdPiece.setText(str(self.selection).strip("(',')"))
					self.clientSelector("client_id_num","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientId.setText(str(self.selection).strip("(',')"))
					self.clientSelector("client_id_delivery","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientIdDate.setText(str(self.selection).strip("(',')"))
					self.clientSelector("client_id_at","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientIdAt.setText(str(self.selection).strip("(',')"))
					self.clientSelector("client_id_by","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientIdBy.setText(str(self.selection).strip("(',')"))
					
					self.clientSelector("client_enter","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientEnter.setText(str(self.selection).strip("(',')"))
					
					self.clientCivil.clear()
					self.clientSelector("client_civil","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientCivil.addItem(str(self.selection).strip("(',')"))
					
					self.clientSelector("client_kids","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientKids.setValue(int(str(self.selection).strip("(',')")))
					
					self.clientSexe.clear()
					self.clientSelector("client_sexe","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientSexe.addItem(str(self.selection).strip("(',')"))
					
					self.clientSelector("client_birth","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientBirth.setText(str(self.selection).strip("(',')"))#birth date
					self.clientSelector("client_birth_place","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientBirthPlace.setText(str(self.selection).strip("(',')"))#birth date
					
					self.clientSelector("client_adress","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientAdress.setText(str(self.selection).strip("(',')"))#birth date
					self.clientSelector("client_job","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientJob.setText(str(self.selection).strip("(',')"))#birth date
					
					self.clientSelector("client_company","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientCompany.setText(str(self.selection).strip("(',')"))
					
					self.clientSelector("client_sejour","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientStay.setValue(int(str(self.selection).strip("(',')")))
					
					self.clientSelector("client_date_in","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientDateIn.setText(str(self.selection).strip("(',')"))
					self.clientSelector("client_date_out","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientDateOut.setText(str(self.selection).strip("(',')"))
					
					self.clientPayMthd.clear()
					self.clientSelector("client_pay_mthd","client_room ='"+str(prov).strip("(',')")+"'")
					self.clientPayMthd.addItem(str(self.selection).strip("(',')"))
				
					self.clientSelector("client_total","client_room ='"+str(prov).strip("(',')")+"'")
					self.viewTotal.display(int(self.selection))
					
				self.roomIndicator.setText("Chambre : "+str(prov).strip("(',')"))
				self.roomIndic.setText(str(prov).strip("(',')"))
			i+=1

	def EditSignals(self):
		# DIALOG SIGNALS :
		self.Edit.Add.clicked.connect(self.EditSlot)
		
		self.Edit.Modif.clicked.connect(self.EditSlot)
		
		self.Edit.Del.clicked.connect(self.EditSlot)
		
		self.Edit.EditOK.clicked.connect(self.EditSlot)
		self.Edit.EditNO.clicked.connect(self.EditSlot)	
		
		self.Edit.modifComboBox.currentTextChanged.connect(self.comboSlots)	
		self.Edit.delComboBox.currentTextChanged.connect(self.comboSlots)	
		
		self.Edit.numModifComboBox.currentTextChanged.connect(self.comboSlots)	
		self.Edit.numDelComboBox.currentTextChanged.connect(self.comboSlots)	
		
	def comboSlots(self):
		self.Edit.roomData()
		
		if self.sender() == self.Edit.modifComboBox :
			query.execute("SELECT room_name FROM Rooms WHERE room_type ='"+self.Edit.modifComboBox.currentText()+"' ORDER BY id")
			rooms = query.fetchall()
			self.Edit.numModifComboBox.clear()
			self.Edit.numModifComboBox.addItem("...")
			i=0
			for room in rooms :
				self.Edit.numModifComboBox.addItem(str(room).strip("(',')"))
		
		if self.sender() == self.Edit.delComboBox :
			self.Edit.numDelComboBox.clear()
			self.Edit.numDelComboBox.addItem("...")
			query.execute("SELECT room_name FROM Rooms WHERE room_type ='"+self.Edit.delComboBox.currentText()+"' ORDER BY id")
			rooms = query.fetchall()
			for room in rooms :
				self.Edit.numDelComboBox.addItem(str(room).strip("(',')"))
				
		if self.sender() == self.Edit.numModifComboBox :
			i=0
			for room in self.Edit.roomNameList :
				if str(room).strip("(',')") == self.Edit.numModifComboBox.currentText() :
					self.Edit.modifName.setText(str(room).strip("(',')"))
					self.Edit.modifType.setText(str(self.Edit.roomTypeList[i]).strip("(',')"))
					self.Edit.modifPrice.setText(str(self.Edit.roomPriceList[i]).strip("(',')"))
				i+=1
				
		if self.sender() == self.Edit.numDelComboBox :
			for room in self.Edit.roomNameList :
				if str(room).strip("(',')") == self.Edit.numDelComboBox.currentText() :
					self.Edit.delName.setText(str(room).strip("(',')"))

	def Adder(self):
		try:
			if self.Edit.newRoomName.text() != "" and self.Edit.newRoomType.text() != "" and self.Edit.newRoomPrice.text() != ""  :
				query.execute("INSERT INTO Rooms (room_name,room_type,room_price,room_state) VALUES ('"+self.Edit.newRoomName.text()+\
				"','"+self.Edit.newRoomType.text()+"',"+self.Edit.newRoomPrice.text()+",0)")
				self.Edit.roomData()
				self.roomCount()
				conn.commit()
				
				self.Edit.newRoomName.setText("")
				self.Edit.newRoomType.setText("")
				self.Edit.newRoomPrice.setText('')
				 
				self.messageFactory.raiseAdder("Chambre")
			
			else :
				self.messageFactory.raiseCaseExcept("toutes les cases")
		except:
			self.messageFactory.raiseCharExcept()	

	def Modifer(self):
		try:
			if self.Edit.modifComboBox.currentText() != "..." and self.Edit.numModifComboBox.currentText() != "" and self.Edit.modifName.text()!="" \
			and self.Edit.modifType.text() != "" and self.Edit.modifPrice.text()!= "":
				query.execute("UPDATE Rooms SET room_name ='"+self.Edit.modifName.text()+"',room_type ='"+self.Edit.modifType.text()+"'\
				, room_price ="+self.Edit.modifPrice.text()+" WHERE room_name = '"+self.Edit.numModifComboBox.currentText()+"'")
				
				conn.commit()
				self.Edit.roomData()
				self.roomCount()
					
				self.Edit.modifName.setText("")
				self.Edit.modifType.setText("")
				self.Edit.modifPrice.setText("")
				self.Edit.numModifComboBox.clear()
				
				self.messageFactory.raiseModifier("Chambre")
			
			else :
				self.messageFactory.raiseCaseExcept("toutes les cases")
		except:
			self.messageFactory.raiseCharExcept()	

	def Deleter(self):
		try:
			if self.Edit.delComboBox.currentText() != "..." and self.Edit.numDelComboBox.currentText() != "" and self.Edit.delName.text()!="" :
				query.execute("DELETE FROM 'Rooms' WHERE room_name ='"+self.Edit.numDelComboBox.currentText()+"'")
				conn.commit()
				self.Edit.roomData()
				self.roomCount()
				self.messageFactory.raiseDeleter("Chambre")
				
				self.Edit.delName.setText("")
				self.Edit.numDelComboBox.clear()
			
			else :
				self.messageFactory.raiseCaseExcept("toutes les cases")
		except:
			self.messageFactory.raiseCharExcept()

	def EditSlot(self) :	
		if self.sender() == self.Edit.Add :
			self.Adder()
				
		if self.sender() == self.Edit.Modif :
			self.Modifer()
			
		if self.sender() == self.Edit.Del:
			self.Deleter()

		if self.sender() == self.Edit.EditOK or self.sender() == self.Edit.EditNO :
			self.Edit.destroy()
		 
#============================================================================================================DESIGN PATTERN EXCEPT FACTORY
class MessageFactory():

	def raiseSuccess(self,data):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("OPERATION REUSSIE !")
		msg.setText(data)
		msg.exec_()	

	def raiseAdder(self,data):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("OPERATION REUSSIE !")
		msg.setText(data+" ajouté(e) avec succés !")
		msg.exec_()	
		
	def raiseModifier(self,data):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("OPERATION REUSSIE !")
		msg.setText(data+" modifié(e) avec succés !")
		msg.exec_()	
		
	def raiseDeleter(self,data):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Information)
		msg.setWindowTitle("OPERATION REUSSIE !")
		msg.setText(data+" suprimé(e) avec succés !")
		msg.exec_()	
		
	def raiseCaseExcept(self,case):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("CASE OBLIGATOIRE VIDE !")
		msg.setText("Veuillez saisir "+case)
		msg.exec_()
		
	def raisePrintExcept(self,doc):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("ECHEC IMPRESSION !")
		msg.setText("Veuillez fermer le document word : '"+str(doc)+".docx'  !")
		msg.exec_()	
		
	def raiseCharExcept(self):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("CARACTERE INDESIRABLE !")
		msg.setText("Veuillez ne pas utiliser de caracteres spéciaux ?;:',.$*.... !")
		msg.exec_()
		
	def raiseIndefinedExcept(self,erreur):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("ERREUR !")
		msg.setText(erreur)
		msg.exec_()
		
	def raiseStockAlert(self,stock):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("ALERT STOCK !")
		msg.setText("Vérifié votre stock de "+stock)
		msg.exec_()
		
	def raiseAlert(self,alert):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)
		msg.setWindowTitle("ALERT !")
		msg.setText(alert)
		msg.exec_()

class ReferenceSelector():
	def __init__(self):
		self.refSelector()
		self.societySelctor()
	
	def refSelector(self):
		query.execute("SELECT ref FROM Ref WHERE id=1")
		self.ref = str(query.fetchone()).strip("(',')")	
		
	def societySelctor(self):
		query.execute("SELECT society_name FROM Society")
		self.society = query.fetchone() 
		self.society = str(self.society).strip("(',')") 
		query.execute("SELECT society_number FROM Society")
		self.numero = query.fetchone() 
		self.numero = str(self.numero).strip("(',')") 
		
#============================================================================================================PRODUCTWIDGET AND UTILITIE
# =============INITSUM DIALOG		
qtInitial= "DESIGN/DIALOGS/iniSum.ui"
Ui_initDialog, QtBaseClass = uic.loadUiType(qtInitial)

class InitDialog(QDialog, Ui_initDialog):#CONFIRM : VALID DIALOG ON VLDButton CLICK()ProdList
	def __init__(self):
		QDialog.__init__(self)
		Ui_initDialog.__init__(self)
		self.setupUi(self)
		
		#===========LOGICALS
		self.Date = time.strftime("%Y-%m-%d")
		
		self.messageFactory = MessageFactory()
		self.dataInit()
		
	def dataInit(self) :
		query.execute("SELECT register_sum_init FROM Register WHERE register_date = '"+self.Date+"'")
		self.sominit = query.fetchone()
		self.sominit = str(self.sominit).strip("(',')")
		self.sumInit.setText(self.sominit)
		
		self.INIT_OK.clicked.connect(self.slot)
		self.INIT_NO.clicked.connect(self.slot)
		
	def slot(self) :
	
		if self.sender() == self.INIT_OK and self.sumInit.text() != "":
			try:
				query.execute("UPDATE Register SET register_sum_init = '"+self.sumInit.text()+"' WHERE register_date = '"+self.Date+"'")
				query.execute("SELECT register_sum_init FROM Register WHERE register_date = '"+self.Date+"'")
				self.sominit = query.fetchone()
				self.sominit = str(self.sominit).strip("(',')")
				self.sumInit.setText(self.sominit)
				
				conn.commit()
					 
				self.messageFactory.raiseAdder("Somme Initiale")
				self.destroy()
			except:
				self.messageFactory.raiseIndefinedExcept("lors de l'initiation")
				self.messageFactory.raiseCharExcept()
		else :
			self.messageFactory.raiseCaseExcept("la case Somme initiale avec une valeur numerique !")
			self.messageFactory.raiseCharExcept()
		
		
		if self.sender() == self.INIT_NO :
			self.destroy()

#=============REGISTER DIALOG
qtRegisterDialog= "DESIGN/DIALOGS/registerCreatorDialog.ui"
Ui_registerCreatorDialog, QtBaseClass = uic.loadUiType(qtRegisterDialog)

class RegisterCreatorDialog(QDialog, Ui_registerCreatorDialog):# EDIT : MODIF Product Catégory DIALOG

	def __init__(self):
		QDialog.__init__(self)
		Ui_registerCreatorDialog.__init__(self)
		self.setupUi(self)

		self.setWindowTitle("Gestion de la Caisse et Des Sommes d'Argent")
		self.referenceSelector = ReferenceSelector()
		self.referenceSelector.refSelector()
		self.referenceSelector.societySelctor()
		
		#===========LOGICALS
		self.Date = time.strftime("%Y-%m-%d")
			
		self.localDateTime = time.strftime("%d-%m-%Y      %H:%M:%S")
		self.localDate = time.strftime("%d-%m-%Y")
		
		#==========INIT INSTANCES
		self.productDialog = ProductCreatorDialog()
		self.registerInit = InitDialog()
		self.messageFactory = MessageFactory()
		
		self.registerData()
		self.registerSignal()
		
		#===========METHODES
	def registerData(self) :
		
		query.execute("SELECT register_date FROM Register ORDER BY id DESC")
		self.lastDate = query.fetchall()
		
		if self.lastDate == [] :
			query.execute("INSERT INTO Register (register_date, register_sum_init, register_recette_total,register_depense_total,register_ajout_total) VALUES ('"\
			+self.Date+"',0,0,0,0)")
			self.recette = 0
			
		elif str(self.lastDate[0]).strip("(',')") != self.Date :
			query.execute("INSERT INTO Register (register_date, register_sum_init, register_recette_total,register_depense_total,register_ajout_total) VALUES ('"\
			+self.Date+"',0,0,0,0)")
			self.recette = 0
			
		query.execute("SELECT register_recette_total FROM Register WHERE register_date = '"+self.Date+"'")
		self.recette = str(query.fetchone()).strip("(',')")
		self.recette = int(self.recette)
		
		# query.execute("SELECT register_sum_init FROM Register WHERE register_date = '"+self.Date+"'")
		query.execute("SELECT register_sum_init FROM Register ORDER BY id DESC")

		query.execute("SELECT register_sum_init FROM Register ORDER BY id DESC")
		self.suminit = str(query.fetchone()).strip("(',')")
		
		query.execute("SELECT register_depense_total FROM Register WHERE register_date = '"+self.Date+"'")
		self.depense = str(query.fetchone()).strip("(',')")
		
		query.execute("SELECT register_ajout_total FROM Register WHERE register_date = '"+self.Date+"'")
		self.ajout = str(query.fetchone()).strip("(',')")
		
		self.sumInit.setText(self.suminit)
		self.sumDay.setText(str(self.recette))
		self.sumDep.setText(self.depense)
		self.sumAdd.setText(str(self.ajout))
		
		self.sumTotal.display(0)
		conn.commit()

	def registerSignal(self):
	
		self.totalButton.clicked.connect(self.registerSlot)
		self.depAdd.clicked.connect(self.registerSlot)
		self.addAdd.clicked.connect(self.registerSlot)
		
		self.registerAnnuler.clicked.connect(self.registerSlot)
		self.registerPrint.clicked.connect(self.registerSlot)
		self.editInit.clicked.connect(self.registerSlot)
		
		self.registerInit.INIT_OK.clicked.connect(self.totalRegister)

	def totalRegister(self):
		try:
			self.regTotal = []
			query.execute("SELECT register_recette_total FROM Register WHERE register_date = '"+self.Date+"'")
			self.recette = str(query.fetchone()).strip("(',')")	
			self.regTotal.append(int(self.recette))
			
			query.execute("SELECT register_sum_init FROM Register ORDER BY id DESC")
			self.suminit = str(query.fetchone()).strip("(',')")
			self.regTotal.append(int(self.suminit))
			
			query.execute("SELECT register_depense_total FROM Register WHERE register_date = '"+self.Date+"'")
			self.depense = str(query.fetchone()).strip("(',')")
			
			query.execute("SELECT register_ajout_total FROM Register WHERE register_date = '"+self.Date+"'")
			self.ajout = str(query.fetchone()).strip("(',')")
			self.regTotal.append(int(self.ajout))
			
			self.sumInit.setText(self.suminit)
			self.sumDay.setText(str(self.recette))
			self.sumDep.setText(self.depense)
			self.sumAdd.setText(str(self.ajout))
			
			self.sumTotal.display(sum(self.regTotal) - int(self.depense))
			self.actual = str(sum(self.regTotal) - int(self.depense))
			self.actual = str(self.actual)	
		except:
			self.messageFactory.raiseIndefinedExcept("au niveau de la caisse")
		
	def registerSlot(self) :	
		self.regTotal = []
		
		if self.sender() == self.totalButton :
			self.totalRegister()
		
		if self.sender() == self.registerAnnuler :
			self.destroy()
		
		if self.sender() == self.depAdd :
			self.newDepense()
			
		if self.sender() == self.editInit :
			self.totalRegister()
			self.registerInit.show()
			
		if self.sender() == self.addAdd :
			self.newDepot()
		
		if self.sender() == self.registerPrint :
			self.totalRegister()
			self.printRegister()

	def newDepot(self):
		try:
			if self.addVal.text() != '0' and self.addName.text() != "" :
				query.execute("INSERT INTO Register_add (add_date, add_depo, add_description,add_value) VALUES ('"\
				+self.Date+"','"+str(self.addName.text())+"','"+str(self.addDesc.text())+"','"+str(self.addVal.text())+"')")
				conn.commit()
				 
				msg = QMessageBox()
				msg.setIcon(QMessageBox.Information)

				msg.setText("Ajout Enregistrée avec succés")
				msg.setWindowTitle("Edition des Ajout !")
				msg.exec_()
		
				self.addName.setText("")
				self.addDesc.setText("")
				self.addVal.setText("0")
		
				query.execute("SELECT add_value FROM Register_add WHERE add_date = '"+self.Date+"'")
				self.addList = []
				self.addList = query.fetchall()
				self.add = []
				i=0
				for add in self.addList :	
					self.add.append(str(add).strip("(',')"))
					self.add[i] = int(self.add[i])
					i+=1
				query.execute("UPDATE Register SET register_ajout_total = "+str(sum(self.add))+" WHERE register_date ='"+self.Date+"'")
				
				query.execute("SELECT * FROM Register")
				self.regTotal.append(sum(self.add))
				conn.commit()
				self.totalRegister()
				
			else : 
				self.messageFactory.raiseCaseExcept("Toutes les cases")
				if  self.addName.text() == '':
					self.messageFactory.raiseCaseExcept("la case Dépositaire")
				if  self.addVal.text() == '0':
					self.messageFactory.raiseCaseExcept("la case Somme à dépenser")
		except:
			self.messageFactory.raiseCharExcept()

	def actualDepense(self):	
		query.execute("SELECT dep_value FROM Register_dep WHERE dep_date = '"+self.Date+"'")
		self.depensesList = []
		self.depensesList = query.fetchall()
		self.dep = []
		i=0
		for dep in self.depensesList :	
			self.dep.append(str(dep).strip("(',')"))
			self.dep[i] = int(self.dep[i])
			i+=1
		query.execute("UPDATE Register SET register_depense_total = "+str(sum(self.dep))+" WHERE register_date ='"+self.Date+"'")
		
		query.execute("SELECT register_depense_total FROM Register WHERE register_date = '"+self.Date+"'")
		self.depense = str(query.fetchone()).strip("(',')")
		self.depense = int(self.depense)
		conn.commit()
		self.totalRegister()
			
	def newDepense(self):
		try:
			if self.depVal.text() != '0' and self.depDesc.text() != "" :
				query.execute("INSERT INTO Register_dep (dep_date, dep_type, dep_description,dep_value) VALUES ('"\
				+self.Date+"','"+self.depTypeCmbBox.currentText()+"','"+self.depDesc.text()+"','"+self.depVal.text()+"')")
				query.execute("SELECT * FROM Register_dep")
				conn.commit()
				 
				
				self.messageFactory.raiseAdder("Dépense")
		
				self.depDesc.setText("")
				self.depVal.setText("0")
				self.totalRegister()
				
				self.actualDepense()
				
			else : 
				self.messageFactory.raiseCaseExcept("Toutes les cases")
				if  self.depDesc.text() == '':
					self.messageFactory.raiseCaseExcept("la case déscription")
				if  self.depVal.text() == '0':
					self.messageFactory.raiseCaseExcept("la case Somme à déposer")
		except:
			self.messageFactory.raiseCharExcept()

	def	printRegister(self) :
		try :			
			TICKET = Document()
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")

			h = TICKET.add_heading("                        "+self.referenceSelector.society+" : 'Mouvement de Caisse'", level=1)
			h.bold = True
			h.italic = True
			p = TICKET.add_paragraph("				             "+self.Date)
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_heading("                                     Somme Actuelle en Caisse : "+str(self.actual)+" DA ", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			#DEPENSES
			p = TICKET.add_heading("DEPENSES : ", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			query.execute("SELECT dep_type FROM Register_dep WHERE dep_date = '"+self.Date+"' ORDER BY id")
			self.typeList = []
			self.typeList = query.fetchall()
			query.execute("SELECT dep_description FROM Register_dep WHERE dep_date = '"+self.Date+"' ORDER BY id")
			self.descList = []
			self.descList = query.fetchall()
			query.execute("SELECT dep_value FROM Register_dep WHERE dep_date = '"+self.Date+"' ORDER BY id")
			self.valList = []
			self.valList = query.fetchall()
			
			pi = 0 
			tab = TICKET.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Type'
			heading_cells[1].text = 'Description'
			heading_cells[2].text = 'Valeur'
			heading_cells[3].text = 'Date'
			pi =0
			for dep in self.typeList :
				cells = tab.add_row().cells
				cells[0].text = str(self.typeList[pi]).strip("(',')")
				cells[1].text = str(self.descList[pi]).strip("(',')")
				cells[2].text = str(self.valList[pi]).strip("(',')") + " DA"
				cells[3].text = self.Date
				pi+=1
				
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			#AJOUTS		
			p = TICKET.add_heading("DEPOTS : ", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			query.execute("SELECT add_depo FROM Register_add WHERE add_date = '"+self.Date+"' ORDER BY id")
			self.typeList = []
			self.typeList = query.fetchall()
			query.execute("SELECT add_description FROM Register_add WHERE add_date = '"+self.Date+"' ORDER BY id")
			self.descList = []
			self.descList = query.fetchall()
			query.execute("SELECT add_value FROM Register_add WHERE add_date = '"+self.Date+"' ORDER BY id")
			self.valList = []
			self.valList = query.fetchall()
			
			pi = 0 
			tab = TICKET.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Type'
			heading_cells[1].text = 'Description'
			heading_cells[2].text = 'Valeur'
			heading_cells[3].text = 'Date'
			pi =0
			for dep in self.typeList :
				cells = tab.add_row().cells
				cells[0].text = str(self.typeList[pi]).strip("(',')")
				cells[1].text = str(self.descList[pi]).strip("(',')")
				cells[2].text = str(self.valList[pi]).strip("(',')") + " DA"
				cells[3].text = self.Date
				pi+=1
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph("SOMME INITIALE EN CAISSE : " +self.suminit+ " DA")
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph("RECETTE DU JOUR  : " +self.recette+ " DA")
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_heading("\t\tReference No : "+str(self.referenceSelector.ref)+". --"+self.referenceSelector.society+"-- "+self.localDateTime, level=3)
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			TICKET.save('DOCUMENTS/CAISSE/Caisse'+str(self.referenceSelector.ref)+"_"+self.localDate+'.docx' )
			os.startfile('DOCUMENTS\CAISSE\Caisse'+str(self.referenceSelector.ref)+"_"+self.localDate+'.docx' , 'print')
			
			conn.commit()
			
		except :
			self.messageFactory.raisePrintExcept("Caisse"+str(self.ref)+"_"+self.localDate)

#==================================================	VALID PROD DIALOG	
	
qtValid= "DESIGN/DIALOGS/VLDDialog.ui"
Ui_validDialog, QtBaseClass = uic.loadUiType(qtValid)

class ValidProdsDialog(QDialog, Ui_validDialog):#CONFIRM : VALID DIALOG ON VLDButton CLICK()
	def __init__(self):
		QDialog.__init__(self)
		Ui_validDialog.__init__(self)
		self.setupUi(self)
		
#=============CATEGORY DIALOG
qtCreatorDialog= "DESIGN/DIALOGS/categoryCreatorDialog.ui"
Ui_categoryCreatorDialog, QtBaseClass = uic.loadUiType(qtCreatorDialog)

class CategoryCreatorDialog(QDialog, Ui_categoryCreatorDialog):# EDIT : MODIF Product Catégory DIALOG

	def __init__(self):
		QDialog.__init__(self)
		Ui_categoryCreatorDialog.__init__(self)
		self.setupUi(self)

		# self.prodWidget = ProductWidget()	
		self.setWindowTitle("Edition du menu et sous-catégories")
		# SETTING TEXT  category_name FROM CATEGORIES -->  CATLIST
		
		self.messageFactory = MessageFactory()
		
		#GETTING CATEGORY NAMES FROM CATEGORY QDIALOG INTO LISTS
		#CATEGORY PRODUCTS NAMES LIST
		self.categoryName = [self.categoryName1, self.categoryName2, self.categoryName3,
		self.categoryName4,  self.categoryName5,  self.categoryName6, self.categoryName7,
		self.categoryName8,  self.categoryName9, self.categoryName10]
  
		#CATEGORY PRODUCTS SUB NAMES 1 LIST
		self.categorySub1 = [self.cat1subCat1, self.cat2subCat1, self.cat3subCat1,
		self.cat4subCat1, self.cat5subCat1, self.cat6subCat1, self.cat7subCat1, 
		self.cat8subCat1, self.cat9subCat1, self.cat10subCat1]

		#CATEGORY PRODUCTS SUB NAMES 2 LIST
		self.categorySub2 = [self.cat1subCat2, self.cat2subCat2, self.cat3subCat2,
		self.cat4subCat2, self.cat5subCat2, self.cat6subCat2, self.cat7subCat2, 
		self.cat8subCat2, self.cat9subCat2, self.cat10subCat2 ]


		#CATEGORY PRODUCTS SUB NAMES 3 LIST
		self.categorySub3 = [self.cat1subCat3, self.cat2subCat3, self.cat3subCat3,
		self.cat4subCat3, self.cat5subCat3, self.cat6subCat3, self.cat7subCat3, 
		self.cat8subCat3, self.cat9subCat3, self.cat10subCat3]	
				
		# SELECTING CATEGORIES FROM DATA --> INTO LIST 
		
		# SELECTING  category_name FROM CATEGORIES -->  CATLIST:
		query.execute("SELECT category_name FROM Categories ORDER by id ASC ")#SELECT : Category Name
		self.catList = list()
		self.catList = query.fetchall()
		i=0
		while i <= 9 :
			self.categoryName[i].setText(str(self.catList[i]).strip("(',')"))
			self.categorytabWidget.setTabText(i,str(self.catList[i]).strip("(',')"))#TAB 1 SETTEXT
			i+=1
	
		
		query.execute("SELECT category_sub_1 FROM Categories ORDER by id ASC ")#SELECT : Category sub 1
		self.subList1 = list()
		self.subList1 = query.fetchall()
		i=0
		while i <= 9 :
			if str(self.subList1[i]).strip("(',')") != "None" :
				self.categorySub1[i].setText(str(self.subList1[i]).strip("(',')"))
			i+=1

			
		query.execute("SELECT category_sub_2 FROM Categories ORDER by id ASC ")#SELECT : Category sub 2
		self.subList2 = list()
		self.subList2 = query.fetchall()
		i=0
		while i <= 9 :
			if str(self.subList2[i]).strip("(',')") != "None" :
				self.categorySub2[i].setText(str(self.subList2[i]).strip("(',')"))
			i+=1

			
		query.execute("SELECT category_sub_3 FROM Categories ORDER by id ASC ")#SELECT : Category sub 3 tabWidget
		self.subList3 = list()
		self.subList3 = query.fetchall()
		i=0
		while i <= 9 :
			if str(self.subList3[i]).strip("(',')") != "None" :
				self.categorySub3[i].setText(str(self.subList3[i]).strip("(',')"))
			i+=1
			
		#modif products cat_name
		
	def AnnulerNewCategorySlot(self):#SELF.DIALOG.DESTROY
		self.destroy()
		
	def validNewCategorySlot(self):#UPDATE * SET Categories Name and Subs BY id, UPDATE SELF.CATNAMES, TABSNAMES in DIALOG AND MAIN
		
		#SELECTING PRODUCT CATEGORY NAMES FROM PRODUCT AND UPDATE FROM CATLIST
		query.execute("SELECT product_category_name FROM Products ORDER BY id")
		prodcatlist = query.fetchall()
		query.execute("SELECT product_subCategory FROM Products ORDER BY id")
		prodsublist = query.fetchall()
		try :
			c=0
			while c <= 9 :
			
				#CATEGORY  UPDATING
				query.execute("UPDATE Categories SET category_name = '"+self.categoryName[c].text()+"' WHERE id = "+str(c+1)+"")
						
				query.execute("UPDATE Categories SET category_sub_1 = '"+self.categorySub1[c].text()+"' WHERE id = "+str(c+1)+"")
					
				query.execute("UPDATE Categories SET category_sub_2 = '"+self.categorySub2[c].text()+"' WHERE id = "+str(c+1)+"")

				query.execute("UPDATE Categories SET category_sub_3 = '"+self.categorySub3[c].text()+"' WHERE id = "+str(c+1)+"") 

			
				#UPDATING PRODUCTS CATEGORY NAMES WITH NEW DATA
				
				for prod in prodcatlist:
					if str(prod).strip("(',')") == str(self.catList[c]).strip("(',')") :
						query.execute("UPDATE Products SET product_category_name = '"+self.categoryName[c].text()+"' WHERE product_category_name ='"+str(self.catList[c]).strip("(',')")+"'")
				
				for prod in prodsublist:
					if str(prod).strip("(',')") == str(self.subList1[c]).strip("(',')") :
						query.execute("UPDATE Products SET product_subCategory = '"+self.categorySub1[c].text()+"' WHERE product_subCategory ='"+str(self.subList1[c]).strip("(',')")+"'")
						query.execute("SELECT * FROM Products WHERE product_subCategory = '"+self.categorySub1[c].text()+"'")
						
					if str(prod).strip("(',')") == str(self.subList2[c]).strip("(',')") :
						query.execute("UPDATE Products SET product_subCategory = '"+self.categorySub2[c].text()+"' WHERE product_subCategory ='"+str(self.subList2[c]).strip("(',')")+"'")
						query.execute("SELECT * FROM Products WHERE product_subCategory = '"+self.categorySub2[c].text()+"'")
						
					if str(prod).strip("(',')") == str(self.subList3[c]).strip("(',')") :
						query.execute("UPDATE Products SET product_subCategory = '"+self.categorySub3[c].text()+"' WHERE product_subCategory ='"+str(self.subList3[c]).strip("(',')")+"'")
						query.execute("SELECT * FROM Products WHERE product_subCategory = '"+self.categorySub3[c].text()+"'")
			
				c+=1
			
			# SELECTING CATEGORIES FROM DATA --> INTO LIST 
			
			# SELECTING  category_name FROM CATEGORIES -->  CATLIST:
			query.execute("SELECT category_name FROM Categories ORDER by id ASC ")#SELECT : Category Name
			self.catList = list()
			self.catList = query.fetchall()
			i=0
			while i <= 9 :
				self.categoryName[i].setText(str(self.catList[i]).strip("(',')"))
				self.categorytabWidget.setTabText(i,str(self.catList[i]).strip("(',')"))#TAB 1 SETTEXT
				i+=1

			
			query.execute("SELECT category_sub_1 FROM Categories ORDER by id ASC ")#SELECT : Category sub 1
			self.subList1 = list()
			self.subList1 = query.fetchall()
			i=0
			while i <= 9 :
				if str(self.subList1[i]).strip("(',')") != "None" :
					self.categorySub1[i].setText(str(self.subList1[i]).strip("(',')"))
				i+=1

				
			query.execute("SELECT category_sub_2 FROM Categories ORDER by id ASC ")#SELECT : Category sub 2
			self.subList2 = list()
			self.subList2 = query.fetchall()
			i=0
			while i <= 9 :
				if str(self.subList2[i]).strip("(',')") != "None" :
					self.categorySub2[i].setText(str(self.subList2[i]).strip("(',')"))
				i+=1

				
			query.execute("SELECT category_sub_3 FROM Categories ORDER by id ASC ")#SELECT : Category sub 3
			self.subList3 = list()
			self.subList3 = query.fetchall()
			i=0
			while i <= 9 :
				if str(self.subList3[i]).strip("(',')") != "None" :
					self.categorySub3[i].setText(str(self.subList3[i]).strip("(',')"))	
				i+=1

			conn.commit()
		except:
			self.messageFactory.raiseCharExcept()

#=============PRODUCT CREATOR DIALOG

qtProductCreator= "DESIGN\DIALOGS\productCreatorDialog.ui"
Ui_productCreatorDialog, QtBaseClass = uic.loadUiType(qtProductCreator)

class ProductCreatorDialog(QDialog, Ui_productCreatorDialog):#EDIT : MODIF Product Name,Price DIALOG
	
	def __init__(self):
		QDialog.__init__(self)
		Ui_productCreatorDialog.__init__(self)
		self.setupUi(self)
		
		self.setWindowTitle("Edition des Produits")
		
		self.prodAnnuler.clicked.connect(self.prodEditAnnuler)
		
		#======INSTANCES
		self.categoryDialog = CategoryCreatorDialog()#EDIT : CATEGORY CHANGES DIALOG 
		self.prodDataList()
		#======SLOT VALiD
		
		self.prodOK.clicked.connect(self.prodEditValid)

		# """ RECUPERATION DE LA CLASSE VIA LISTES"""
		query.execute("SELECT category_name FROM Categories ORDER BY id")
		self.categoryList = list()
		self.categoryList = query.fetchall()

	def prodDataList(self):# i=0
		
		query.execute("SELECT product_name FROM Products ORDER BY id ASC")
		self.prodNameList = list()
		self.prodNameList = query.fetchall()
		
		query.execute("SELECT product_price FROM Products ORDER BY id ASC")
		self.prodPriceList = list()
		self.prodPriceList = query.fetchall()
		
		query.execute("SELECT product_qnt FROM Products ORDER BY id ASC")
		self.prodQntList = list()
		self.prodQntList = query.fetchall()
		
		query.execute("SELECT product_comptable FROM Products ORDER BY id ASC")
		self.prodComptList = list()
		self.prodComptList = query.fetchall()
		
		query.execute("SELECT product_stock FROM Products ORDER BY id ASC")
		self.prodStockableList = list()
		self.prodStockableList = query.fetchall()
		# print(query.fetchall())
		
		query.execute("SELECT product_alert FROM Products ORDER BY id ASC")
		self.prodAlertList = list()
		self.prodAlertList = query.fetchall()
		# print(query.fetchall())
		
		query.execute("SELECT product_stockable FROM Products ORDER BY id ASC")
		self.prodStockList = list()
		self.prodStockList = query.fetchall()
		
		query.execute("SELECT product_total FROM Products ORDER BY id ASC")
		self.prodTotalList = list()
		self.prodTotalList = query.fetchall()
		
	def prodEditValid(self):	            		
		self.close()

	def prodEditAnnuler(self): 
		self.close()	
		
#========== PRODUCT WIDGET UIC CONVERT & Load

qtProductWidget = "DESIGN/WIDGETS/ProductWidget.ui" # Enter file here.

Ui_ProductWidget, QtBaseClass = uic.loadUiType(qtProductWidget)

#==========INIT CLASS
#============================================================================================================PRODUCT WIDGET 

class ProductWidget(QWidget, Ui_ProductWidget):

	def __init__(self):
	
		QWidget.__init__(self)#QWidget init
		Ui_ProductWidget.__init__(self)#ui init
		self.setupUi(self)#UI init
				
		#==========Logical & Réusables Variables
		self.referenceSelector = ReferenceSelector()
			
		self.localDateTime = time.strftime("%d-%m-%Y      %H:%M:%S")
		self.localDate = time.strftime("%d-%m-%Y")
		self.Date = time.strftime("%Y-%m-%d")
		
		self.society = self.referenceSelector.society 
		self.numero = self.referenceSelector.numero  
		
		query.execute("SELECT user FROM User WHERE id = 1")
		self.user = query.fetchone()
		self.user= str(self.user).strip("(',')")
				
		self.ref = self.referenceSelector.ref
		self.viewRefNum.setText(str(self.ref))#Reference of sold TICKETS 
		
		self.prodId = 0

		#==========INIT CLASSES OF TYPE DIALOGS FROM MODULE "CreatorDialog":
		
		self.categoryDialog = CategoryCreatorDialog()#EDIT : CATEGORY CHANGES DIALOG
		
		self.productDialog = ProductCreatorDialog()#EDIT : CHANGE Prod Name,Price DIALOG
		
		self.validVLD = ValidProdsDialog()#CONFIRM : TOTAL AND VALID DIALOG
		
		self.register = RegisterCreatorDialog()
		
		self.messageFactory = MessageFactory()
		
		#==========DELEGATES INSTANCES :
		
		
		self.categoryDataManager()#DATAMANAGER : Categories.db
		
		self.categorySignalClicked()#SIGNAL : EMIT FOR CATEGORY WHEN CLICKED
		
		self.categorySlotClicked()#SLOT : Work ON Category WHEN CLICKED
		
		# self.productsDataManager()#DATAMANAGER : Products.db
		
		self.productsSignalClicked()#SIGNAL : EMIT wheN PRODUCT IS CLICKED
		
		self.validationButtonsSignalClicked()#SIGNAL : EMIT WhEN VALIDATIONBUTTON IS CLICKED
		
		self.tableData()
		
#=========================================	tabWidget		
	def categoryDataManager(self):	#SELECT Categories.names|subs --> catList|subListi --> menuprod_catName, tabs, dialog(catName, subicatsub)"""
									#UPDATE prod_Category FROM catList by id """
		
		self.menuProd_categoryList=[self.menuProd_category_1, self.menuProd_category_2, self.menuProd_category_3,
				self.menuProd_category_4, self.menuProd_category_5, self.menuProd_category_6, self.menuProd_category_7,
				self.menuProd_category_8, self.menuProd_category_9, self.menuProd_category_10]
#CREATE A prod_List[prod_buttons]
		self.prod_List = list()
		
		self.prod_List = [self.prod_1, self.prod_2, self.prod_3, self.prod_4, self.prod_5, self.prod_6, 
		self.prod_7, self.prod_8, self.prod_9, self.prod_10, self.prod_11, self.prod_12, self.prod_13, 
		self.prod_14, self.prod_15, self.prod_16, self.prod_17, self.prod_18, self.prod_19, self.prod_20, 
		self.prod_21, self.prod_22, self.prod_23, self.prod_24, self.prod_25, self.prod_26, self.prod_27, 
		self.prod_28, self.prod_29, self.prod_30, self.prod_31, self.prod_32, self.prod_33, self.prod_34, 
		self.prod_35, self.prod_36, self.prod_37, self.prod_38, self.prod_39, self.prod_40, self.prod_41, 
		self.prod_42, self.prod_43, self.prod_44, self.prod_45, self.prod_46, self.prod_47, self.prod_48, 
		self.prod_49, self.prod_50, self.prod_51, self.prod_52, self.prod_53, self.prod_54, self.prod_55, 
		self.prod_56, self.prod_57, self.prod_58, self.prod_59, self.prod_60, self.prod_61, self.prod_62, 
		self.prod_63, self.prod_64, self.prod_65, self.prod_66, self.prod_67, self.prod_68, self.prod_69, 
		self.prod_70, self.prod_71, self.prod_72, self.prod_73, self.prod_74, self.prod_75]	
		c=0
		while c <= 9:
		
			self.categoryDialog.catList[c] = self.categoryDialog.categoryName[c].text()
			self.menuProd_categoryList[c].setText(str(self.categoryDialog.catList[c]).strip("(',')"))
			self.menuProd_categoryList[c].setToolTip(str(self.categoryDialog.catList[c]).strip("(',')"))
			self.menuProd_categoryList[c].setStatusTip(str(self.categoryDialog.catList[c]).strip("(',')")+"...")
			c+=1	
			
	def categorieCreatorSignal(self):#SIGNAL : CALL DIALOG ui
		self.categoryDialog.catOK.clicked.connect(self.categoryDialog.validNewCategorySlot)#SIGNAL : VALID CHANGES ON CATEGORIES and SUbCategories	
		self.categoryDialog.catOK.clicked.connect(self.validNewCategorySlot)#SIGNAL : VALID CHANGES ON CATEGORIES and SUbCategories
		self.categoryDialog.catAnnuler.clicked.connect(self.categoryDialog.AnnulerNewCategorySlot)#Valid and Save changes Function
		self.categoryDialog.show()
		
	def validNewCategorySlot(self):#UPDATE * SET Categories Name and Subs BY id, UPDATE SELF.CATNAMES, TABSNAMES in DIALOG AND MAIN
						
		c=0
		while c <= 9 :
		
			#CATEGORY  UPDATING
			
			self.menuProd_categoryList[c].setText(self.categoryDialog.categoryName[c].text())#Menu Category 1 SETTEXT
			self.categoryDialog.catList[c] = self.categoryDialog.categoryName[c].text()
			
			self.categoryDialog.subList1[c] = self.categoryDialog.categorySub1[c].text()
			self.categoryDialog.categorySub1[c].setText(str(self.categoryDialog.subList1[c]).strip("(',')"))#Dialog sub1 SETTEXT
		
			self.categoryDialog.subList2[c] = self.categoryDialog.categorySub2[c].text()
			self.categoryDialog.categorySub2[c].setText(str(self.categoryDialog.subList2[c]).strip("(',')"))#Dialog sub1 SETTEXT
		
			self.categoryDialog.subList3[c] = self.categoryDialog.categorySub3[c].text()
			self.categoryDialog.categorySub3[c].setText(str(self.categoryDialog.subList3[c]).strip("(',')"))#Dialog sub1 SETTEXT  
			
			c+=1

		self.categoryDialog.destroy()
					
	def categorySignalClicked(self): # CONNECT CATEGORIES BUTTONS SIGNAL
	
		self.menuProd_addCat.clicked.connect(self.categorieCreatorSignal)#SIGNAL : CATEGORY DIALOG
		
		#SIGNAL : MENU CATEGORY CLICKED
		i = 0
		while i<=9:
			self.menuProd_categoryList[i].clicked.connect(self.categorySlotClicked)
			i+=1
		
	def categorySlotClicked(self):#SLOT : SET TO TABWIDGET : TabText, VALUES TO BUTTONS, FROM DB , TEXT FROM DIALOG to catNames
				
		pd=0
		while pd <= 74 :		
			self.prod_List[pd].setText("")
			pd+=1
		
		c = 0		
		while c <= 9 :
			if self.sender() == self.menuProd_categoryList[c] :
				self.categoryIndicator.setText(str(self.categoryDialog.catList[c]).strip("(',')"))
				
				if str(self.categoryDialog.subList1[c]).strip("(',')") != "None":
					self.tabWidget.setTabText(0,str(self.categoryDialog.subList1[c]).strip("(',')"))#TAB 1 SETTEXT
				else :
					self.tabWidget.setTabText(0,"")
					
				if str(self.categoryDialog.subList2[c]).strip("(',')") != "None":
					self.tabWidget.setTabText(1,str(self.categoryDialog.subList2[c]).strip("(',')"))#TAB 2 SETTEXT
				else :
					self.tabWidget.setTabText(1,"")
					
				if str(self.categoryDialog.subList3[c]).strip("(',')") != "None":
					self.tabWidget.setTabText(2,str(self.categoryDialog.subList3[c]).strip("(',')"))#TAB 3 SETTEXT
				else :
					self.tabWidget.setTabText(2,"")
			
				query.execute("SELECT product_name FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList1[c]).strip("(',')")+"'")
				subListS1 = list()
				subListS1 = query.fetchall()
				query.execute("SELECT product_name FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList2[c]).strip("(',')")+"'")
				subListS2 = list()
				subListS2 = query.fetchall()
				query.execute("SELECT product_name FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList3[c]).strip("(',')")+"'")
				subListS3 = list()
				subListS3 = query.fetchall()
			
				query.execute("SELECT product_price FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList1[c]).strip("(',')")+"'")
				subListP1 = list()
				subListP1 = query.fetchall()
				query.execute("SELECT product_price FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList2[c]).strip("(',')")+"'")
				subListP2 = list()
				subListP2 = query.fetchall()
				query.execute("SELECT product_price FROM Products WHERE product_subCategory = '"+str(self.categoryDialog.subList3[c]).strip("(',')")+"'")
				subListP3 = list()
				subListP3 = query.fetchall()
				
				p1=0
				p2=25
				p3=50
				pixmap ="IMAGES\Icons\chef-hat.png"
				for prod in subListS1 :					
					self.prod_List[p1].setText(str(prod).strip("(',')"))
					self.prod_List[p1].setToolTip(str(prod).strip("(',')")+"\n"+str(subListP1[p1]).strip("(',')")+" DA") 
					if self.prod_List[p1].text() != "" :
						self.prod_List[p1].setStyleSheet("QPushButton{background-color:darkgrey ; color :white ; border: 0px solid #2f4f4f  ; border-radius: 0.5px; height:35px;}\
						\n QPushButton:hover{color:white;background-color:#f96332;}\
						\nQPushButton:focus{border:0px solid #f96332;color:#f96332;background-color:#ecf0f5;}")
						
						self.prod_List[p1].setIcon(QIcon(".\IMAGES\Icons\chef-hat.png"))
					else:
						self.prod_List[p1].setStyleSheet("QPushButton{border:0px solid #f96332;color:#f96332;background-color:#ecf0f5;}")
					p1+=1
				p=0
				for prod in subListS2 :		
					self.prod_List[p2].setText(str(prod).strip("(',')"))
					if self.prod_List[p2].text() != "" :
						self.prod_List[p2].setStyleSheet("QPushButton{background-color:darkgrey ; color :white ; border: 0px solid #2f4f4f  ; border-radius: 0.5px; height:35px;}\
						\n QPushButton:hover{color:white;background-color:#f96332;}\
						\nQPushButton:focus{border:0px solid #f96332;color:#f96332;background-color:#ecf0f5;}")
					else:
						self.prod_List[p2].setStyleSheet("QPushButton{border:0px solid #f96332;color:#f96332;background-color:#ecf0f5;}")
					self.prod_List[p2].setToolTip(str(prod).strip("(',')")+"\n"+str(subListP2[p]).strip("(',')")+" DA") 
					p2+=1
					p+=1
				p=0
				for prod in subListS3 :	
					self.prod_List[p3].setText(str(prod).strip("(',')"))
					if self.prod_List[p3].text() != "" :
						self.prod_List[p3].setStyleSheet("QPushButton{background-color:darkgrey ; color :white ; border: 0px solid #2f4f4f  ; border-radius: 0.5px; height:35px;}\
						\n QPushButton:hover{color:white;background-color:#f96332;}\
						\nQPushButton:focus{border:0px solid #f96332;color:#f96332;background-color:#ecf0f5;}")
					else:
						self.prod_List[p3].setStyleSheet("QPushButton{border:0px solid #f96332;color:#f96332;background-color:#ecf0f5;}")
					self.prod_List[p3].setToolTip(str(prod).strip("(',')")+"\n"+str(subListP3[p]).strip("(',')")+" DA") 
					p3+=1
					p+=1
		
			c+=1
			
		
		self.prodDesign()

	def prodDesign(self):
		
		self.prod_ListText = [self.prod_1.text(), self.prod_2.text(), self.prod_3.text(), self.prod_4.text(), self.prod_5.text(), self.prod_6.text(), 
		self.prod_7.text(), self.prod_8.text(), self.prod_9.text(), self.prod_10.text(), self.prod_11.text(), self.prod_12.text(), self.prod_13.text(), 
		self.prod_14.text(), self.prod_15.text(), self.prod_16.text(), self.prod_17.text(), self.prod_18.text(), self.prod_19.text(), self.prod_20.text(), 
		self.prod_21.text(), self.prod_22.text(), self.prod_23.text(), self.prod_24.text(), self.prod_25.text(), self.prod_26.text(), self.prod_27.text(), 
		self.prod_28.text(), self.prod_29.text(), self.prod_30.text(), self.prod_31.text(), self.prod_32.text(), self.prod_33.text(), self.prod_34.text(), 
		self.prod_35.text(), self.prod_36.text(), self.prod_37.text(), self.prod_38.text(), self.prod_39.text(), self.prod_40.text(), self.prod_41.text(), 
		self.prod_42.text(), self.prod_43.text(), self.prod_44.text(), self.prod_45.text(), self.prod_46.text(), self.prod_47.text(), self.prod_48.text(), 
		self.prod_49.text(), self.prod_50.text(), self.prod_51.text(), self.prod_52.text(), self.prod_53.text(), self.prod_54.text(), self.prod_55.text(), 
		self.prod_56.text(), self.prod_57.text(), self.prod_58.text(), self.prod_59.text(), self.prod_60.text(), self.prod_61.text(), self.prod_62.text(), 
		self.prod_63.text(), self.prod_64.text(), self.prod_65.text(), self.prod_66.text(), self.prod_67.text(), self.prod_68.text(), self.prod_69.text(), 
		self.prod_70.text(), self.prod_71.text(), self.prod_72.text(), self.prod_73.text(), self.prod_74.text(), self.prod_75.text()]	
		i=0
		for prod in self.prod_ListText :
			if self.prod_List[i].text() != "" :
				self.prod_List[i].setStyleSheet("QPushButton{background-color:darkgrey ; color :white ; border: 0px solid #2f4f4f  ; border-radius: 0.5px; height:35px;}\
				\n QPushButton:hover{color:white;background-color:#f96332;}\
				\nQPushButton:focus{border:0px solid #f96332;color:#f96332;background-color:#ecf0f5;}")
			i+=1
		self.prod_ListText = []
#=========================================				
	def productsSignalClicked(self)	: # CONNECT CATEGORIES BUTTONS SIGNAL
		
		i = 0
		while i <= 74:
			self.prod_List[i].clicked.connect(self.productsSlotClicked)
			i+=1
			
		self.editProduct.clicked.connect(self.productEditSlot)
		
	def productsSlotClicked(self) :	#SLOT : SEE WHO, CREATE variable = prodQntList[i], on click prodNum+=1, FILL TABLEWIDGET
		sender = self.sender()
		pc1 = 0
		query.execute("SELECT table_name FROM Tables ORDER BY ID ASC")
		self.tabNameList = []
		self.tabNameList = query.fetchall()
			
		for tab in self.tabNameList :
			if str(tab).strip("(',')") == self.viewTabNum.text() :
				for prod in self.productDialog.prodNameList :#CATEGORY 1
				
					if sender.text() == str(prod).strip("(',')") and str(prod).strip("(',')") != '' :#and self.prodCatList[pc1] == self.catList[0]:#CATEGORY 1
						# if str(prod).strip
						#GETTING QNT
						self.quantite = str(self.productDialog.prodQntList[pc1]).strip("(,)") 
						self.quantite = int(self.quantite) + 1
						self.productDialog.prodQntList[pc1] = self.quantite
						
						#GETTING price
						self.price = str(self.productDialog.prodPriceList[pc1]).strip("(,)") 
						self.price = int(self.price)
						
						#GETTING total
						self.total = self.quantite * self.price
						self.productDialog.prodTotalList[pc1] =  self.total
						
						#GETTING compta
						self.comptable = str(self.productDialog.prodComptList[pc1]).strip("(,)") 
						self.comptable = int(self.comptable)
						self.comptable +=1
						self.productDialog.prodComptList[pc1] = self.comptable

						#GETTING stock
						self.stockable = str(self.productDialog.prodStockList[pc1]).strip("(,)") 
						self.stockable = int(self.stockable) 
						self.stockable -=1
						self.productDialog.prodStockList[pc1] = self.stockable
						
						query.execute("INSERT INTO Tables_content (table_name, table_prod, table_price, table_qnt, table_prodTotal) VALUES (\
						'"+self.viewTabNum.text()+"','"+str(self.productDialog.prodNameList[pc1]).strip("(',')")+"',\
						"+str(self.productDialog.prodPriceList[pc1]).strip("(',')")+",1,0)")
						
						query.execute("SELECT table_prod FROM Tables_content WHERE table_name ='"+self.viewTabNum.text()+"' ORDER BY id")
						self.tabProd = []
						self.tabProd = query.fetchall()
						
						self.TotalList = list()
						query.execute("SELECT table_price FROM Tables_content WHERE table_name ='"+self.viewTabNum.text()+"' ORDER BY id")
						self.tabPrice = []
						self.tabPrice = query.fetchall()
						
						query.execute("SELECT table_qnt FROM Tables_content WHERE table_name ='"+self.viewTabNum.text()+"' ORDER BY id")
						self.tabQnt = []
						self.tabQnt = query.fetchall()
						
						query.execute("UPDATE Tables SET table_state='REMPLIE' WHERE table_name = '"+self.viewTabNum.text()+"'")
						
						# self.stateIndicator.setStyleSheet("background:white;border:none;;color:Blue;")
						conn.commit()	
						self.ProdList.clear()
						i=0
						self.ProdList.addItem("")
						for prod in self.tabProd:
							self.ProdList.addItem(str(i+1)+")                    "+str(prod).strip("(',')") +"  *  "+str(self.tabQnt[i]).strip("(',')") )
							i+=1
							
						self.calculTotal()
							
					pc1+=1

#=========================================				
	def validationButtonsSignalClicked(self):#CONNECT SELF.VALIDATION BUTTONS TO thEIR SLOT
					
		self.totalButton.clicked.connect(self.validationButtonsSlotClicked)
		self.undoButton.clicked.connect(self.validationButtonsSlotClicked)
		self.chefButton.clicked.connect(self.validationButtonsSlotClicked)
		self.validButton.clicked.connect(self.validationButtonsSlotClicked)

	#VALIDATION FUNCTION CALLED
	def calculTotal(self) :
		iT = 0
		self.TotalList = []
		self.prodQnt = 0
		
		query.execute("SELECT table_prod FROM Tables_content WHERE table_name ='"+self.viewTabNum.text()+"' ORDER BY id")
		self.tabProd = []
		self.tabProd = query.fetchall()
		
		query.execute("SELECT table_price FROM Tables_content WHERE table_name ='"+self.viewTabNum.text()+"' ORDER BY id")
		self.tabPrice = []
		self.tabPrice = query.fetchall()
		
		query.execute("SELECT table_qnt FROM Tables_content WHERE table_name ='"+self.viewTabNum.text()+"' ORDER BY id")
		self.tabQnt = []
		self.tabQnt = query.fetchall()
		
		for prod in self.tabProd :
			self.tabQnt[iT] = int(str(self.tabQnt[iT]).strip("(',')"))
			self.prodQnt = self.prodQnt + 1
			self.tabPrice[iT] = int(str(self.tabPrice[iT]).strip("(',')"))
			if self.tabQnt[iT] != 0  :
				self.totali = self.tabQnt[iT] * self.tabPrice[iT]
				self.totali=int(self.totali)
				self.TotalList.append(self.totali)
			iT+=1	
		self.viewTotal.display(sum(self.TotalList))

	def prodListView(self) :
			
		self.ProdList.clear()
		i=0
		self.ProdList.addItem("")
		for prod in self.tabProd:
			self.ProdList.addItem(str(i+1)+")                    "+str(prod).strip("(',')") +"  *  "+str(self.tabQnt[i]).strip("(',')") )
			i+=1

	def annulerTotal(self):
		query.execute("DELETE FROM 'Tables_content' WHERE table_name = '"+self.viewTabNum.text()+"'")
					
		query.execute("UPDATE Tables SET table_state='VIDE' WHERE table_name = '"+self.viewTabNum.text()+"'")
		self.valider = 0
		
		self.viewTabNum.setText('A Emporter')
		self.ProdList.clear()
	
		query.execute("SELECT ref FROM Ref WHERE id = 1")
		
		if str(self.ref) != str(query.fetchone()).strip("(',')") : 	
			self.ref = int(self.ref)-1
			self.viewRefNum.setText(str(self.ref))
				
		self.viewTotal.display(0)
		self.validVLD.VLD_Montant_rendu.setText("0")
		self.validVLD.VLD_Montant_recu.setText("0")
		
		iTu = 0
		for prod in self.productDialog.prodNameList :
			qnt = str(self.productDialog.prodQntList[iTu]).strip("(',')")
			qnt = int(qnt)
				
			self.quantite = str(self.productDialog.prodQntList[iTu]).strip("(,)") 
			self.quantite = int(self.quantite)
			
			if self.quantite != 0:
				self.comptable = str(self.productDialog.prodComptList[iTu]).strip("(,)") 	
				self.comptable = int(self.comptable) - self.quantite	
				self.productDialog.prodComptList[iTu] = self.comptable 
						
				self.stockable = str(self.productDialog.prodStockList[iTu]).strip("(,)")  
				self.stockable = int(self.stockable) + self.quantite	
				self.productDialog.prodStockList[iTu] = self.stockable 	
				
				# GETTING QNT
				self.quantite = 0
				self.productDialog.prodQntList[iTu] = self.quantite
				self.productDialog.prodTotalList[iTu] = 0
				
				self.TotalList = self.productDialog.prodTotalList
			iTu+=1
	
		query.execute("SELECT register_recette_total FROM Register WHERE register_date = '"+self.register.Date+"'")
		self.recetteTest = str(query.fetchone()).strip("(',')")
		
		if str(self.recetteTest).strip("(,)")  != str(self.register.recette).strip("(,)")  :
			self.register.recette = int(self.recetteTest)
			self.register.sumDay.setText(str(self.register.recette))			

	def chefPrint(self):
		try :
			TICKETchef = Document()
			
			p = TICKETchef.add_paragraph(" ___________________________________________________________________________________________________ ")

			h = TICKETchef.add_heading("				 "+self.society+"", level=1)
			h.bold = True
			h.italic = True
			p = TICKETchef.add_paragraph("				 "+self.localDateTime)
			
			p = TICKETchef.add_paragraph(" ___________________________________________________________________________________________________ ")
					
			
			tab = TICKETchef.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = ''
			heading_cells[1].text = 'PRODUITS COMMANDES'
			heading_cells[2].text = 'QUANTITES'
			heading_cells[3].text = ''
			
			pi = 0
			try:
				for prod in self.productDialog.prodNameList :
					query.execute("SELECT DISTINCT table_prod FROM Tables_content WHERE table_name ='"+self.viewTabNum.text()+"'\
						AND table_prod = '"+str(prod).strip("(',')")+"' ORDER BY id")
					self.tabprod = query.fetchone()
					i=0
					
					if str(prod).strip("(',')") == str(self.tabprod).strip("(',')") :
						query.execute("SELECT table_qnt FROM Tables_content WHERE table_name ='"+self.viewTabNum.text()+"'\
						AND table_prod = '"+str(self.tabprod).strip("(',')")+"' ORDER BY id")
						self.qnt = []
						self.qnt = query.fetchall()
						for qnt in self.qnt :
							self.qnt[i] = str(self.qnt[i]).strip("(',')")
							self.qnt[i] = int(self.qnt[i])
							i+=1
						if sum(self.qnt) != 0 :
							cells = tab.add_row().cells
							cells[1].text = str(self.tabprod).strip("(',')") 
							cells[2].text = str(sum(self.qnt))
					# pi+=1
			except:
				pass
			
			p = TICKETchef.add_paragraph(" ___________________________________________________________________________________________________ ")
		
			p = TICKETchef.add_heading("				         Table : "+self.viewTabNum.text()+" ", level = 2)
			
			p = TICKETchef.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKETchef.add_paragraph(" ___________________________________________________________________________________________________ ")

			TICKETchef.save('DOCUMENTS/TICKETS/Chef/TICKETchef-'+str(self.ref)+'.docx' )
			#os.startfile('DOCUMENTS\TICKETS\Chef\TICKETchef-'+str(self.ref)+'.docx' , 'print')
		except:
			self.messageFactory.raisePrintExcept("TICKETchef-"+str(self.ref))

	def validCommand(self) :
		try:
			self.register.recette += sum(self.TotalList)
			self.register.sumDay.setText("")
				
			query.execute("SELECT ref FROM Ref WHERE id = 1")
			if str(self.ref) == str(query.fetchone()).strip("(',')") :
				self.ref = int(self.ref) + 1 
				self.viewRefNum.setText(str(self.ref))
				
			self.ProdList.clear()
				
			TICKET = Document()
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")

			h = TICKET.add_heading("                                                 "+self.society+"", level=1)
			h.bold = True
			h.italic = True
			p = TICKET.add_paragraph("                                                            "+self.localDateTime)
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
					
			pi = 0
			tab = TICKET.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Nom produit'
			heading_cells[1].text = 'PRIX Produit'
			heading_cells[2].text = 'Quantité produit'
			heading_cells[3].text = 'Montant Produit'
			
			try:
				for prod in self.productDialog.prodNameList :
					query.execute("SELECT DISTINCT table_prod FROM Tables_content WHERE table_name ='"+self.viewTabNum.text()+"'\
						AND table_prod = '"+str(prod).strip("(',')")+"' ORDER BY id")
					self.tabprod = query.fetchone()
					i=0
						
					if str(prod).strip("(',')") == str(self.tabprod).strip("(',')") :
						query.execute("SELECT table_qnt FROM Tables_content WHERE table_name ='"+self.viewTabNum.text()+"'\
						AND table_prod = '"+str(self.tabprod).strip("(',')")+"' ORDER BY id")
						self.qnt = []
						self.qnt = query.fetchall()
						for qnt in self.qnt :
							self.qnt[i] = str(self.qnt[i]).strip("(',')")
							self.qnt[i] = int(self.qnt[i])
							i+=1
							
					query.execute("SELECT DISTINCT table_price FROM Tables_content WHERE table_name ='"+self.viewTabNum.text()+"'\
						AND table_prod = '"+str(prod).strip("(',')")+"' ORDER BY id")
					self.tablePrice = query.fetchone()
						
					if sum(self.qnt) != 0:
						cells = tab.add_row().cells
						cells[0].text = str(self.tabprod).strip("(',')") 
						cells[1].text = str(self.tablePrice).strip("(',')") + "DA"
						cells[2].text = str(sum(self.qnt))
						self.totalProd = sum(self.qnt) * int(str(self.tablePrice).strip("(',')"))
						cells[3].text = str(self.totalProd).strip("(',')") + "DA"
					pi+=1
			except:
				pass
				
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			try :	
				p = TICKET.add_heading("\t\t\t\tTotal à payer : "+str(sum(self.TotalList))+" DA ", level=2)
				p.bold = True
			except :
				return 0
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
		
			p = TICKET.add_paragraph("\t\t\t\t Reçu : "+self.validVLD.VLD_Montant_recu.text()+" DA, Rendu : "+self.validVLD.VLD_Montant_rendu.text()+" DA ")
			p = TICKET.add_paragraph("\t\t\t\tCaissier : "+str(self.user))
			p = TICKET.add_paragraph("\t\t\t\tTable No : "+self.viewTabNum.text()+" ")
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_heading("                                           Ticket No : "+str(self.ref)+". --"+self.society+"-- "+self.numero, level=3)
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			TICKET.save('DOCUMENTS/TICKETS/Tickets/TICKET-'+str(self.ref)+"_"+self.localDate+'.docx' )
			#os.startfile('DOCUMENTS\TICKETS\Tickets\TICKET-'+str(self.ref)+"_"+self.localDate+'.docx' , 'print')
					
			self.viewTotal.display(0)
			self.validVLD.VLD_Montant_rendu.setText("0")
			self.validVLD.VLD_Montant_recu.setText("0")
			
			iTu = 0
			for prod in self.productDialog.prodNameList :
					
				self.quantite = str(self.productDialog.prodQntList[iTu]).strip("(,)") 
				self.quantite = int(self.quantite)
				
				if self.quantite != 0:
					
					# GETTING QNT
					self.quantite = 0
					self.productDialog.prodQntList[iTu] = self.quantite
					self.productDialog.prodTotalList[iTu] = 0
				
				iTu+=1

			
			query.execute("SELECT ref FROM Ref WHERE id = 1")
			
			if str(self.ref) != str(query.fetchone()).strip("(',')") : 	
			
				query.execute("UPDATE Ref SET ref = "+str(self.ref)+" WHERE id = 1")
				
			query.execute("UPDATE Register SET register_recette_total = "+str(self.register.recette)+" WHERE register_date ='"+self.Date+"'")
			query.execute("SELECT register_recette_total FROM Register WHERE register_date ='"+self.Date+"'")
		
			query.execute("DELETE FROM 'Tables_content' WHERE table_name = '"+self.viewTabNum.text()+"'")
		
			self.viewTabNum.setText('A Emporter')
			
			iTu = 0
			
			for prod in self.productDialog.prodNameList :
			
				self.quantite = str(self.productDialog.prodQntList[iTu]).strip("(,)") 
				self.quantite = int(self.quantite)
				
				if self.quantite != 0:
					
					self.quantite = 0
					self.productDialog.prodQntList[iTu] = self.quantite
				
					self.total = 0
					self.productDialog.prodTotalList[iTu] = self.total
					
				iTu +=1
				
			i=0
			for prod in self.productDialog.prodNameList :
				query.execute("UPDATE Products SET product_qnt= 0 WHERE product_name = '"+str(self.productDialog.prodNameList[i]).strip("(',')")+"'")	
				query.execute("UPDATE Products SET product_comptable =  '"+str(self.productDialog.prodComptList[i]).strip("(',')")+"' WHERE product_name = '"+str(self.productDialog.prodNameList[i]).strip("(',')")+"'")
				query.execute("UPDATE Products SET product_stockable =  '"+str(self.productDialog.prodStockList[i]).strip("(',')")+"' WHERE product_name = '"+str(self.productDialog.prodNameList[i]).strip("(',')")+"'")
				query.execute("UPDATE Products SET product_total = 0 WHERE product_name = '"+str(self.productDialog.prodNameList[i]).strip("(',')")+"'")	
				i+=1
							
			query.execute("UPDATE Tables SET table_state='VIDE' WHERE table_name = '"+self.viewTabNum.text()+"'")
			
			conn.commit()
		except:
			self.messageFactory.raisePrintExcept("TICKET-"+str(self.ref)+"_"+self.localDate)

	def validationButtonsSlotClicked(self):#VALIDATION, UnDO, 
		sender = self.sender()
		iTu = 0
		iT = 0
		if sender == self.totalButton :#CREAT A LIST TOTAL, APPEND when ITEM != 0, SUM it, if SUM != 0 --> SELF.Ref++ --> data,DIALOGVLD get text andCALCUL, 
			self.calculTotal()
			if sum(self.TotalList) != 0 :
				
				self.validVLD.VLD_Montant_total.setText(str(sum(self.TotalList)))
				self.validVLD.show()
				
				self.validVLD.VLD_OK.clicked.connect(self.rendurecuSlot)

		if sender == self.undoButton :#ANNULER
			self.annulerTotal()
			
		if sender == self.chefButton and str(self.viewTotal.intValue()) != '0':
			self.chefPrint()
		
		if sender == self.validButton  and str(self.viewTotal.intValue()) != '0':
			self.validCommand()

	#DIALOG : DIALOG ALERT DELEGATES

	def rendurecuSlot(self):
		try :
			self.validVLD.VLD_Montant_rendu.setText(str(int(self.validVLD.VLD_Montant_recu.text()) - sum(self.TotalList) ))
		except :
			return 0 

	#EDIT PRODUCTS

	def productAdder(self):
		try :
			if  self.productDialog.newProdName.text() != '' and self.productDialog.newProdPrice.text() != ''\
			and self.productDialog.categoryAddComboBox.currentText() != ''	and self.productDialog.subCategoryAddComboBox.currentText() != ''\
			and self.productDialog.subCategoryAddComboBox.currentText() != '...' :
				if self.productDialog.stockCheck.checkState() != 0 :	
					if self.productDialog.newProdStock.text() != "" and self.productDialog.newProdStock.text() != "0" :
						query.execute("INSERT INTO Products (product_category_name,product_subCategory,product_name,product_price,product_qnt,\
						product_comptable,product_stockable,product_total,product_comptable_cst, product_stock,product_alert) VALUES\
						('{0}','{1}','{2}','{3}','{4}','{5}','{6}','{7}','{8}','{9}','{10}')"\
						.format(self.productDialog.categoryAddComboBox.currentText(), self.productDialog.subCategoryAddComboBox.currentText(), \
						self.productDialog.newProdName.text(),self.productDialog.newProdPrice.text(),0,0,self.productDialog.newProdStock.text(),\
						0,0,1,self.productDialog.newProdAlert.text()))
							
						conn.commit()

						self.productDialog.newProdName.setText("")	
						self.productDialog.newProdPrice.setText("")	
						self.productDialog.newProdStock.setText("0")
						self.productDialog.newProdAlert.setText("0")
						self.productDialog.categoryAddComboBox.clear()
						self.productDialog.subCategoryAddComboBox.clear()
						
						self.productDialog.prodDataList()
						 
						self.messageFactory.raiseAdder("Produit")
				
						self.productDialog.categoryAddComboBox.clear()
						self.prodDesign()
						i=0
						self.productDialog.categoryAddComboBox.addItem("")
						for prod in self.productDialog.categoryDialog.catList :
							self.productDialog.categoryAddComboBox.addItem(str(prod).strip("(',')"))
							i+=1
					else :
						self.messageFactory.raiseCaseExcept("La case Stock et le niveau d'alert")
				else :
					query.execute("INSERT INTO Products (product_category_name,product_subCategory,product_name,product_price,product_qnt,\
					product_comptable,product_stockable,product_total,product_comptable_cst, product_stock,product_alert) VALUES\
					('{0}','{1}','{2}','{3}','{4}','{5}','{6}','{7}','{8}','{9}','{10}')"\
					.format(self.productDialog.categoryAddComboBox.currentText(), self.productDialog.subCategoryAddComboBox.currentText(), \
					self.productDialog.newProdName.text(),self.productDialog.newProdPrice.text(),0,0,0,0,0,0,0))
							
					conn.commit()

					self.productDialog.newProdName.setText("")	
					self.productDialog.newProdPrice.setText("")
					self.productDialog.categoryAddComboBox.clear()
					self.productDialog.subCategoryAddComboBox.clear()
					
					self.productDialog.prodDataList()
					 
					self.messageFactory.raiseAdder("Produit")
			
					self.productDialog.categoryAddComboBox.clear()
					self.prodDesign()
					i=0
					self.productDialog.categoryAddComboBox.addItem("")
					for prod in self.productDialog.categoryDialog.catList :
						self.productDialog.categoryAddComboBox.addItem(str(prod).strip("(',')"))
						i+=1
					
			else :  
				self.messageFactory.raiseCaseExcept("Toutes les cases")
		except:
			self.messageFactory.raiseCharExcept()

	def productModifier(self):
		try:
			if  self.productDialog.newNameModif.text() != '' and self.productDialog.subCategoryModifComboBox.currentText() != '' and self.productDialog.subCategoryModifComboBox.currentText() != '...' :
					
				if self.productDialog.newNameModif.text() != self.productDialog.prodModifComboList.currentText() :
					query.execute("UPDATE Products SET product_name = '"+str(self.productDialog.newNameModif.text())+"' WHERE product_name = '"+self.productDialog.prodModifComboList.currentText()+"' AND product_category_name = '"+self.productDialog.categoryModifComboBox.currentText()+"'")

				elif self.productDialog.newPriceModif.text() != "" :
					query.execute("UPDATE Products SET product_price = '"+str(self.productDialog.newPriceModif.text())+"' WHERE product_name = '"+self.productDialog.prodModifComboList.currentText()+"' AND product_category_name = '"+self.productDialog.categoryModifComboBox.currentText()+"'")
				
				elif self.productDialog.newNameModif.text() != self.productDialog.prodModifComboList.currentText() and  self.productDialog.newPriceModif.text() != "" :
					query.execute("UPDATE Products SET product_name = '"+str(self.productDialog.newNameModif.text())+"',product_price ='"+self.productDialog.newPriceModif.text()+"' WHERE product_name = '"+self.productDialog.prodModifComboList.currentText()+"' AND product_category_name = '"+self.productDialog.categoryModifComboBox.currentText()+"'")
			
				conn.commit()
				
				self.productDialog.newNameModif.setText("")	
				self.productDialog.newPriceModif.setText("")
				
				self.productDialog.categoryModifComboBox.clear()
				self.productDialog.subCategoryModifComboBox.clear()
				
				self.productDialog.prodDataList()
				
				self.messageFactory.raiseModifier("Produit")
				self.prodDesign()
				i=0
				self.productDialog.categoryModifComboBox.addItem("")
				for prod in self.productDialog.categoryDialog.catList :
					self.productDialog.categoryModifComboBox.addItem(str(prod).strip("(',')"))
					i+=1	

					
			else :  
				self.messageFactory.raiseCaseExcept("toutes les cases")

		except:
			self.messageFactory.raiseCharExcept()

	def productDeleter(self):
		try:
			if  self.productDialog.delProdName.text() != '' and self.productDialog.subCategoryDelComboBox.currentText() != '' and self.productDialog.subCategoryDelComboBox.currentText() != '...' :
																					
				query.execute("DELETE FROM 'Products' WHERE product_name = '"+self.productDialog.delProdName.text()+"' AND product_category_name = '"+self.productDialog.categoryDelComboBox.currentText()+"'")
					
				conn.commit()
			
				self.productDialog.delProdName.setText("")	
				
				self.productDialog.categoryDelComboBox.clear()
				self.productDialog.subCategoryDelComboBox.clear()
				
				self.productDialog.prodDataList()
				
				self.messageFactory.raiseDeleter("Produit")
			
				self.productDialog.categoryDelComboBox.clear()
				self.prodDesign()
				i=0
				self.productDialog.categoryDelComboBox.addItem("")
				for prod in self.productDialog.categoryDialog.catList :
					self.productDialog.categoryDelComboBox.addItem(str(prod).strip("(',')"))
					i+=1
				
			else :  
				self.messageFactory.raiseCaseExcept("toutes les cases")
	
		except:
			self.messageFactory.raiseCharExcept()

	def comboInit(self)	:
		
		self.productDialog.prodAdd.clicked.connect(self.prodEditOK)
		self.productDialog.prodModif.clicked.connect(self.prodEditOK)
		self.productDialog.prodDel.clicked.connect(self.prodEditOK)
			
		i=0
		self.productDialog.categoryAddComboBox.clear()
		self.productDialog.categoryAddComboBox.addItem("...")
		self.productDialog.categoryModifComboBox.clear()
		self.productDialog.categoryModifComboBox.addItem("...")
		self.productDialog.categoryDelComboBox.clear()
		self.productDialog.categoryDelComboBox.addItem("...")
		while i <= 9:
			self.productDialog.categoryAddComboBox.addItem(self.categoryDialog.categoryName[i].text())
			self.productDialog.categoryModifComboBox.addItem(self.categoryDialog.categoryName[i].text())
			self.productDialog.categoryDelComboBox.addItem(self.categoryDialog.categoryName[i].text())
			i+=1
			
	def productEditSlot(self):
		self.comboInit()
		self.productDialog.show()
		
		self.productDialog.categoryAddComboBox.currentTextChanged.connect(self.prodSubCatSlot)
		self.productDialog.categoryModifComboBox.currentTextChanged.connect(self.prodSubCatSlot)
		self.productDialog.categoryDelComboBox.currentTextChanged.connect(self.prodSubCatSlot)
			
	def prodEditOK(self):	
		#ADD PRODUCT
		if self.sender() == self.productDialog.prodAdd:
			self.productAdder()
		
		#MODIF PRODUCT
		if self.sender() == self.productDialog.prodModif:
			self.productModifier()
	
		#DELETE PRODUCT
		if self.sender() == self.productDialog.prodDel:
			self.productDeleter()

	def subItemer(self,catCombo,catlist,combo,listsub1,listsub2,listsub3):
			i=0
			combo.clear()
			combo.addItem("...")
			
			while i <= 9:
				if catCombo.currentText() == str(catlist[i]).strip("(',')"):
					
					if listsub1[i].text() != '':
						combo.addItem(str(listsub1[i].text()).strip("(',')"))
						
					if listsub2[i].text() != '':
						combo.addItem(str(listsub2[i].text()).strip("(',')"))
						
					if listsub3[i].text() != '':
						combo.addItem(str(listsub3[i].text()).strip("(',')"))
				i+=1
			combo.currentTextChanged.connect(self.prodSubProdSlot)
				
	def prodSubCatSlot(self):
		
		if self.sender() == self.productDialog.categoryAddComboBox :
			self.subItemer(self.productDialog.categoryAddComboBox, self.categoryDialog.catList,self.productDialog.subCategoryAddComboBox,\
			self.categoryDialog.categorySub1,self.categoryDialog.categorySub2,self.categoryDialog.categorySub3)
		
		if self.sender() == self.productDialog.categoryModifComboBox :
			self.subItemer(self.productDialog.categoryModifComboBox, self.categoryDialog.catList,self.productDialog.subCategoryModifComboBox,\
			self.categoryDialog.categorySub1,self.categoryDialog.categorySub2,self.categoryDialog.categorySub3)
	
		if self.sender() == self.productDialog.categoryDelComboBox :
			self.subItemer(self.productDialog.categoryDelComboBox, self.categoryDialog.catList,self.productDialog.subCategoryDelComboBox,\
			self.categoryDialog.categorySub1,self.categoryDialog.categorySub2,self.categoryDialog.categorySub3)				

	def comboListItemer(self,combolist,subcombo):#JE SUIS ICI	
		#MODIF
		subcombo.clear()
		subcombo.addItem("...")
		
		query.execute("SELECT product_name FROM Products WHERE product_subCategory = '"+combolist.currentText()+"'")

		for prod in query.fetchall() :
			subcombo.addItem(str(prod).strip("(',')"))
			
		subcombo.currentTextChanged.connect(self.prodModifSetting)
		
	def prodSubProdSlot(self):
		if self.sender() == self.productDialog.subCategoryModifComboBox :
			self.comboListItemer(self.productDialog.subCategoryModifComboBox,self.productDialog.prodModifComboList)
		if self.sender() == self.productDialog.subCategoryDelComboBox :
			self.comboListItemer(self.productDialog.subCategoryDelComboBox,self.productDialog.prodDelComboList)
			
	def prodModifSetting(self):
		if self.sender() == self.productDialog.prodModifComboList:
			self.productDialog.newNameModif.setText(self.productDialog.prodModifComboList.currentText())
		if self.sender() == self.productDialog.prodDelComboList:
			self.productDialog.delProdName.setText(self.productDialog.prodDelComboList.currentText())
		
	def tableData(self) :
		query.execute("SELECT table_name FROM Tables ORDER BY ID ASC")
		self.tabNameList = []
		self.tabNameList = query.fetchall()
		
		query.execute("SELECT table_place FROM Tables ORDER BY ID ASC")
		self.tabplaceList = []
		self.tabplaceList = query.fetchall()
		
		query.execute("SELECT table_state FROM Tables ORDER BY ID ASC")
		self.tabstateList = []
		self.tabstateList = query.fetchall()
		
		self.i=0
		for tab in self.tabNameList :
			self.tablesListWidget.addItem(str(tab).strip("(',')"))
			self.i+=1	
			
#============================================================================================================ACCOUNTINGWIDGET 

#========== ACCOUNT WIDGET UIC CONVERT & Load
#======================================================

qtCreatorFile = "DESIGN/DIALOGS/accountCreatorDialog.ui" # Enter file here.

Ui_AccountCreatorDialog, QtBaseClass = uic.loadUiType(qtCreatorFile)

class AccountCreatorDialog(QDialog, Ui_AccountCreatorDialog):
	def __init__(self):
		QDialog.__init__(self)
		Ui_AccountCreatorDialog.__init__(self)
		self.setupUi(self)
		
		self.setWindowTitle("Comptabilité Générale.")
		self.Date = time.strftime("%Y-%m-%d")
		self.localDateTime = time.strftime("%d-%m-%Y      %H:%M:%S")
		self.localDate = time.strftime("%d-%m-%Y")
		
		self.referenceSelector = ReferenceSelector()
		self.society = self.referenceSelector.society 
		self.numero = self.referenceSelector.numero  
		
		self.messageFactory = MessageFactory()
		
		self.accountSignals()
		self.receiptBDD()
	
	def dateSelecterBet(self,type,desc,value,date,table,dater,title,list,col,to,dateTo):
		# try:
		query.execute(" SELECT "+value+" FROM "+table+" WHERE "+date+"='"+str(dater)+"' ORDER BY id ASC")
		depValue = query.fetchall()
	
		self.depTot = []
		
		pi = 0
		for dep in depValue :
			dep = str(dep).strip("(',')")
			dep = int(dep)
			self.depTot.append(dep)
			pi+=1
	
		query.execute(" SELECT "+type+", "+desc+", "+value+", "+date+"  FROM "+table+" WHERE client_date_out BETWEEN'"+str(dater)+"' AND 'client_date_out' ORDER BY "+value+" DESC")

		list.setColumnCount(col) 
		list.setRowCount(0)
		for row, form in enumerate(query):
			list.insertRow(row)
			for column, item in enumerate(form):
				#print(str(item))
				list.setItem(row, column,QTableWidgetItem(str(item)))
				
		to.setText(str(sum(self.depTot)))	
		dateTo.setText(title +" "+dater)
		# except:
			# self.messageFactory.raiseIndefinedExcept("\nVérifiez d'avoir saisit une date valide")

	def dateSelecter(self,type,desc,value,date,table,dater,title,list,col,to,dateTo):
		try:
			query.execute(" SELECT "+value+" FROM "+table+" WHERE "+date+"='"+str(dater)+"' ORDER BY id ASC")
			depValue = query.fetchall()
		
			self.depTot = []
			
			pi = 0
			for dep in depValue :
				dep = str(dep).strip("(',')")
				dep = int(dep)
				self.depTot.append(dep)
				pi+=1
		
			query.execute(" SELECT "+type+", "+desc+", "+value+", "+date+"  FROM "+table+" WHERE "+date+"='"+str(dater)+"' ORDER BY "+value+" DESC")

			list.setColumnCount(col) 
			list.setRowCount(0)
			for row, form in enumerate(query):
				list.insertRow(row)
				for column, item in enumerate(form):
					#print(str(item))
					list.setItem(row, column,QTableWidgetItem(str(item)))
					
			to.setText(str(sum(self.depTot)))	
			dateTo.setText(title +" "+dater)
		except:
			self.messageFactory.raiseIndefinedExcept("\nVérifiez d'avoir saisit une date valide")

	def dateSelectorGlobal(self,beforValue,value,date,table,title,list,col,to,Dateto) :
	
		query.execute(" SELECT  "+str(value)+" FROM "+str(table)+" ORDER BY id ASC")
		depValue = query.fetchall()
		list.clearContents()
	
		addTot = []
		
		pi = 0
		for dep in depValue :
			dep = str(dep).strip("(',')")
			dep = int(dep)
			addTot.append(dep)
			pi+=1
		
		list.clearContents()
		
		query.execute(" SELECT "+str(beforValue)+", "+str(value)+", "+str(date)+"\
		FROM "+str(table)+" ORDER BY "+str(date)+" DESC")

		list.setColumnCount(col) 
		list.setRowCount(0) 
		# self.providerHistory.providerTableWidget.setRowCount(0)
		for row, form in enumerate(query.fetchall()):
			list.insertRow(row)
			for column, item in enumerate(form):
				#print(str(item))
				list.setItem(row, column,QTableWidgetItem(str(item)))
				
		to.setText(str(sum(addTot)))	
		Dateto.setText(str(title))
		addTot = []			


	def accountSignals(self) :
		self.depGlobal.clicked.connect(self.accountSlots)
		self.depDay.clicked.connect(self.accountSlots)
		
		self.addGlobal.clicked.connect(self.accountSlots)
		self.addDay.clicked.connect(self.accountSlots)
		
		self.salaryGlobal.clicked.connect(self.accountSlots)
		self.salaryDay.clicked.connect(self.accountSlots)
		
		self.transGlobal.clicked.connect(self.accountSlots)
		self.transDay.clicked.connect(self.accountSlots)
		
		self.avanceGlobal.clicked.connect(self.accountSlots)
		self.avanceDay.clicked.connect(self.accountSlots)
		
		self.roomGlobal.clicked.connect(self.accountSlots)
		self.roomDay.clicked.connect(self.accountSlots)
		
		self.comptaPrintDay.clicked.connect(self.accountSlots)
		self.comptaPrintGlobal.clicked.connect(self.accountSlots)
		
		self.dateDep.clicked.connect(self.accountSlots)
		self.dateAdd.clicked.connect(self.accountSlots)
		self.dateProv.clicked.connect(self.accountSlots)
		self.dateSalary.clicked.connect(self.accountSlots)
		self.dateAvance.clicked.connect(self.accountSlots)
		self.dateRoom.clicked.connect(self.accountSlots)
		
		self.comptaValid.clicked.connect(self.accountSlots)
		
		self.receiptPrint.clicked.connect(self.accountSlots)
	
	def accountSlots(self)	:
		if self.sender() == self.depDay :
			self.dateSelecter("dep_type","dep_description","dep_value","dep_date","Register_dep",self.Date,"DEPENSES DU : ",\
			self.deptabWidget,4,self.depenseTotal,self.depenseDate)
			
		if self.sender() == self.depGlobal :
			self.dateSelectorGlobal("dep_type, dep_description","dep_value","dep_date","Register_dep",\
			"HISTORIQUE DEPENSES",self.deptabWidget,4,self.depenseTotal,self.depenseDate)
			
		if self.sender() == self.addDay :
			self.dateSelecter("add_depo","add_description","add_value","add_date","Register_add",self.Date,"DEPOTS DU : ",\
			self.addtabWidget,4,self.addTotal,self.addDate)
			
		if self.sender() == self.addGlobal :
			self.dateSelectorGlobal("add_depo, add_description","add_value","add_date","Register_add",\
			"HISTORIQUE DEPOTS",self.addtabWidget,4,self.addTotal,self.addDate)
			
		if self.sender() == self.transDay :
			self.dateSelecter("transact_name,transact_category","transact_products,transact_quantite","transact_total",\
			"transact_date_transact","Providers_History",self.Date,"TRANSACTIONS DU : ",self.providerTableWidget,6,\
			self.transTotal,self.transDate)	
			
		if self.sender() == self.transGlobal :
			self.dateSelectorGlobal("transact_name, transact_category,transact_products,transact_quantite","transact_total","transact_date_transact\
			","Providers_History","HISTORIQUE TRANSACTIONS",self.providerTableWidget,6,self.transTotal,self.transDate)
			
		if self.sender() == self.salaryDay :
			self.dateSelecter("worker_category, worker_name","worker_start","worker_salary","worker_salary_date","Worker_salaire",\
			self.Date,"SALAIRES DU : ",self.salaryTableWidget,5,self.salaryTotal,self.salaryDate)
			
		if self.sender() == self.salaryGlobal :
			self.dateSelectorGlobal("worker_category, worker_name,worker_start","worker_salary","worker_salary_date","Worker_salaire",\
			"HISTORIQUE SALAIRES",self.salaryTableWidget,5,self.salaryTotal,self.salaryDate)			

		if self.sender() == self.avanceDay :
			self.dateSelecter("worker_category, worker_name","worker_start","worker_advance","worker_advance_date","Worker_avance",\
			self.Date,"AVANCES DU : ",self.avanceTableWidget,5,self.avanceTotal,self.avanceDate)
			
		if self.sender() == self.avanceGlobal :
			self.dateSelectorGlobal("worker_category, worker_name,worker_start","worker_advance","worker_advance_date","Worker_avance",\
			"HISTORIQUE AVANCES",self.avanceTableWidget,5,self.avanceTotal,self.avanceDate)
			
		if self.sender() == self.roomDay :
			self.dateSelecter("client_room, client_name","client_sejour","client_total","client_date_in","Clients",\
			self.Date,"CHAMBRES RESERVEES LE : ",self.roomTableWidget,5,self.roomTotal,self.roomDate)
			
		if self.sender() == self.roomGlobal :
			self.dateSelectorGlobal("client_room, client_name,client_sejour","client_total","client_date_in","Clients",\
			"HISTORIQUE CHAMBRES ",self.roomTableWidget,5,self.roomTotal,self.roomDate)

			
		if self.sender() == self.dateDep :
			self.dateSelecter("dep_type","dep_description","dep_value","dep_date","Register_dep",self.depDater.text(),"DEPENSES DU : ",\
			self.deptabWidget,4,self.depenseTotal,self.depenseDate)

		if self.sender() == self.dateAdd :
			self.dateSelecter("add_depo","add_description","add_value","add_date","Register_add",self.addDater.text(),"DEPOTS DU : ",\
			self.addtabWidget,4,self.addTotal,self.addDate)

		if self.sender() == self.dateProv :
			self.dateSelecter("transact_name,transact_category","transact_products,transact_quantite","transact_total",\
			"transact_date_transact","Providers_History",self.provDater.text(),"TRANSACTIONS DU : ",self.providerTableWidget,6,\
			self.transTotal,self.transDate)

		if self.sender() == self.dateSalary :
			self.dateSelecter("worker_category, worker_name","worker_start","worker_salary","worker_salary_date","Worker_salaire",\
			self.salaryDater.text(),"SALAIRES DU : ",self.salaryTableWidget,5,self.salaryTotal,self.salaryDate)

		if self.sender() == self.dateAvance :
			self.dateSelecter("worker_category, worker_name","worker_start","worker_advance","worker_advance_date","Worker_avance",\
			self.avanceDater.text(),"AVANCES DU : ",self.avanceTableWidget,5,self.avanceTotal,self.avanceDate)

		if self.sender() == self.dateRoom :
			self.dateSelecter("client_room, client_name","client_sejour","client_total","client_date_in","Clients",\
			self.roomDater.text(),"CHAMBRES RESERVEES LE : ",self.roomTableWidget,5,self.roomTotal,self.roomDate)

			
		if self.sender() == self.comptaPrintDay :
			self.receiptBDD()
			self.accountPrintDay()
			
		if self.sender() == self.comptaPrintGlobal :
			self.receiptBDD()
			self.dateSelectorGlobal("dep_type, dep_description","dep_value","dep_date","Register_dep",\
			"DEPENSES GLOBALES",self.deptabWidget,4,self.depenseTotal,self.depenseDate)
			self.dateSelectorGlobal("add_depo, add_description","add_value","add_date","Register_add",\
			"DEPOTS GLOBAL",self.addtabWidget,4,self.addTotal,self.addDate)
			self.accountPrintGlobal()
			
		if self.sender() == self.receiptPrint :
			self.receiptBDD()
			self.printReceiptGlobal()
			
		if self.sender() == self.comptaValid :
			self.destroy()

	
	def receiptBDD(self) :
	
		query.execute(" SELECT  product_name FROM Products ORDER BY id ASC")
		self.prodNameList = query.fetchall()
		query.execute(" SELECT  product_comptable FROM Products  ORDER BY id ASC")
		self.prodComptList = query.fetchall()
		query.execute(" SELECT  product_price  FROM Products  ORDER BY id ASC")
		self.prodPriceList = query.fetchall()
	
		self.tot = []
		
		pi = 0
		for prod in self.prodNameList :
			self.prx = str(self.prodPriceList[pi]).strip("(',')")
			self.prx = int(self.prx)
			self.cmp = str(self.prodComptList[pi]).strip("(',')")
			self.cmp = int(self.cmp)
			
			self.prd = self.prx * self.cmp
			query.execute("UPDATE Products SET product_total = "+str(self.prd)+" WHERE product_name = '"+str(self.prodNameList[pi]).strip("(',')")+"'")
			self.tot.append(self.prd)
			pi+=1
	
		query.execute(" SELECT product_category_name , product_subCategory , product_name,product_price,product_comptable,product_total \
		FROM Products  ORDER BY product_total DESC")

		self.receiptTableWidget.setColumnCount(6) 
		self.receiptTableWidget.setRowCount(0)
		i=0
		for row, form in enumerate(query):
			if str(self.tot[i]) != "0" :
				self.receiptTableWidget.insertRow(row)
				for column, item in enumerate(form):
					self.receiptTableWidget.setItem(row, column,QTableWidgetItem(str(item)))
			i+=1
				
		self.receiptTotal.setText(str(sum(self.tot)))	
		self.tot = []
			
	def accountPrintDay(self) :	
		try:
			self.regTotal = []
				
			query.execute("SELECT register_recette_total FROM Register WHERE register_date = '"+self.Date+"'")
			self.recette = str(query.fetchone()).strip("(',')")	
			self.regTotal.append(int(self.recette))
			
			query.execute("SELECT register_sum_init FROM Register ORDER BY id DESC")
			self.suminit = str(query.fetchone()).strip("(',')")
			self.regTotal.append(int(self.suminit))
			
			query.execute("SELECT register_depense_total FROM Register WHERE register_date = '"+self.Date+"'")
			self.depense = str(query.fetchone()).strip("(',')")
			
			query.execute("SELECT register_ajout_total FROM Register WHERE register_date = '"+self.Date+"'")
			self.ajout = str(query.fetchone()).strip("(',')")
			self.regTotal.append(int(self.ajout))
			
			self.actual = str(sum(self.regTotal) - int(self.depense))
			self.actual = str(self.actual)
			
			TICKET = Document()
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")

			h = TICKET.add_heading("                        "+self.society+" : Comptabilité Journée", level=1)
			h.bold = True
			h.italic = True
			p = TICKET.add_paragraph("				             "+self.localDate)
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_heading("                                     Somme Actuelle en Caisse : "+str(self.actual)+" DA ", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			#DEPENSES
			p = TICKET.add_heading("DEPENSES : ", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			query.execute("SELECT dep_type FROM Register_dep WHERE dep_date = '"+self.Date+"' ORDER BY id")
			self.typeList = []
			self.typeList = query.fetchall()
			query.execute("SELECT dep_description FROM Register_dep WHERE dep_date = '"+self.Date+"' ORDER BY id")
			self.descList = []
			self.descList = query.fetchall()
			query.execute("SELECT dep_value FROM Register_dep WHERE dep_date = '"+self.Date+"' ORDER BY id")
			self.valList = []
			self.valList = query.fetchall()
			
			pi = 0 
			tab = TICKET.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Type'
			heading_cells[1].text = 'Description'
			heading_cells[2].text = 'Valeur'
			heading_cells[3].text = 'Date'
			pi =0
			val = []
			for dep in self.typeList :
				cells = tab.add_row().cells
				cells[0].text = str(self.typeList[pi]).strip("(',')")
				cells[1].text = str(self.descList[pi]).strip("(',')")
				cells[2].text = str(self.valList[pi]).strip("(',')") + " DA"
				cells[3].text = self.Date
				val.append(int(str(self.valList[pi]).strip("(',')")))
				pi+=1
				
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_heading("TOTAL : "+str(sum(val))+" DA", level=2)
			p.bold = False
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			#AJOUTS		
			p = TICKET.add_heading("DEPOTS : ", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			query.execute("SELECT add_depo FROM Register_add WHERE add_date = '"+self.Date+"' ORDER BY id")
			self.typeList = []
			self.typeList = query.fetchall()
			query.execute("SELECT add_description FROM Register_add WHERE add_date = '"+self.Date+"' ORDER BY id")
			self.descList = []
			self.descList = query.fetchall()
			query.execute("SELECT add_value FROM Register_add WHERE add_date = '"+self.Date+"' ORDER BY id")
			self.valList = []
			self.valList = query.fetchall()
			
			pi = 0 
			tab = TICKET.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Type'
			heading_cells[1].text = 'Description'
			heading_cells[2].text = 'Valeur'
			heading_cells[3].text = 'Date'
			pi =0
			val=[]
			for dep in self.typeList :
				cells = tab.add_row().cells
				cells[0].text = str(self.typeList[pi]).strip("(',')")
				cells[1].text = str(self.descList[pi]).strip("(',')")
				cells[2].text = str(self.valList[pi]).strip("(',')") + " DA"
				cells[3].text = self.Date
				val.append(int(str(self.valList[pi]).strip("(',')")))
				pi+=1
				
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_heading("TOTAL : "+str(sum(val))+" DA", level=2)
			p.bold = False
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			#SALAIRE	
			p = TICKET.add_heading("SAILAIRES : ", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			query.execute("SELECT worker_category FROM Worker_salaire WHERE worker_salary_date = '"+self.Date+"' ORDER BY id")
			self.worker = []
			self.worker = query.fetchall()
			query.execute("SELECT worker_name FROM Worker_salaire WHERE worker_salary_date = '"+self.Date+"' ORDER BY id")
			self.cat = []
			self.cat = query.fetchall()
			query.execute("SELECT worker_salary FROM Worker_salaire WHERE worker_salary_date = '"+self.Date+"' ORDER BY id")
			self.sal = []
			self.sal = query.fetchall()
			
			pi = 0 
			tab = TICKET.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Categorie'
			heading_cells[1].text = 'Employé(e)'
			heading_cells[2].text = 'Salaire'
			heading_cells[3].text = 'Date'
			pi =0
			for dep in self.worker :
				cells = tab.add_row().cells
				cells[0].text = str(self.worker[pi]).strip("(',')")
				cells[1].text = str(self.cat[pi]).strip("(',')")
				cells[2].text = str(self.sal[pi]).strip("(',')") + " DA"
				cells[3].text = self.Date
				pi+=1
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			#AVANCES	
			p = TICKET.add_heading("AVANCES : ", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			query.execute("SELECT worker_category FROM Worker_avance WHERE worker_advance_date = '"+self.Date+"' ORDER BY id")
			self.workerS = []
			self.workerS = query.fetchall()
			query.execute("SELECT worker_name FROM Worker_avance WHERE worker_advance_date = '"+self.Date+"' ORDER BY id")
			self.catS = []
			self.catS = query.fetchall()
			query.execute("SELECT worker_advance FROM Worker_avance WHERE worker_advance_date = '"+self.Date+"' ORDER BY id")
			self.salS = []
			self.salS = query.fetchall()
			
			pi = 0 
			tab = TICKET.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Categorie'
			heading_cells[1].text = 'Employé(e)'
			heading_cells[2].text = 'Avance'
			heading_cells[3].text = 'Date'
			pi =0
			for dep in self.workerS :
				cells = tab.add_row().cells
				cells[0].text = str(self.workerS[pi]).strip("(',')")
				cells[1].text = str(self.catS[pi]).strip("(',')")
				cells[2].text = str(self.salS[pi]).strip("(',')") + " DA"
				cells[3].text = self.Date
				pi+=1
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_heading("SOMME INITIALE EN CAISSE : " +self.suminit+ " DA", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_heading("RECETTE DU JOUR  : " +self.recette+ " DA", level=2)
			p.bold = True
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_heading("\t\tCompanie :  --"+self.society+"-- "+self.localDateTime, level=3)
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			TICKET.save('DOCUMENTS/COMPTABILITE/GLOBALE/compta_jour_'+self.localDate+'.docx' )
			# os.startfile('DOCUMENTS\COMPTABILITE\GLOBALE\compta_jour_'+self.localDate+'.docx' , 'print')
			
			conn.commit()
		except:
			self.messageFactory.raisePrintExcept("compta_jour_"+self.localDate)

	def accountPrintGlobal(self) :	
		try:
			TICKET = Document()
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")

			h = TICKET.add_heading("                        "+self.society+" : Comptabilité Globale", level=1)
			h.bold = True
			h.italic = True
			p = TICKET.add_paragraph("				             "+self.localDate)
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			#DEPENSES
			p = TICKET.add_heading("DEPENSES : "+self.depenseTotal.text()+" DA", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			query.execute("SELECT dep_type FROM Register_dep  ORDER BY dep_date DESC")
			self.typeList = []
			self.typeList = query.fetchall()
			query.execute("SELECT dep_description FROM Register_dep ORDER BY dep_date DESC")
			self.descList = []
			self.descList = query.fetchall()
			query.execute("SELECT dep_value FROM Register_dep ORDER BY dep_date DESC")
			self.valList = []
			self.valList = query.fetchall()
			query.execute("SELECT dep_date FROM Register_dep ORDER BY dep_date DESC")
			self.datin = []
			self.datin = query.fetchall()
			
			pi = 0 
			tab = TICKET.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Type'
			heading_cells[1].text = 'Description'
			heading_cells[2].text = 'Valeur'
			heading_cells[3].text = 'Date'
			pi =0
			for dep in self.typeList :
				cells = tab.add_row().cells
				cells[0].text = str(self.typeList[pi]).strip("(',')")
				cells[1].text = str(self.descList[pi]).strip("(',')")
				cells[2].text = str(self.valList[pi]).strip("(',')") + " DA"
				cells[3].text = str(self.datin[pi]).strip("(',')")
				pi+=1
				
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			#AJOUTS		
			p = TICKET.add_heading("DEPOTS : "+self.addTotal.text()+" DA", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			query.execute("SELECT add_depo FROM Register_add ORDER BY add_date DESC")
			self.typeList = []
			self.typeList = query.fetchall()
			query.execute("SELECT add_description FROM Register_add ORDER BY add_date DESC")
			self.descList = []
			self.descList = query.fetchall()
			query.execute("SELECT add_value FROM Register_add ORDER BY add_date DESC")
			self.valList = []
			self.valList = query.fetchall()
			query.execute("SELECT add_date FROM Register_add ORDER BY add_date DESC")
			self.datin = []
			self.datin = query.fetchall()
			
			pi = 0 
			tab = TICKET.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Dépositaire'
			heading_cells[1].text = 'Description'
			heading_cells[2].text = 'Valeur'
			heading_cells[3].text = 'Date'
			pi =0
			for dep in self.typeList :
				cells = tab.add_row().cells
				cells[0].text = str(self.typeList[pi]).strip("(',')")
				cells[1].text = str(self.descList[pi]).strip("(',')")
				cells[2].text = str(self.valList[pi]).strip("(',')") + " DA"
				cells[3].text = str(self.datin[pi]).strip("(',')")
				pi+=1
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
		
			#SALAIRE	
			p = TICKET.add_heading("SAILAIRES : ", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			query.execute("SELECT worker_category FROM Worker_salaire ORDER BY worker_salary_date DESC")
			self.worker = []
			self.worker = query.fetchall()
			query.execute("SELECT worker_name FROM Worker_salaire ORDER BY worker_salary_date DESC")
			self.cat = []
			self.cat = query.fetchall()
			query.execute("SELECT worker_salary FROM Worker_salaire ORDER BY worker_salary_date DESC")
			self.sal = []
			self.sal = query.fetchall()
			query.execute("SELECT worker_salary_date FROM Worker_salaire ORDER BY worker_salary_date DESC")
			self.aDate = []
			self.aDate = query.fetchall()
			
			pi = 0 
			tab = TICKET.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Categorie'
			heading_cells[1].text = 'Employé(e)'
			heading_cells[2].text = 'Salaire'
			heading_cells[3].text = 'Date'
			pi =0
			for dep in self.worker :
				cells = tab.add_row().cells
				cells[0].text = str(self.worker[pi]).strip("(',')")
				cells[1].text = str(self.cat[pi]).strip("(',')")
				cells[2].text = str(self.sal[pi]).strip("(',')") + " DA"
				cells[3].text = str(self.aDate[pi]).strip("(',')")
				pi+=1
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			#AVANCES	
			p = TICKET.add_heading("AVANCES : ", level=2)
			p.bold = True
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			query.execute("SELECT worker_category FROM Worker_avance ORDER BY worker_advance_date DESC")
			self.workerS = []
			self.workerS = query.fetchall()
			query.execute("SELECT worker_name FROM Worker_avance ORDER BY worker_advance_date DESC")
			self.catS = []
			self.catS = query.fetchall()
			query.execute("SELECT worker_advance FROM Worker_avance  ORDER BY worker_advance_date DESC")
			self.salS = []
			self.salS = query.fetchall()
			query.execute("SELECT worker_advance_date FROM Worker_avance  ORDER BY worker_advance_date DESC")
			self.aDate = []
			self.aDate = query.fetchall()
			
			pi = 0 
			tab = TICKET.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Categorie'
			heading_cells[1].text = 'Employé(e)'
			heading_cells[2].text = 'Avance'
			heading_cells[3].text = 'Date'
			pi =0
			for dep in self.workerS :
				cells = tab.add_row().cells
				cells[0].text = str(self.workerS[pi]).strip("(',')")
				cells[1].text = str(self.catS[pi]).strip("(',')")
				cells[2].text = str(self.salS[pi]).strip("(',')") + " DA"
				cells[3].text = str(self.aDate[pi]).strip("(',')")
				pi+=1
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_heading("RECETTE GLOBALE  : " +str(sum(self.tot))+ " DA", level=2)
			p.bold = True
			
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKET.add_heading("\t\tCompanie :  --"+self.society+"-- "+self.localDateTime, level=3)
			p = TICKET.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			TICKET.save('DOCUMENTS/COMPTABILITE/GLOBALE/compta_global_'+self.localDate+'.docx' )
			# os.startfile('DOCUMENTS\COMPTABILITE\GLOBALE\compta_global_'+self.localDate+'.docx' , 'print')
			
			conn.commit()
		except:
			self.messageFactory.raisePrintExcept("compta_global_"+self.localDate)
					
	def printReceiptGlobal(self) : 
		try:
			TICKETreceipt = Document()
			
			p = TICKETreceipt.add_paragraph(" ___________________________________________________________________________________________________ ")

			h = TICKETreceipt.add_heading("\t\t\t\t     "+self.society+"", level=1)
			h.bold = True
			h.italic = True
			p = TICKETreceipt.add_paragraph("\t\t\t\t    "+self.localDateTime)
			
			p = TICKETreceipt.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			p = TICKETreceipt.add_heading("\t\t\t\t      RECETTE GENERALE ",level=2)			
			
			p = TICKETreceipt.add_paragraph(" ___________________________________________________________________________________________________ ")
					
			pi = 0
			tab = TICKETreceipt.add_table(1,4)
			heading_cells = tab.rows[0].cells
			heading_cells[0].text = 'Nom produit'
			heading_cells[1].text = 'PRIX Produit'
			heading_cells[2].text = 'Quantité produit'
			heading_cells[3].text = 'Montant Produit'
			
			self.tot = list()
			
			for prod in self.prodNameList :
				self.prx = str(self.prodPriceList[pi]).strip("(',')")
				self.prx = int(self.prx)
				self.cmp = str(self.prodComptList[pi]).strip("(',')")
				self.cmp = int(self.cmp)
				
				self.prd = self.prx * self.cmp
				self.tot.append(self.prd)
				
				
				if self.cmp != 0:
					cells = tab.add_row().cells
					cells[0].text = str(self.prodNameList[pi]).strip("(',')")
					cells[1].text = str(self.prodPriceList[pi]).strip("(',')") + "DA"
					cells[2].text = str(self.prodComptList[pi]).strip("(',')")
					cells[3].text = str(self.prd) + "DA"
				pi+=1
				
			p = TICKETreceipt.add_paragraph(" ___________________________________________________________________________________________________ ")
				
			p = TICKETreceipt.add_heading("\t                                  Recette totale : "+"  *** "+str(sum(self.tot))+" DA ***", level=2)
			p.bold = True
			
			p = TICKETreceipt.add_paragraph(" ___________________________________________________________________________________________________ ")
			p = TICKETreceipt.add_paragraph(" ___________________________________________________________________________________________________ ")
			
			TICKETreceipt.save('DOCUMENTS/COMPTABILITE/RECETTE/Recette-'+self.localDate+'.docx' )
			#os.startfile('DOCUMENTS\COMPTABILITE\RECETTE\Recette-'+self.localDate+'.docx' , 'print')	
		except:
			self.messageFactory.raisePrintExcept("Recette-"+self.localDate)	
	
#============================================================================================================SETTINGS
	
qtValid= "DESIGN/DIALOGS/settingCreatorDialog.ui"
Ui_SettingCreatorDialog, QtBaseClass = uic.loadUiType(qtValid)

class SettingCreatorDialog(QDialog, Ui_SettingCreatorDialog):#CONFIRM : VALID DIALOG ON VLDButton CLICK()
	def __init__(self):
		QDialog.__init__(self)
		Ui_SettingCreatorDialog.__init__(self)
		self.setupUi(self)
		
		self.referenceSelector = ReferenceSelector()
		
		self.messageFactory = MessageFactory()
		
		self.settingSignal()
		self.selection()
		
	def selection(self):
		self.society = self.referenceSelector.society 
		self.numero = self.referenceSelector.numero  
		
		self.societyName.setText(self.society)
		self.societyNumber.setText(self.numero)
		
		query.execute("SELECT user FROM User WHERE id = 1")
		self.user = query.fetchone()
		self.user= str(self.user).strip("(',')")
		self.who.setText(self.user)
		
		query.execute("SELECT user_name FROM Login ORDER BY id ASC")
		self.users=[]
		self.users = query.fetchall()
		query.execute("SELECT user_password FROM Login ORDER BY id ASC")
		self.passwords=[]
		self.passwords=query.fetchall()
		
		self.modSetComboBox.clear()
		self.modSetComboBox.addItem("...")
		self.delSetComboBox.clear()
		self.delSetComboBox.addItem("...")
		
		i=0
		for user in self.users :
			i+=1
			self.userNbr.setText(str(i))
			self.modSetComboBox.addItem(str(user).strip("(',')"))
			self.delSetComboBox.addItem(str(user).strip("(',')"))
					
	def settingSignal(self) :
		self.logAdd.clicked.connect(self.settingSlot)
		
		self.logModif.clicked.connect(self.settingSlot)
		self.modSetComboBox.currentTextChanged.connect(self.comboSetting)
		
		self.logDel.clicked.connect(self.settingSlot)
		self.delSetComboBox.currentTextChanged.connect(self.comboSetting)
		
		self.settingInfo.clicked.connect(self.settingSlot)
		
		self.setOK.clicked.connect(self.settingSlot)
		self.setAnnuler.clicked.connect(self.settingSlot)
		
	def settingSlot(self):
		if self.sender() == self.logAdd :
			self.settingAdd()
			self.selection()
			
		if self.sender() == self.logModif :
			self.settingModif()
			self.selection()
			
		if self.sender() == self.logDel :
			self.settingDel()
			self.selection()			
		
		if self.sender() == self.settingInfo :
			self.socityInfo()
			self.selection()		
		
		if self.sender() == self.setOK :
			self.selection()	
			self.destroy()
		
		if self.sender() == self.setAnnuler :	
			self.destroy()
			
	def comboSetting(self):
			
		self.modUserName.setText(self.modSetComboBox.currentText())
		self.delUserName.setText(self.delSetComboBox.currentText())
		
	def settingAdd(self) :
		try:
			if self.newUserName.text() != "" and self.newPassWord.text() != "" :
				query.execute("INSERT INTO Login (user_name, user_password) VALUES('"+self.newUserName.text()+"','"+self.newPassWord.text()+"')")
				conn.commit()
				self.newUserName.setText("")
				self.newPassWord.setText("")
				
				self.messageFactory.raiseAdder("Utilisateur")
				self.selection()
				
			else :
				self.messageFactory.raiseCaseExcept("toute les case")
				if self.newUserName.text() != "" :
					self.messageFactory.raiseCaseExcept("le nom d'utilisaeur")
				if self.newPassWord.text() != "":
					self.messageFactory.raiseCaseExcept("le mot de passe")
		except:
			self.messageFactory.raiseCharExcept()
		
	def settingModif(self) :
		try:
			if self.modUserName.text() != "" and self.modPassWord.text() != "" and self.who.text() == self.modSetComboBox.currentText() and \
				self.modUserName.text() != "..." and self.modSetComboBox.currentText() != "..." :
				query.execute("UPDATE Login SET user_name='"+self.modUserName.text()+"', user_password='"+self.modPassWord.text()+"' WHERE user_name = '"+self.modSetComboBox.currentText()+"'")
				conn.commit()
				
				self.messageFactory.raiseModifier("Utilisateur")
				
				self.modUserName.setText("")
				self.modPassWord.setText("")
				self.selection()
				
			else :
				self.messageFactory.raiseCaseExcept("toutes les cases")
				if self.modUserName.text() == "" :
					self.messageFactory.raiseCaseExcept("le nom d'utilisateur")
				if self.modPassWord.text() == "" :
					self.messageFactory.raiseCaseExcept("le mot de passe")
				# elif self.modSetComboBox.currentText() != self.user :
					# msg.setText("seul "+self.modSetComboBox.currentText()+" est abilité à modifier se compte !")
					
				self.modUserName.setText("")
				self.modPassWord.setText("")
		except:
			self.messageFactory.raiseCharExcept()
			
	def settingDel(self) :
		try:
			if self.delUserName.text() != "" and self.delUserName.text() != "..." and self.delSetComboBox.currentText() != "..."  and\
			self.delSetComboBox.currentText() == self.user :
				query.execute("DELETE FROM 'Login' WHERE user_name = '"+self.delSetComboBox.currentText()+"'")
				conn.commit()
				
				self.messageFactory.raiseDeleter("Utilisateur")
				
				self.delUserName.setText("")
				self.selection()
				
			else :
				self.messageFactory.raiseCaseExcept("toutes les cases")
				if self.delUserName.text() == "" :
					self.messageFactory.raiseCaseExcept("le Nom d'utilisateur")
				# if self.user != "admin" :
					# msg.setText("seul admin est abilité a changer les comptes !")
				# elif self.delSetComboBox.currentText() != self.user :
					# msg.setText("seul "+self.delSetComboBox.currentText()+" est abilité à modifier se compte !")
					
				self.delUserName.setText("")
		except:
			self.messageFactory.raiseCharExcept()
		
	def socityInfo(self):
		try:
			if self.societyName.text() != "" and self.societyNumber.text() != "" :
				query.execute("UPDATE Society SET society_name='"+self.societyName.text()+"', society_number='"+self.societyNumber.text()+"' WHERE id = 1 ")
				conn.commit()
							
				self.messageFactory.raiseModifier("Information")
				
				self.modUserName.setText("")
				self.modPassWord.setText("")
				
			else :
				self.messageFactory.raiseCaseExcept("toutes les cases")
				if self.societyName.text() == "" :
					self.messageFactory.raiseCaseExcept("le nom de la society")
				elif self.societyNumber.text() == "" :
					self.messageFactory.raiseCaseExcept("le numéro de la society")
		except:
			self.messageFactory.raiseCharExcept()
	
#============================================================================================================CONNECTION 		
	
qtValid= "DESIGN/DIALOGS/connecterDialog.ui"
Ui_ConnectDialog, QtBaseClass = uic.loadUiType(qtValid)

class ConnectDialog(QDialog, Ui_ConnectDialog):#CONFIRM : VALID DIALOG ON VLDButton CLICK()
	def __init__(self):
		QDialog.__init__(self)
		Ui_ConnectDialog.__init__(self)
		self.setupUi(self)
		
		self.disconnect.clicked.connect(self.connSlot)
		
	def connSlot(self):
		if self.sender() == self.disconnect :
			self.destroy()
			
	def login(self,window) :
		query.execute("SELECT user_name FROM Login ORDER BY id")
		self.users=[]
		self.users = query.fetchall()
		query.execute("SELECT user_password FROM Login ORDER BY id")
		self.passwords=[]
		self.passwords=query.fetchall()
		i=0
		j=0
		while i < len(self.users) :
			while j < len(self.passwords) :
				if self.userName.text() == str(self.users[i]).strip("(',')")  and self.userPassword.text() == str(self.passwords[j]).strip("(',')") :
					query.execute("UPDATE User SET user='"+self.userName.text()+"' WHERE id = 1")
					conn.commit()
					window.show()
					self.userName.setText("")
					self.userPassword.setText("")
					self.destroy()
				j+=1
				i+=1	
								
#============================================================================================================TABLE CREATORDIALOG

qtTableCreator= "DESIGN/DIALOGS/tableCreatorDialog.ui"
Ui_tableCreatorDialog, QtBaseClass = uic.loadUiType(qtTableCreator)

class TableCreatorDialog(QDialog, Ui_tableCreatorDialog):#EDIT : MODIF Product Name,Price DIALOG
	
	def __init__(self):
		QDialog.__init__(self)
		Ui_tableCreatorDialog.__init__(self)
		self.setupUi(self)
		
		self.setWindowTitle("Edition des Tables")
		 
#============================================================================================================PROVIDERS 
#============PROVIDER DIALOG		
qtProviderCreatorDialog= "DESIGN/DIALOGS/providerCreatorDialog.ui"
Ui_providerCreatorDialog, QtBaseClass = uic.loadUiType(qtProviderCreatorDialog)

class ProviderCreatorDialog(QDialog, Ui_providerCreatorDialog):#CONFIRM : VALID DIALOG ON VLDButton CLICK()
	def __init__(self):
		QDialog.__init__(self)
		Ui_providerCreatorDialog.__init__(self)
		self.setupUi(self)	
		
		self.providerData()
		
	def providerData(self):
		#SELECT Providers INTO LISTS:
		query.execute("SELECT provider_name FROM Providers ORDER BY id ")
		self.providerNameList = list()
		self.providerNameList = query.fetchall()
		
		query.execute("SELECT provider_phone FROM Providers ORDER BY id ")
		self.providerPhoneList = list()
		self.providerPhoneList = query.fetchall()
		
		
		query.execute("SELECT provider_category FROM Providers ORDER BY id ")
		self.providerCategoryList = list()
		self.providerCategoryList = query.fetchall()
		
		
		query.execute("SELECT provider_product1 FROM Providers ORDER BY id ")
		self.providerProductList1 = list()
		self.providerProductList1 = query.fetchall()
		
		
		query.execute("SELECT provider_product2 FROM Providers ORDER BY id ")
		self.providerProductList2 = list()
		self.providerProductList2 = query.fetchall()
		
		
		query.execute("SELECT provider_product3 FROM Providers ORDER BY id ")
		self.providerProductList3 = list()
		self.providerProductList3 = query.fetchall()
		
		
		query.execute("SELECT provider_product4 FROM Providers ORDER BY id ")
		self.providerProductList4 = list()
		self.providerProductList4 = query.fetchall()
		
		
		query.execute("SELECT provider_product5 FROM Providers ORDER BY id ")
		self.providerProductList5 = list()
		self.providerProductList5 = query.fetchall()
		
		
		query.execute("SELECT provider_price1 FROM Providers ORDER BY id ")
		self.providerPriceList1 = list()
		self.providerPriceList1 = query.fetchall()
		
		
		query.execute("SELECT provider_price2 FROM Providers ORDER BY id ")
		self.providerPriceList2 = list()
		self.providerPriceList2 = query.fetchall()
		
		
		query.execute("SELECT provider_price3 FROM Providers ORDER BY id ")
		self.providerPriceList3 = list()
		self.providerPriceList3 = query.fetchall()
		
		
		query.execute("SELECT provider_price4 FROM Providers ORDER BY id ")
		self.providerPriceList4 = list()
		self.providerPriceList4 = query.fetchall()
		
		
		query.execute("SELECT provider_price5 FROM Providers ORDER BY id ")
		self.providerPriceList5 = list()
		self.providerPriceList5 = query.fetchall()	
		
		#SET PRODUCTS AND PRICES INTO LISTS FROM DIALOG
		
		self.providersProductList = list()
		self.providersProductPriceList = list()
		
		self.providersProductList = [self.dialog_product1,self.dialog_product2,
		self.dialog_product3,self.dialog_product4,
		self.dialog_product5]
		
		self.providersProductPriceList = [self.dialog_productPrice1,self.dialog_productPrice2,
		self.dialog_productPrice3,
		self.dialog_productPrice4,self.dialog_productPrice5]

#============PROVIDER CLASS		

qtProvidersWidget= "DESIGN/WIDGETS/providersWidget.ui"
Ui_providerWidget, QtBaseClass = uic.loadUiType(qtProvidersWidget)

class ProviderWidget(QWidget, Ui_providerWidget):

	def __init__(self):
	
		QWidget.__init__(self)
		Ui_providerWidget.__init__(self)
		self.setupUi(self)
				
		#==========Logical & Réusables Variables
		
		self.setWindowTitle("Gestion Des Fournisseur.")
		self.Date = time.strftime("%Y-%m-%d")
		self.localDateTime = time.strftime("%d-%m-%Y      %H:%M:%S")
		self.localDate = time.strftime("%d-%m-%Y")
		self.referenceSelector = ReferenceSelector()
		self.society = self.referenceSelector.society 
		self.numero = self.referenceSelector.numero  
	
		query.execute("SELECT ref FROM Ref WHERE id=2")
		self.ref = str(query.fetchone()).strip("(',')")
		self.viewProviderRefNum.setText(str(self.ref))
		
		#==========INIT CLASSES OF TYPE DIALOGS FROM MODULE "CreatorDialog":
		self.providerEdit = ProviderCreatorDialog()
		self.register = RegisterCreatorDialog()
		
		#==========DELEGATES INSTANCES :
		self.messageFactory = MessageFactory()
		
		self.providerDataManager()#DATAMANAGER
		
		self.validButtonsSignals()#SIGNAL
		
		self.providersListSignal()#SIGNAL : QListWidget item clicked
		
		self.spinBoxSignal()#Signal : Mths payement
		
		self.EditSignals()
		
		self.providerCount()
		
	def reference(self)	:#SELECT REF
		query.execute("SELECT ref FROM Ref WHERE id=2")
		self.ref = str(query.fetchone()).strip("(',')")	
		
	def providerDataManager(self):#SELECT PROVIDERS LIST THEM; PRODUCTLIST,	QNTLIST, PRICELIST
			
		#SET PRODUCTS AND PRICES INTO LISTS FROM SELF
		
		query.execute("SELECT DISTINCT provider_category FROM Providers ORDER BY id ")
		self.providerCat = list()
		self.providerCat = query.fetchall()
		
		for type in self.providerCat:
			self.providerProductType.clear()
			self.providerProductType.addItem("")	
			self.providerProductType.addItem(str(type).strip("(',')"))	
		
		self.ProductList = []
		self.ProductPriceList = []
		self.productQuantite = []
		
		self.ProductList = [self.product1,self.product2,self.product3,self.product4,self.product5]
		
		self.ProductPriceList = [self.productPrice1,self.productPrice2,self.productPrice3,self.productPrice4,self.productPrice5]
		
		self.productQuantite = [self.productQuantite1, self.productQuantite2, self.productQuantite3,
		self.productQuantite4, self.productQuantite5]
			
	def validButtonsSignals(self):
		
		#SELF SIGNALS :
		self.providerEditer.clicked.connect(self.validEditSlots)#Bouton d'edition
		
		self.validButton.clicked.connect(self.validSlot)
		self.nextButton.clicked.connect(self.validSlot)
		self.undoButton.clicked.connect(self.validSlot)

	def totalTransaction(self):
		self.reference()
		
		if str(self.ref).strip("(',')") == self.viewProviderRefNum.text():
			self.viewProviderRefNum.setText(str(int(self.ref) + 1))
			
		self.TOTAL = []	
		
		self.ProdList.clear()
		self.ProdList.addItem("__________________________________")
	
		i = 0
		self.prods = 0
		while i <= 4 :
		
			if self.ProductPriceList[i].text() != '0' and self.productQuantite[i].text() != '0' and self.productQuantite[i].text() != "":
			
				self.price = str(self.ProductPriceList[i].text())
				self.price = int(self.price)
				
				self.quantite = str(self.productQuantite[i].text())
				self.quantite = int(self.quantite)
				
				self.totali = self.price * self.quantite
				
				self.TOTAL.append(self.totali)
				
				self.prods+=int(self.quantite)
				
				self.ProdList.addItem(str(self.ProductList[i].text())+" : "+str(self.price)+" DA * "+str(self.quantite))
				self.ProdList.addItem("__________________________________")
			i+=1
		self.viewProviderTotal.display(sum(self.TOTAL))
		self.provProducts.setText(str(self.prods))
		self.totalProviderCash.setText(str(sum(self.TOTAL)))
		self.totalProviderCredit.setText(str(sum(self.TOTAL)))
		
		if sum(self.TOTAL) != 0 :
			self.viewProviderTotal.display(sum(self.TOTAL))
			self.provProducts.setText(str(self.prods))
			self.totalProviderCash.setText(str(sum(self.TOTAL)))
			self.totalProviderCredit.setText(str(sum(self.TOTAL)))
			
			self.ProdList.addItem("")
			self.ProdList.addItem("TOTAL TRANSACTION : "+str(sum(self.TOTAL))+" DA ")
			self.ProdList.addItem("")
			self.ProdList.addItem("__________________________________")	
		
		self.TOTAL=[]
		
	def annulTransaction(self):
		self.reference()
		self.checkBoxCash.setChecked(False)
		self.checkBoxCredit.setChecked(False)

		if str(self.ref).strip("(',')") != self.viewProviderRefNum.text():
			self.viewProviderRefNum.setText(str(int(self.viewProviderRefNum.text())-1))
		
		self.ProdList.clear()
		i = 0
		self.prods = 0
		while i <= 4 :
			
			self.ProductList[i].setText("")
		
			self.price = 0
			self.ProductPriceList[i].setText(str(self.price))
			
			self.quantite = 0
			self.productQuantite[i].setValue(self.quantite)  
				
			i+=1
			
		self.viewProviderTotal.display(0)
		self.provProducts.setText(str(self.prods))
		self.totalProviderCash.setText("0")
		self.totalProviderCredit.setText("0")
		
		self.providerName.setText("")
		self.providerNumber.setText("")

	def validTransaction(self):
		try:
			self.statBoxcach= self.checkBoxCash.isChecked()
			self.statBoxcredit = self.checkBoxCredit.isChecked()
			self.statBoxcheck = self.checkBoxCheck.isChecked()
			if self.statBoxcach == False and self.statBoxcredit == False and self.statBoxcheck == False :
				self.messageFactory.raiseCaseExcept("La méthode de payement et au moin un produit")
			else:
				self.reference()
				self.prodTotal = []
				self.prodTotalList = []
				
				COUPON = Document()
				
				p = COUPON.add_paragraph(" ___________________________________________________________________________________________________ ")

				h = COUPON.add_heading("\t\t\t\t     "+self.society+"", level=1)
				h.bold = True
				h.italic = True
				p = COUPON.add_paragraph("\t\t\t\t       "+self.numero)
				
				p = COUPON.add_paragraph(" ___________________________________________________________________________________________________ ")
						
				pi = 0
				tab_prod = COUPON.add_table(1,4)
				heading_cells = tab_prod.rows[0].cells
				heading_cells[0].text = 'FOURNISSEUR'
				heading_cells[1].text = 'NUMERO '
				heading_cells[2].text = 'CATEGORY '
				heading_cells[3].text = ''
				
				cells = tab_prod.add_row().cells
				cells[0].text = self.providerName.text()
				cells[1].text = self.providerNumber.text()
				cells[2].text = self.providerProductType.currentText()
				cells[3].text = ''
					
				p = COUPON.add_paragraph(" ___________________________________________________________________________________________________ ")
				p = COUPON.add_paragraph(" ___________________________________________________________________________________________________ ")
						
				pi = 0
				tab_prod = COUPON.add_table(1,4)
				heading_cells = tab_prod.rows[0].cells
				heading_cells[0].text = 'PRODUIT'
				heading_cells[1].text = 'PRIX '
				heading_cells[2].text = 'QUANTITE '
				heading_cells[3].text = 'MONTANT'
				
				while pi <= 4 :
				
					self.prod = str(self.ProductPriceList[pi].text()).strip("(',')")
					self.prod = int(self.prod)
						
					self.quantite = str(self.productQuantite[pi].value())
					self.quantite = int(self.quantite)
					
					total = self.prod * self.quantite
					
					self.prodTotal.append(total)
					
					if self.quantite != 0:
						cells = tab_prod.add_row().cells
						cells[0].text = str(self.ProductList[pi].text()).strip("(',')")
						cells[1].text = str(self.ProductPriceList[pi].text()).strip("(',')") + "DA"
						cells[2].text = str(self.productQuantite[pi].text()).strip("(',')")
						cells[3].text = str(self.prodTotal[pi]) + "DA"
					pi+=1
					
				p = COUPON.add_paragraph(" ___________________________________________________________________________________________________ ")
				try :	
					p = COUPON.add_heading("Total à payer : "+str(sum(self.prodTotal))+" DA ", level=3)
					p.bold = True
				except :
					return 0
				
				p = COUPON.add_paragraph(" ___________________________________________________________________________________________________ ")
				self.mthd = ''
				p = COUPON.add_heading("Méthode de payement : " +self.mthd)
				
				if self.checkBoxCash.checkState() != 0:
					self.mthd = 'Payement en Cash'
				
				if self.checkBoxCheck.checkState() != 0:
					self.mthd = 'Payement par Chèque'
					p = COUPON.add_paragraph("Chèque No : "+self.checkNo.text()+"\nBanque : "+self.checkBanque.text()+".")
					p = COUPON.add_paragraph("Compte No : "+self.checkCompteNo.text()+"\nVirement prévu le : "+self.checkDateEdit.text()+".")
				
				if self.checkBoxCredit.checkState() != 0:
					self.mthd = 'Payement à Crédit'
					p = COUPON.add_paragraph("Somme à payer : "+self.totalProviderCredit.text()+" DA\nAvance : "+self.creditAvance.text()+" DA")
					p = COUPON.add_paragraph("Somme Restante : "+self.creditRemain.text()+" DA\nVirement restant prévu le : "+self.creditVirement.text()+".")
				
				p = COUPON.add_paragraph(" ___________________________________________________________________________________________________ ")
				p = COUPON.add_paragraph(" ___________________________________________________________________________________________________ ")
				p = COUPON.add_heading("Bon No : "+str(self.ref)+" --"+self.society+"-- "+self.localDateTime, level=3)
				p = COUPON.add_paragraph(" ___________________________________________________________________________________________________ ")
				COUPON.add_page_break()
				
				COUPON.save('DOCUMENTS/BONS/Bon-'+str(self.ref)+"_"+self.localDate+'.docx' )
				os.startfile('DOCUMENTS\BONS\Bon-'+str(self.ref)+"_"+self.localDate+'.docx' , 'print')
			
				# self.viewProviderRefNum.setText(str(self.ref))
				
				query.execute("SELECT ref FROM Ref WHERE id = 2")
				
				query.execute("UPDATE Ref SET ref = "+str(self.viewProviderRefNum.text())+" WHERE id = 2")
				
				self.transactType = "Fournitures | "+self.providerProductType.currentText()
				i=0
				for prod in self.productQuantite:
					self.transactDesc = str(prod.value())+" * "+self.ProductList[i].text()
					
					self.price = str(self.ProductPriceList[i].text())			
					self.price = str(self.price)
					self.price = int(self.price)
					
					# self.quantite = int(self.productQuantite[i].text())
					
					self.total = self.price * prod.value()
					
					if prod.value() != 0:
						query.execute('INSERT INTO Providers_History\
						(transact_name, transact_category, transact_phone, transact_products, transact_quantite, transact_total, transact_date_transact)\
						VALUES("{0}","{1}" ,"{2}" ,"{3}","{4}","{5}","{6}" );'\
						.format(self.providerName.text(),self.providerProductType.currentText(),self.providerNumber.text(),str(self.ProductList[i].text()),\
						str(prod.value()),str(self.total),self.Date))
				
						if self.statBoxcach == True :
							# print(prod.value())
							if prod.value() != 0 :
								query.execute("INSERT INTO Register_dep (dep_date, dep_type, dep_description,dep_value) VALUES ('"\
								+self.Date+"','"+str(self.transactType)+"','"+str(self.transactDesc)+"','"+str(self.total)+"')")
							i+=1
							
				conn.commit()
				self.register.actualDepense()
				self.checkBoxCash.setChecked(False)
				self.checkBoxCredit.setChecked(False)
				i = 0
				self.prods = 0
				while i <= 4 :
					
					self.ProductList[i].setText("")
				
					self.price = 0
					self.ProductPriceList[i].setText(str(self.price))
					
					self.quantite = 0
					self.productQuantite[i].setValue(self.quantite)  
						
					i+=1
					
				self.viewProviderTotal.display(0)
				self.provProducts.setText(str(self.prods))
				self.totalProviderCash.setText("0")
				self.totalProviderCredit.setText("0")
				
				self.providerName.setText("")
				self.providerNumber.setText("")
						
				self.ProdList.clear()
				
				self.prodTotal = []
				self.prodTotalList = []	
				
				self.providerCount()
			
		except:
			self.messageFactory.raisePrintExcept("Bon-"+str(self.ref)+"_"+self.localDate)
		
	def validSlot(self) :

		if self.sender() == self.validButton :
			self.totalTransaction()			

		if self.sender() == self.undoButton :	
			self.annulTransaction()
				
		if self.sender() == self.nextButton :
			self.validTransaction()
	
	def providersListSignal(self):
		
		self.providerListWidget.itemClicked.connect(self.listwidgetClickedSlot)

	def listwidgetClickedSlot(self, item):

		i=0
		p=0
		for prov in self.providerEdit.providerNameList :
			if item.text() == str(prov).strip("(',')") +" | "+str(self.providerEdit.providerCategoryList[i]).strip("(',')"):
		
				self.providerName.setText(str(prov).strip("(',')"))
				self.providerNumber.setText(str(self.providerEdit.providerPhoneList[i]).strip("(',')"))
				self.providerProductType.clear()
				self.providerDataManager()
				self.providerProductType.setItemText(0,str(self.providerEdit.providerCategoryList[i]).strip("(',')"))
		
				self.ProductList[0].setText(str(self.providerEdit.providerProductList1[i]).strip("(',')"))
				self.ProductList[1].setText(str(self.providerEdit.providerProductList2[i]).strip("(',')"))
				self.ProductList[2].setText(str(self.providerEdit.providerProductList3[i]).strip("(',')"))
				self.ProductList[3].setText(str(self.providerEdit.providerProductList4[i]).strip("(',')"))
				self.ProductList[4].setText(str(self.providerEdit.providerProductList5[i]).strip("(',')"))
		
				self.ProductPriceList[0].setText(str(self.providerEdit.providerPriceList1[i]).strip("(',')"))
				self.ProductPriceList[1].setText(str(self.providerEdit.providerPriceList2[i]).strip("(',')"))
				self.ProductPriceList[2].setText(str(self.providerEdit.providerPriceList3[i]).strip("(',')"))
				self.ProductPriceList[3].setText(str(self.providerEdit.providerPriceList4[i]).strip("(',')"))
				self.ProductPriceList[4].setText(str(self.providerEdit.providerPriceList5[i]).strip("(',')"))
			i+=1
			
	def spinBoxSignal(self):
		self.checkBoxCredit.clicked.connect(self.spinCreditSlot)
		
	def spinCreditSlot(self):
		
		if self.totalProviderCredit.text() != "0":
			self.restant = int(self.totalProviderCredit.text()) - int(self.creditAvance.text())
			self.creditRemain.setText(str(self.restant))

	def EditSignals(self):
		#DIALOG SIGNALS :
		self.providerEdit.provAdd.clicked.connect(self.validEditSlots)
		
		self.providerEdit.provModif.clicked.connect(self.validEditSlots)
		
		self.providerEdit.provDel.clicked.connect(self.validEditSlots)
		
		self.providerEdit.providerEditOK.clicked.connect(self.validEditSlots)
		self.providerEdit.providerEditNO.clicked.connect(self.validEditSlots)	
	
	def providerCount(self):	
		#SET ProvidersLists AS QLISTWIDGET ITEMS :
		self.providerEdit.providerData()
		self.providerListWidget.clear()
		self.transactListWidget.clear()
		
		self.providerEdit.providerModifCombo.clear()
		self.providerEdit.providerDelCombo.clear()
		self.providerEdit.providerModifCombo.addItem("...")
		self.providerEdit.providerDelCombo.addItem("...")
		
		i=0
		for prov  in self.providerEdit.providerNameList :
			self.providerListWidget.addItem(str(prov).strip("(',')") +" | "+str(self.providerEdit.providerCategoryList[i]).strip("(',')"))
			
			self.providerNbr.setText(str(i+1))
			self.providerEdit.providerNbr.setText(str(i+1))
			
			self.providerEdit.providerModifCombo.addItem(str(prov).strip("(',')"))
			self.providerEdit.providerDelCombo.addItem(str(prov).strip("(',')"))
			i+=1
		
		query.execute("SELECT transact_name FROM Providers_History WHERE transact_date_transact ='"+self.Date+\
		"' ORDER BY transact_date_transact DESC")
		self.transactName = query.fetchall()
		query.execute("SELECT transact_category FROM Providers_History WHERE transact_date_transact ='"+self.Date+\
		"' ORDER BY transact_date_transact DESC")
		self.transactCat = query.fetchall()
		query.execute("SELECT transact_date_transact FROM Providers_History WHERE transact_date_transact ='"+self.Date+\
		"' ORDER BY transact_date_transact DESC")
		self.transactDate = query.fetchall()
		
		i=0
		for prov  in self.transactName :
			self.transactListWidget.addItem(str(prov).strip("(',')") +"  |  "+str(self.transactCat[i]).strip("(',')")+"  |  "\
			+str(self.transactDate[i]).strip("(',')"))
			i+=1
		
		self.transactName = []
		self.transactCat = []
		self.transactDate = []

	def editCombo(self):
		if self.sender() == self.providerEdit.providerModifCombo :
				
			self.Products = []
			self.Products = [self.providerEdit.product_modif1, self.providerEdit.product_modif2,\
			self.providerEdit.product_modif3, self.providerEdit.product_modif4, self.providerEdit.product_modif5]
			
			self.Prices = []
			self.Prices = [self.providerEdit.price_modif1,self.providerEdit.price_modif2,\
			self.providerEdit.price_modif3,self.providerEdit.price_modif4,self.providerEdit.price_modif5]
			self.provProduct = []
			query.execute("SELECT provider_product1 FROM Providers WHERE provider_name = '"+self.providerEdit.providerModifCombo.currentText()+"'")
			self.provProduct.append(query.fetchone())
			query.execute("SELECT provider_product2 FROM Providers WHERE provider_name = '"+self.providerEdit.providerModifCombo.currentText()+"'")
			self.provProduct.append(query.fetchone())
			query.execute("SELECT provider_product3 FROM Providers WHERE provider_name = '"+self.providerEdit.providerModifCombo.currentText()+"'")
			self.provProduct.append(query.fetchone())
			query.execute("SELECT provider_product4 FROM Providers WHERE provider_name = '"+self.providerEdit.providerModifCombo.currentText()+"'")
			self.provProduct.append(query.fetchone())
			query.execute("SELECT provider_product5 FROM Providers WHERE provider_name = '"+self.providerEdit.providerModifCombo.currentText()+"'")
			self.provProduct.append(query.fetchone())
			
			self.provPrice = []
			query.execute("SELECT provider_price1 FROM Providers WHERE provider_name = '"+self.providerEdit.providerModifCombo.currentText()+"'")
			self.provPrice.append(query.fetchone())
			query.execute("SELECT provider_price2 FROM Providers WHERE provider_name = '"+self.providerEdit.providerModifCombo.currentText()+"'")
			self.provPrice.append(query.fetchone())
			query.execute("SELECT provider_price3 FROM Providers WHERE provider_name = '"+self.providerEdit.providerModifCombo.currentText()+"'")
			self.provPrice.append(query.fetchone())
			query.execute("SELECT provider_price4 FROM Providers WHERE provider_name = '"+self.providerEdit.providerModifCombo.currentText()+"'")
			self.provPrice.append(query.fetchone())
			query.execute("SELECT provider_price5 FROM Providers WHERE provider_name = '"+self.providerEdit.providerModifCombo.currentText()+"'")
			self.provPrice.append(query.fetchone())
			
			query.execute("SELECT DISTINCT provider_name FROM Providers WHERE provider_name = '"+self.providerEdit.providerModifCombo.currentText()+"'")
			providerName = query.fetchone()
			
			query.execute("SELECT DISTINCT provider_phone FROM Providers WHERE provider_name = '"+self.providerEdit.providerModifCombo.currentText()+"'")
			providerPhone = query.fetchone()
			
			query.execute("SELECT DISTINCT provider_category FROM Providers WHERE provider_name = '"+self.providerEdit.providerModifCombo.currentText()+"'")
			providerCategory = query.fetchone()
			
			
			self.providerEdit.newProviderName.setText(str(providerName).strip("(',')"))
			self.providerEdit.newProviderNum.setText(str(providerPhone).strip("(',')"))
			self.providerEdit.newProviderCategory.setText(str(providerCategory).strip("(',')"))
			
			i=0
			for prov in self.provProduct :
				self.Products[i].setText(str(prov).strip("(',')"))
				# print(self.Products[i].text())
				# print(str(prov))
				i+=1
			j=0
			for prov in self.provPrice :
				self.Prices[j].setText(str(prov).strip("(',')"))
				# print(self.Prices[j].text())
				# print(str(prov))
				j+=1
	
		if self.sender() == self.providerEdit.providerDelCombo :
			
			query.execute("SELECT DISTINCT provider_name FROM Providers WHERE provider_name = '"+self.providerEdit.providerDelCombo.currentText()+"'")
			providerName = query.fetchone()
			
			query.execute("SELECT DISTINCT provider_category FROM Providers WHERE provider_name = '"+self.providerEdit.providerDelCombo.currentText()+"'")
			providerCategory = query.fetchone()
			
			
			self.providerEdit.delProviderName.setText(str(providerName).strip("(',')")+" | "+str(providerCategory).strip("(',')"))
				
	def providerAdder(self):
		try:	
			self.newProducts = []
			self.newProducts = [self.providerEdit.dialog_product1,self.providerEdit.dialog_product2,\
			self.providerEdit.dialog_product3,self.providerEdit.dialog_product4,self.providerEdit.dialog_product5]
			
			self.newPrices = []
			self.newPrices = [self.providerEdit.dialog_productPrice1,self.providerEdit.dialog_productPrice2,\
			self.providerEdit.dialog_productPrice3,self.providerEdit.dialog_productPrice4,self.providerEdit.dialog_productPrice5]
		
			INSERTER1 = "INSERT INTO Providers (provider_name, provider_category, provider_phone, provider_date_add,"
			INSERTER2 = "provider_product1, provider_product2,provider_product3, provider_product4, provider_product5,"
			INSERTER3 = "provider_price1, provider_price2,provider_price3, provider_price4, provider_price5)"
			
			VALUES1 = " VALUES ('"+self.providerEdit.dialog_providerName.text()+"','"+self.providerEdit.dialog_providerCategory.text() +\
				"','"+self.providerEdit.dialog_providerNumber.text()+"','"+self.localDate+"',"
				
			VALUES2 = "'"+str(self.newProducts[0].text())+"','"+str(self.newProducts[1].text())+"','"+str(self.newProducts[2].text())+\
				"','"+str(self.newProducts[3].text())+"','"+str(self.newProducts[4].text())+"'"
			
			VALUES3 = ","+str(self.newPrices[0].text())+","+str(self.newPrices[1].text())+","+str(self.newPrices[2].text())+\
				","+str(self.newPrices[3].text())+","+str(self.newPrices[4].text())+")"
			
			if self.providerEdit.dialog_providerName.text() != '' and self.providerEdit.dialog_providerNumber.text() != "" and self.providerEdit.dialog_providerCategory.text() != "":
				
				query.execute(INSERTER1+INSERTER2+INSERTER3+VALUES1+VALUES2+VALUES3)
				
				conn.commit()
				
				self.providerCount()
				self.providerEdit.providerData()
				
				self.providerEdit.dialog_providerName.setText("")
				self.providerEdit.dialog_providerNumber.setText("")
				self.providerEdit.dialog_providerCategory.setText("")
				
				for prod in self.newProducts :
					prod.setText("")
				
				for prod in self.newPrices :
					prod.setText("0")
				
				self.newProducts = []
				self.newPrices = []
				
				self.messageFactory.raiseAdder("Fournisseur")
					
			else :  
				self.messageFactory.raiseCaseExcept("toutes les cases")
				if self.providerEdit.dialog_providerName.text() == '':
					self.messageFactory.raiseCaseExcept("le Nom Fournisseur")
				elif self.providerEdit.dialog_providerNumber.text() == '':
					self.messageFactory.raiseCaseExcept("le Numero Fournisseur")
				elif self.providerEdit.dialog_providerCategory.text() == '':
					self.messageFactory.raiseCaseExcept("le Type de produits Fournit")
		except:
			self.messageFactory.raiseCharExcept()
		
	def providerModifier(self):
		try:	
			if self.providerEdit.newProviderName.text() != '' and self.providerEdit.newProviderCategory.text() != "" and self.providerEdit.newProviderNum.text() != "" and self.providerEdit.providerModifCombo.currentText() != "...":
				
				
				VALUE1 = self.providerEdit.newProviderName.text()
				VALUE2 = self.providerEdit.newProviderCategory.text()
				VALUE3 = self.providerEdit.newProviderNum.text()
				
				UPDATER1 = "UPDATE Providers SET provider_name = '"+VALUE1+"', provider_category='"+VALUE2+"', provider_phone='"+VALUE3+"'\
				WHERE provider_name = '"+self.providerEdit.providerModifCombo.currentText()+"'"
				query.execute(UPDATER1)
				
				i=1
				for prov in self.Products :
					query.execute("UPDATE Providers SET provider_product"+str(i)+"='"+self.Products[i-1].text()+"'\
					WHERE provider_name = '"+self.providerEdit.providerModifCombo.currentText()+"'")
					query.execute("UPDATE Providers SET provider_price"+str(i)+"="+self.Prices[i-1].text()+"\
					WHERE provider_name = '"+self.providerEdit.providerModifCombo.currentText()+"'")
					i+=1
					
				conn.commit()
				
				self.providerCount()
				self.providerEdit.providerData()
				
				self.providerEdit.newProviderName.setText("")
				self.providerEdit.newProviderCategory.setText("")
				self.providerEdit.newProviderNum.setText("")
				
				for prod in self.Products :
					prod.setText("")
				for prod in self.Prices :
					prod.setText("")
					
				self.messageFactory.raiseModifier("Fournisseur")
					
			else :  
				self.messageFactory.raiseCaseExcept("toutes les cases")
				if self.providerEdit.self.providerEdit.providerModifCombo.currentText() == '...':
					self.messageFactory.raiseCaseExcept("le Fournisseur")
				if self.providerEdit.newProviderName.text() == '':
					self.messageFactory.raiseCaseExcept("le Nom du Fournisseur")
				elif self.providerEdit.newProviderNum.text() == '':
					self.messageFactory.raiseCaseExcept("le Numero Fournisseur")
				elif self.providerEdit.newProviderCategory.text() == '':
					self.messageFactory.raiseCaseExcept("le type de produit Fournit")
				
		except:
			self.messageFactory.raiseCharExcept()
		
	def providerDeleter(self):
		try:	
			if self.providerEdit.delProviderName.text() != '' and self.providerEdit.providerDelCombo.currentText() != "...":
				
				query.execute("DELETE FROM 'Providers' WHERE provider_name = '"+self.providerEdit.providerDelCombo.currentText()+"'")
					
				conn.commit()
				
				self.providerCount()
				self.providerEdit.providerData()
				
				self.providerEdit.delProviderName.setText("")
				
				self.messageFactory.raiseDeleter("Fournisseur")
					
			else :  
				self.messageFactory.raiseCaseExcept("toutes les cases")
				if self.providerEdit.providerDelCombo.currentText() == '...':
					self.messageFactory.raiseCaseExcept("le Fournisseur")
				if self.providerEdit.delProviderName.text() == '':
					self.messageFactory.raiseCaseExcept("le Nom du Fournisseur")
				msg.exec_()	
			
		except:
			self.messageFactory.raiseCharExcept()
				
	def validEditSlots(self):
	
		if self.sender() == self.providerEditer :
				
			self.providerEdit.providerModifCombo.currentTextChanged.connect(self.editCombo)
			self.providerEdit.providerDelCombo.currentTextChanged.connect(self.editCombo)
	
			self.providerEdit.show()
			self.providerEdit.dialog_providerName.setText("")
			self.providerEdit.dialog_providerNumber.setText("")
			self.providerEdit.dialog_providerCategory.setText("")
			
			for item in self.providerEdit.providersProductList:
				item.setText("")
			for item in self.providerEdit.providersProductPriceList:
				item.setText("0")
	
		if self.sender() == self.providerEdit.provAdd:			
			self.providerAdder()
			
		if self.sender() == self.providerEdit.provModif:			
			self.providerModifier()	
			
		if self.sender() == self.providerEdit.provDel:			
			self.providerDeleter()	
			
		if self.sender() == self.providerEdit.providerEditOK:			
			self.providerEdit.destroy()
						
		if self.sender() == self.providerEdit.providerEditNO:
			self.providerEdit.close()	
		 
#============================================================================================================WORKERS 
#============WORKERS DIALOG	
			
#=============CATEGORY DIALOG
qtwCreatorDialog= "DESIGN/DIALOGS/workerCatCreatorDialog.ui"
Ui_wcategoryCreatorDialog, QtBaseClass = uic.loadUiType(qtwCreatorDialog)

class wcategoryCreatorDialog(QDialog, Ui_wcategoryCreatorDialog):# EDIT : MODIF Product Catégory DIALOG

	def __init__(self):
		QDialog.__init__(self)
		Ui_wcategoryCreatorDialog.__init__(self)
		self.setupUi(self)
	
		# INSTANCE
		self.messageFactory = MessageFactory()
		#SELF
		self.setWindowTitle("Edition des catégories")
		# SETTING TEXT  category_name FROM CATEGORIES -->  CATLIST
  
		#CATEGORY NAMES QEDIT LIST
		self.categoryText=[]
		self.categoryText = [self.dialogCat1, self.dialogCat2, self.dialogCat3,self.dialogCat4, self.dialogCat5,\
		self.dialogCat6, self.dialogCat7, self.dialogCat8, self.dialogCat9, self.dialogCat10]
		
		#DELEGATES
		self.categorySelector()	
		self.signals()
		
	def categorySelector(self) :
		query.execute("SELECT category_name FROM Workers_category ORDER BY id ASC")
		self.wCatList = query.fetchall()
		i=0
		for cat in self.wCatList :
			self.categoryText[i].setText(str(cat).strip("(',')"))
			self.categorytabWidget.setTabText(i,str(cat).strip("(',')"))
			i+=1

	def signals(self):
		self.catOK.clicked.connect(self.slots)
		self.catAnnuler.clicked.connect(self.slots)
	
	def categoryValid(self):
		try :
			i=1
			for cat in self.categoryText :
				query.execute("UPDATE Workers_category set category_name ='"+cat.text()+"' WHERE id = "+str(i)+"")
				i+=1
					
			conn.commit()
			self.categorySelector()
			self.messageFactory.raiseAdder("Catégorie")
			self.destroy()
		except:
			self.messageFactory.raiseCharExcept()
			
	def slots(self):
		if self.sender() == self.catAnnuler :
			self.destroy()
		if self.sender() == self.catOK :
			self.categoryValid()

#=============PRODUCT CREATOR DIALOG

qtworkerCreatorCreator= "DESIGN\DIALOGS\workerCreatorDialog.ui"
Ui_workerCreatorDialog, QtBaseClass = uic.loadUiType(qtworkerCreatorCreator)

class WorkerCreatorDialog(QDialog, Ui_workerCreatorDialog):#EDIT : MODIF Product Name,Price DIALOG
	def __init__(self):
		QDialog.__init__(self)
		Ui_workerCreatorDialog.__init__(self)
		self.setupUi(self)
		
		self.messageFactory=MessageFactory()
		self.catDialog = wcategoryCreatorDialog()
		self.setWindowTitle("Gestion du Personnel")
		
		self.Date = time.strftime("%Y-%m-%d")
		
		self.catSelector()
		self.catDialog.categorySelector()
		self.signals()

	def signals(self):
		self.workerAdd.clicked.connect(self.slots)
		self.workerModif.clicked.connect(self.slots)
		self.workerDel.clicked.connect(self.slots)
		
		self.workOK.clicked.connect(self.slots)
		self.workAnnuler.clicked.connect(self.slots)

	def workerAdder(self):
		try:
			if self.addComboBox.currentText() != "..." and self.worker.text() != "" and self.workerTel.text() != "" and self.workerSalary.text() != "" :
				query.execute("INSERT INTO Workers(worker_name,worker_category,worker_number,worker_salary,worker_salary_type,worker_start)\
				VALUES ('"+self.worker.text()+"','"+self.addComboBox.currentText()+"','"+self.workerTel.text()+"',"+self.workerSalary.text()+",\
				'"+self.workerPayType.currentText()+"','"+self.Date+"')")
				conn.commit()
				
				self.messageFactory.raiseAdder("Employé")
				
				self.worker.setText("")
				self.workerTel.setText("")
				self.workerSalary.setText("0")
				
				self.catSelector()
			else :
				self.messageFactory.raiseCaseExcept("La category d'employé désiré")
				self.messageFactory.raiseCaseExcept("Toutes les cases")
		except:
			self.messageFactory.raiseCharExcept()

	def workerModifier(self):
		try:
			if self.modifComboBox.currentText() != "..." and self.modifComboList.currentText() != "..." and self.modifWorker.text() != ""\
			and self.modifTel.text() != "" and self.modifSalary.text() != "" :
				query.execute("UPDATE Workers SET worker_name='"+self.modifWorker.text()+"', worker_number='"+self.modifTel.text()+"',\
				worker_salary="+self.modifSalary.text()+", worker_salary_type='"+self.modifPay.currentText()+"' \
				WHERE worker_name='"+self.modifComboList.currentText()+"'")
			
				conn.commit()
				
				self.messageFactory.raiseModifier("Employé")
				
				self.modifWorker.setText("")
				self.modifTel.setText("")
				self.modifSalary.setText("0")
				
				self.catSelector()
			else :
				self.messageFactory.raiseCaseExcept("La category et le Nom d'employé")
				self.messageFactory.raiseCaseExcept("Toutes les cases")
		except:
			self.messageFactory.raiseCharExcept()

	def workerDeleter(self):
		try:
			if self.delComboBox.currentText() != "..." and self.delComboList.currentText() != "..." and self.delWorker.text() != "" :
				query.execute("DELETE FROM 'Workers' WHERE worker_name='"+self.delComboList.currentText()+"'")
			
				conn.commit()
				
				self.messageFactory.raiseDeleter("Employé")
				
				self.delWorker.setText("")
				
				self.catSelector()
			else :
				self.messageFactory.raiseCaseExcept("La category et le Nom d'employé")
				self.messageFactory.raiseCaseExcept("Toutes les cases")
		except:
			self.messageFactory.raiseCharExcept()
			
	def slots(self):
		if self.sender() == self.workerAdd :
			self.workerAdder()
		if self.sender() == self.workerModif :
			self.workerModifier()
		if self.sender() == self.workerDel :
			self.workerDeleter()
		
		if self.sender() == self.workAnnuler : 
			self.destroy()
			
		if self.sender() == self.workOK: 
			self.destroy()
	
	def catSelector(self):
		self.addComboBox.clear()
		self.addComboBox.addItem('...')
		self.modifComboBox.clear()
		self.modifComboBox.addItem('...')
		self.delComboBox.clear()
		self.delComboBox.addItem('...')
		
		self.catDialog.categorySelector()
		for cat in self.catDialog.categoryText :
			if cat.text() != "":
				self.addComboBox.addItem(cat.text())
				self.modifComboBox.addItem(cat.text())
				self.delComboBox.addItem(cat.text())
		#SIGNALS
		self.modifComboBox.currentTextChanged.connect(self.comboSlots)
		self.delComboBox.currentTextChanged.connect(self.comboSlots)
		
		self.modifComboList.currentTextChanged.connect(self.comboSlots)
		self.delComboList.currentTextChanged.connect(self.comboSlots)
	
	def comboSlots(self):
		if self.sender() == self.modifComboBox :
			query.execute("SELECT worker_name FROM Workers WHERE worker_category = '"+self.modifComboBox.currentText()+"'")
			workerList = query.fetchall()
			self.modifComboList.clear()
			self.modifComboList.addItem("...")
			for worker in workerList :
				self.modifComboList.addItem(str(worker).strip("(',')"))
			
		if self.sender() == self.delComboBox :
			query.execute("SELECT worker_name FROM Workers WHERE worker_category = '"+self.delComboBox.currentText()+"'")
			workerList = query.fetchall()
			self.delComboList.clear()
			self.delComboList.addItem("...")
			for worker in workerList :
				self.delComboList.addItem(str(worker).strip("(',')"))
			
		if self.sender() == self.modifComboList :
			query.execute("SELECT worker_name FROM Workers WHERE worker_name = '"+self.modifComboList.currentText()+"'")
			worker = query.fetchone()
			query.execute("SELECT worker_number FROM Workers WHERE worker_name = '"+self.modifComboList.currentText()+"'")
			workerTel = query.fetchone()
			query.execute("SELECT worker_salary FROM Workers WHERE worker_name = '"+self.modifComboList.currentText()+"'")
			workerSalary = query.fetchone()
			
			self.modifWorker.setText(str(worker).strip("(',')"))
			self.modifTel.setText(str(workerTel).strip("(',')"))
			self.modifSalary.setText(str(workerSalary).strip("(',')"))
			
		if self.sender() == self.delComboList :
			query.execute("SELECT worker_name FROM Workers WHERE worker_name = '"+self.delComboList.currentText()+"'")
			worker = query.fetchone()
			
			self.delWorker.setText(str(worker).strip("(',')"))
	
#========== PRODUCT WIDGET UIC CONVERT & Load

qtWorkerWidget = "DESIGN/WIDGETS/WorkersWidget.ui" # Enter file here.

Ui_WorkerWidget, QtBaseClass = uic.loadUiType(qtWorkerWidget)

#==========INIT CLASS

class WorkerWidget(QWidget, Ui_WorkerWidget):

	def __init__(self):
	
		QWidget.__init__(self)#QWidget init
		Ui_WorkerWidget.__init__(self)#ui init
		self.setupUi(self)#UI init
				
		#==========Logical & Réusables Variables
			
		self.localDateTime = time.strftime("%d-%m-%Y      %H:%M:%S")
		self.localDate = time.strftime("%d-%m-%Y")
		
		self.Date = time.strftime("%Y-%m-%d")
		
		self.society = "Resto Manager Sys" 
		self.numero = "xx-xx-xx-xx-xx"  

		#==========INIT CLASSES OF TYPE DIALOGS FROM MODULE "CreatorDialog":
		
		self.categoryDialog = wcategoryCreatorDialog()#EDIT : CATEGORY DIALOG
		self.workerDialog = WorkerCreatorDialog()#EDIT : Personnel DIALOG
		self.messageFactory = MessageFactory()
		self.register = RegisterCreatorDialog()
		
		self.categoryDialog.categorySelector()
		
		#MENU CAT LIST
		self.menuCategory = [self.category_1,self.category_2,self.category_3,self.category_4,self.category_5,\
		self.category_6,self.category_7,self.category_8,self.category_9,self.category_10]
		
		#DELEGATES
		self.buttonSignals()
		self.workerSelecter()
		self.menuSelector()#SELECTOR AND SIGNAL : CATEGORY & WORKERLIST
	
	def workerSelecter(self):
		query.execute("SELECT worker_name FROM Workers ORDER BY worker_category")
		self.workerSelect=query.fetchall()
		query.execute("SELECT worker_category FROM Workers ORDER BY worker_category")
		self.workerCatSelect=query.fetchall()
	
	def menuSelector(self):
		self.categoryDialog.categorySelector()
		i=0
		for cat in self.categoryDialog.categoryText :
			self.menuCategory[i].setText(str(cat.text()))
			i+=1
		#SIGNAL : MENU CATEGORY CLICKED
		i = 0
		while i<=9:
			self.menuCategory[i].clicked.connect(self.menuSlots)
			i+=1
		self.workerList.itemClicked.connect(self.workerSlots)
			
	def menuSlots(self):
		self.workerList.clear()
		for cat in self.menuCategory:
			if self.sender() == cat :	
				# print(str(cat))
				query.execute("SELECT worker_name FROM Workers WHERE worker_category ='"+str(cat.text())+"' ORDER BY id ASC")
				workers = query.fetchall()
				i=0
				for work in workers :	
					self.workerList.addItem(str(work).strip("(',')"))
					i+=1
				self.viewRefWorker.setText(str(i))

	def workerSlots(self, item):
		self.workerSelecter()
		for prov  in self.workerSelect :
			if item.text() == str(prov).strip("(',')"):	
				query.execute("SELECT worker_category FROM Workers WHERE worker_name = '"+str(prov).strip("(',')")+"'")
				workerType_=str(query.fetchone()).strip("(',')")
				query.execute("SELECT worker_number FROM Workers WHERE worker_name = '"+str(prov).strip("(',')")+"'")
				workerTel_=str(query.fetchone()).strip("(',')")
				query.execute("SELECT worker_salary FROM Workers WHERE worker_name = '"+str(prov).strip("(',')")+"'")
				workerSalary_=str(query.fetchone()).strip("(',')")
				query.execute("SELECT worker_salary_type FROM Workers WHERE worker_name = '"+str(prov).strip("(',')")+"'")
				workerSalaryType_=str(query.fetchone()).strip("(',')")
				query.execute("SELECT worker_start FROM Workers WHERE worker_name = '"+str(prov).strip("(',')")+"'")
				workerStart_=str(query.fetchone()).strip("(',')")
				query.execute("SELECT worker_payday FROM Workers WHERE worker_name = '"+str(prov).strip("(',')")+"'")
				workerPayDay_=str(query.fetchone()).strip("(',')")
				query.execute("SELECT worker_advance FROM Workers WHERE worker_name = '"+str(prov).strip("(',')")+"'")
				workerAdvance_=str(query.fetchone()).strip("(',')")
				query.execute("SELECT worker_advance_date FROM Workers WHERE worker_name = '"+str(prov).strip("(',')")+"'")
				workerAdvanceDate_=str(query.fetchone()).strip("(',')")
				
				self.worker.setText(str(prov).strip("(',')"))
				self.workerType.setText(workerType_)
				self.workerTel.setText(workerTel_)
				self.workerSalary.setText(workerSalary_)
				self.workerSalaryType.setText(workerSalaryType_)
				self.workerDate.setText(workerStart_)
				self.workerPayDay.setText(workerPayDay_)
				self.workerAvance.setText(workerAdvance_)
				self.workerAvanceDate.setText(workerAdvanceDate_)

	def avanceFunc(self):
		try:
			if self.workerAvance.text()!= "None" and self.workerAvance.text() != "" and self.workerAvance.text() != "0" :
				query.execute("UPDATE Workers SET worker_advance='"+self.workerAvance.text()+"', worker_advance_date='"+self.Date+"' \
				WHERE worker_name='"+self.worker.text()+"'")
				
				self.messageFactory.raiseAdder("Avance ")
				self.workerAvanceDate.setText(self.Date)
				
				if self.payMthd.currentText() == "Caisse" :
					query.execute("INSERT INTO Register_dep (dep_date, dep_type, dep_description, dep_value) VALUES ('"\
					+self.Date+"','Avance','Avance | "+self.worker.text()+"','"+self.workerAvance.text()+"')")
					self.register.actualDepense()
				
				query.execute("INSERT INTO Worker_avance (worker_category, worker_name, worker_start, worker_advance,worker_advance_date) VALUES (\
				'"+self.workerType.text()+"', '"+self.worker.text()+"', '"+self.Date+"', '"+self.workerAvance.text()+"', '"+self.Date+"')")
				conn.commit()
			
			else :
				self.messageFactory.raiseCaseExcept("La case Avance")
		except:
			self.messageFactory.raiseCaseExcept()
	
	def salaryFunc(self):
		if self.workerAvance.text()!= "None" and self.workerAvance.text() != "" and self.workerAvance.text() != "0" :
			salary = str(int(self.workerSalary.text()) - int(self.workerAvance.text()))
			
			query.execute("UPDATE Workers SET worker_payday='"+self.Date+"' WHERE worker_name='"+self.worker.text()+"'")
			
			self.messageFactory.raiseAdder("salaire")
			self.workerPayDay.setText(self.Date)
			
			if self.payMthd.currentText() == "Caisse" :
				query.execute("INSERT INTO Register_dep (dep_date, dep_type, dep_description, dep_value) VALUES ('"\
				+self.Date+"','Salaire','Salaire | "+self.worker.text()+"','"+salary+"')")
				self.register.actualDepense()
			
			query.execute("INSERT INTO Worker_salaire (worker_category, worker_name, worker_start, worker_salary,worker_salary_date) VALUES (\
			'"+self.workerType.text()+"', '"+self.worker.text()+"', '"+self.Date+"', '"+salary+"', '"+self.Date+"')")
			conn.commit()
		
		else :	
			query.execute("UPDATE Workers SET worker_payday='"+self.Date+"' WHERE worker_name='"+self.worker.text()+"'")
			
			self.messageFactory.raiseAdder("salaire")
			self.workerPayDay.setText(self.Date)
			
			if self.payMthd.currentText() == "Caisse" :
				query.execute("INSERT INTO Register_dep (dep_date, dep_type, dep_description, dep_value) VALUES ('"\
				+self.Date+"','Salaire','Salaire | "+self.worker.text()+"','"+self.workerSalary.text()+"')")
				self.register.actualDepense()
			
			query.execute("INSERT INTO Worker_salaire (worker_category, worker_name, worker_start, worker_salary,worker_salary_date) VALUES (\
			'"+self.workerType.text()+"', '"+self.worker.text()+"', '"+self.Date+"', '"+self.workerSalary.text()+"', '"+self.Date+"')")
			conn.commit()

	def workerLister(self):
		self.workerSelecter()
		self.workerList.clear()
		i=0
		for work in self.workerSelect :	
			self.workerList.addItem(str(i+1)+") "+str(self.workerCatSelect[i]).strip("(',')")+" | "+str(work).strip("(',')"))
			i+=1
		self.viewRefWorker.setText(str(i))
					
	def buttonSignals(self):
		self.workerCatEdit.clicked.connect(self.buttonSlots)
		self.workerEdit.clicked.connect(self.buttonSlots)
		
		#VALIDATION BUTTONS
		self.avanceButton.clicked.connect(self.buttonSlots)
		self.salaireButton.clicked.connect(self.buttonSlots)
		self.lister.clicked.connect(self.buttonSlots)
		
	def buttonSlots(self):
		if self.sender() == self.workerCatEdit :
			self.categoryDialog.show()
		if self.sender() == self.workerEdit :
			self.workerDialog.show()
			
		if self.sender() == self.avanceButton :
			self.avanceFunc()
			
		if self.sender() == self.salaireButton :
			self.salaryFunc()
			
		if self.sender() == self.lister :
			self.workerLister()
			
#============================================================================================================StockIniter 

#========== ACCOUNT WIDGET UIC CONVERT & Load
#======================================================

qtCreatorFile = "DESIGN/DIALOGS/stockIniter.ui" # Enter file here.

Ui_StockIniter, QtBaseClass = uic.loadUiType(qtCreatorFile)

class StockIniter(QDialog, Ui_StockIniter):
	def __init__(self):
		QDialog.__init__(self)
		Ui_StockIniter.__init__(self)
		self.setupUi(self)	
		self.setWindowTitle("INITIALISATION DU STOCK.")	
#============================================================================================================StockIniter 

#========== ACCOUNT WIDGET UIC CONVERT & Load
#======================================================

qtCreatorFile = "DESIGN/DIALOGS/stockManage.ui" # Enter file here.

Ui_StockManage, QtBaseClass = uic.loadUiType(qtCreatorFile)

class StockManage(QDialog, Ui_StockManage):
	def __init__(self):
		QDialog.__init__(self)
		Ui_StockManage.__init__(self)
		self.setupUi(self)	
		self.setWindowTitle("INITIALISATION DU STOCK.")		
#============================================================================================================STOCK WIDGET 

#========== ACCOUNT WIDGET UIC CONVERT & Load
#======================================================

qtCreatorFile = "DESIGN/DIALOGS/stockCreatorDialog.ui" # Enter file here.

Ui_StockCreatorDialog, QtBaseClass = uic.loadUiType(qtCreatorFile)

class StockCreatorDialog(QDialog, Ui_StockCreatorDialog):
	def __init__(self):
		QDialog.__init__(self)
		Ui_StockCreatorDialog.__init__(self)
		self.setupUi(self)
		
		self.setWindowTitle("GESTION DU STOCK.")
		
		self.messageFactory = MessageFactory()
		self.products = ProductCreatorDialog()
		self.stockIniter = StockIniter()
		self.stockManage = StockManage()
		
		self.stockListing()
		self.stockAlerter()
		self.signals()
		
	def signals(self) :
		self.stockinit.clicked.connect(self.stockInit)
		self.stockValid.clicked.connect(self.stockInit)
		self.stockIniter.initValid.clicked.connect(self.stockInit)
		
		self.stockManage.stockValid.clicked.connect(self.stockManagerValid)
		
	def stockAlerter(self):
		self.products.prodDataList()
		i =0
		for stock in self.products.prodStockList :
			if str(self.products.prodStockableList[i]).strip("(',')") != "0" :
				if int(str(stock).strip("(',')")) <= int(str(self.products.prodAlertList[i]).strip("(',')")) :
					self.messageFactory.raiseStockAlert(str(self.products.prodNameList[i]).strip("(',')"))
			i+=1
			
	def stockListing(self):
		query.execute("SELECT product_category_name,product_name, product_stockable,product_alert FROM Products WHERE product_stock = 1 ORDER BY product_category_name ASC")
		
		self.stockTabWidget.clearContents()
		self.stockTabWidget.setColumnCount(4) 
		self.stockTabWidget.setRowCount(0)
		for row, form in enumerate(query):
			self.stockTabWidget.insertRow(row)
			for column, item in enumerate(form):
				self.stockTabWidget.setItem(row, column,QTableWidgetItem(str(item)))
		self.stockTabWidget.itemDoubleClicked.connect(self.stockManager)
				
	def stockManager(self, item):
		i=0
		for prod in self.products.prodNameList :
			if item.text() == str(prod).strip("(',')") and str(self.products.prodStockableList[i]).strip("(',')") != "0" :
				self.stockManage.show()
				self.stockManage.stockValue.setText(str(self.products.prodStockList[i]).strip("(',')"))
				self.stockManage.stockAlert.setText(str(self.products.prodAlertList[i]).strip("(',')"))
				self.stockManage.stockProd.setText(str(self.products.prodNameList[i]).strip("(',')"))
				break
			i+=1				

	def stockManagerValid(self):
		try:
			if self.stockManage.stockValue.text() != "0" and self.stockManage.stockAlert.text() != "0" :
				query.execute("UPDATE Products SET product_stockable="+self.stockManage.stockValue.text()+",\
				product_alert="+self.stockManage.stockAlert.text()+" WHERE product_name ='"+self.stockManage.stockProd.text()+"'")
				conn.commit()
				self.products.prodDataList()
				self.stockListing()
				self.stockAlerter()
				self.stockManage.destroy()
			else :
				self.messageFactory.raiseCaseExcept("Une valeur differente de 0")
		except:
			self.messageFactory.raiseCharExcept()

	def stockInit(self) :
		if self.sender() == self.stockValid:
			self.destroy()
		if self.sender() == self.stockinit:
			self.stockIniter.show()
		if self.sender() == self.stockIniter.initValid:
			try:
				if self.stockIniter.stockValue.text() != "0" :
					query.execute("UPDATE Products SET product_stockable ="+self.stockIniter.stockValue.text()+"")
					conn.commit()
					
					self.products.prodDataList()
					self.stockListing()
					self.stockAlerter()
					
					self.stockIniter.destroy()
				else :
					self.messageFactory.raiseCaseExcept("Une valeur différente de 0 pour le stock")
			except :
				self.messageFactory.raiseCharExcept()
			
#============================================================================================================MAINWINDOW
#========== PRODUCT WIDGET UIC CONVERT & Load
#======================================================

qtCreatorFile = "DESIGN/WIDGETS/mainWindow.ui" # Enter file here.

Ui_MainWindow, QtBaseClass = uic.loadUiType(qtCreatorFile)
		
#============================================================================================================ STATISTICS

qtCreatorFile = "DESIGN/WIDGETS/StatWidget.ui" # Enter file here.

Ui_StatWidget, QtBaseClass = uic.loadUiType(qtCreatorFile)
class StatWidget(QWidget, Ui_StatWidget):
	def __init__(self):
		QWidget.__init__(self)#QWidget init
		Ui_StatWidget.__init__(self)#ui init
		self.setupUi(self)#UI init
		
		#INIT LIST CAT STATISTICS
		self.statCat = [self.statCat1,self.statCat2,self.statCat3,self.statCat4,self.statCat5,self.statCat6,self.statCat7,self.statCat8]
		#delegates
		self.statSignals()
		#INSTANCES
		self.products=ProductCreatorDialog()
		self.register=RegisterCreatorDialog()
		
	def statSignals(self):
		for cat in self.statCat :
			cat.clicked.connect(self.catSlots)
			
	def catSlots(self):
		self.statInit()
		if self.sender() == self.statCat[0]:
			self.stat(self.products.prodNameList,self.products.prodComptList,"PRODUITS","EVOLUTION VENTES PRODUITS","Produits","Quantités",\
			0,self.products.prodNameList,self.products.prodComptList,"Null")
		if self.sender() == self.statCat[1]:
			self.stat(self.recetteDate,self.recette,"RECETTES","EVOLUTION DES RECETTE","Dates","Recettes (DA)",\
			0,self.recetteDate,self.recette,"Null")
		if self.sender() == self.statCat[2]:
			self.stat(self.recetteDate,self.suminit,"SOMMESS INITIALES","EVOLUTION DES SOMMES INITIALES","Dates","Sommes Initiales (DA)",\
			0,self.recetteDate,self.suminit,"Null")
		if self.sender() == self.statCat[3]:
			self.stat(self.depenseDate,self.depense,"LES DEPENSES GENERALES","EVOLUTION DES DEPENSES GLOBALES","Dates","Depenses (DA)",\
			0,self.registerDate,self.depenseCaisse,"LES DEPENSES VIA CAISSE")
		if self.sender() == self.statCat[4]:
			self.stat(self.ajoutDate,self.ajout,"LES DEPOTS","EVOLUTION DES DEPOTS","Dates","Déposition (DA)\
			",0,self.ajoutDate,self.ajout,"null")
		if self.sender() == self.statCat[5]:
			self.stat(self.transDate,self.transValue,"TRANSACTION ET FOURNITURES","EVOLUTION DES TRANSACTION","Dates","Transactions (DA)",\
			0,self.transDate,self.transValue,"LES DEPENSES VIA CAISSE")
		if self.sender() == self.statCat[6]:
			self.stat(self.avanceDate,self.avance,"LES AVANCES","EVOLUTION DES AVANCES","Dates","Avances (DA)",\
			0,self.avanceDate,self.avance,"LES DEPENSES VIA CAISSE")
		if self.sender() == self.statCat[7]:
			self.stat(self.salaryDate,self.salary,"LES SALAIRES","EVOLUTION DES SALAIRES","Dates","Salaire (DA)",\
			0,self.salaryDate,self.salary,"LES DEPENSES VIA CAISSE")
		
	def statInit(self):
		#REGISTER
		query.execute("SELECT register_recette_total FROM Register ORDER BY id ASC")
		self.recette = query.fetchall()
		query.execute("SELECT register_date FROM Register ORDER BY id ASC")
		self.recetteDate = query.fetchall()	
		query.execute("SELECT register_sum_init FROM Register ORDER BY id ASC")
		self.suminit = query.fetchall()
		#DEPENSES		
		query.execute("SELECT dep_date FROM Register_dep ORDER BY id ASC")
		self.depenseDate = query.fetchall()
		query.execute("SELECT dep_value FROM Register_dep ORDER BY id ASC")
		self.depense = query.fetchall()	
		query.execute("SELECT register_depense_total FROM Register ORDER BY id ASC")
		self.depenseCaisse = query.fetchall()
		query.execute("SELECT register_date FROM Register ORDER BY id ASC")
		self.registerDate = query.fetchall()
		#DEPOTS
		query.execute("SELECT add_date FROM Register_add ORDER BY id ASC")
		self.ajoutDate = query.fetchall()
		query.execute("SELECT add_value FROM Register_add ORDER BY id ASC")
		self.ajout = query.fetchall()	
		#TRANSACTION
		query.execute(" SELECT  transact_total FROM Providers_History ORDER BY id ASC")
		self.transValue = query.fetchall()
		query.execute(" SELECT  transact_date_transact FROM Providers_History ORDER BY id ASC")
		self.transDate = query.fetchall()
		#SALARY
		query.execute(" SELECT  worker_salary FROM Worker_salaire ORDER BY id ASC")
		self.salary = query.fetchall()
		query.execute(" SELECT  worker_salary_date FROM Worker_salaire ORDER BY id ASC")
		self.salaryDate = query.fetchall()
		#AVANCES
		query.execute(" SELECT  worker_advance FROM Worker_avance ORDER BY id ASC")
		self.avance = query.fetchall()
		query.execute(" SELECT  worker_advance_date FROM Worker_avance ORDER BY id ASC")
		self.avanceDate = query.fetchall()
		
	def stat(self,listx,listy, labelP,titre,labelx,labely,test,listz,listk,labelP2):
		x=[]
		y=[]
		i=0
		for item in listx:
			absis = str(listx[i]).strip("(',')")
			try:
				x.append(int(absis))
			except:
				x.append(absis)
			absis = str(listy[i]).strip("(',')")
			try:
				y.append(int(absis))
			except:
				y.append(absis)
			i+=1
		if test != 0:
			z=[]
			k=[]
			i=0
			for item in listz:
				absis = str(listz[i]).strip("(',')")
				try:
					z.append(int(absis))
				except:
					z.append(absis)
				absis = str(listy[i]).strip("(',')")
				try:
					k.append(int(absis))
				except:
					k.append(absis)
				i+=1

		plt.plot(x, y, label=labelP)  # Plot some data on the (implicit) axes.
		if test != 0 :
			plt.plot(z, k, label=labelP2)  # etc.
		# plt.plot(x, x**3, label='cubic')
		plt.xlabel(labelx)
		plt.ylabel(labely)
		plt.title(titre)
		plt.legend()
		plt.show()	

#============================================================================================================ MAIN WINDOW
class MainWindow(QMainWindow, Ui_MainWindow):
	def __init__(self):
		QMainWindow.__init__(self)
		Ui_MainWindow.__init__(self)
		self.setupUi(self)
		
		self.setWindowTitle("Hotel Manager Solution.")
		self.Date = time.strftime("%Y-%m-%d")
		
		self.declaredUi()
		self.initUi()

		#============TABLE PART
		self.tableData()
		self.tableSignals()
		self.tableEditSignals()
		
		query.execute("SELECT table_name FROM Tables ORDER BY ID ASC")
		self.tabNameList = []
		self.tabNameList = query.fetchall()

	def declaredUi(self):
		self.messageFactory = MessageFactory()
		self.connecter = ConnectDialog()
		
		self.hotel = RoomWidget()
		self.home = ProductWidget()
		self.register = RegisterCreatorDialog()
		self.TableEdit = TableCreatorDialog()
		self.Account = AccountCreatorDialog()
		self.settings = SettingCreatorDialog()
		self.provider = ProviderWidget()
		self.worker = WorkerWidget()
		self.stock = StockCreatorDialog()
		self.stat = StatWidget()

		self.registerInit = InitDialog()
		
		self.hotel.setGeometry(0,0,1200,700)
		self.home.setGeometry(0,0,1200,700)
		self.provider.setGeometry(0,0,1200,700)
		self.worker.setGeometry(0,0,1200,700)
		self.stat.setGeometry(0,0,1200,700)
				
	def initUi(self):
		self.setCentralWidget(self.dockWidget)
		self.dockWidget.setWidget(self.hotel)
		
		self.HotelAct.triggered.connect(self.toolBarCalls)
		self.HomeAct.triggered.connect(self.toolBarCalls)
		self.CaisseAct.triggered.connect(self.toolBarCalls)
		self.TablesAct.triggered.connect(self.toolBarCalls)
		self.ComptaAct.triggered.connect(self.toolBarCalls)
		self.EditAct.triggered.connect(self.toolBarCalls)
		self.ProviderAct.triggered.connect(self.toolBarCalls)
		self.WorkerAct.triggered.connect(self.toolBarCalls)
		self.StockAct.triggered.connect(self.toolBarCalls)
		self.StatAct.triggered.connect(self.toolBarCalls)
		
		self.registerInit.INIT_OK.clicked.connect(self.closeWindow)
		self.registerInit.INIT_NO.clicked.connect(self.closeWindow)
		
		self.connecter.connOk.clicked.connect(self.toolBarCalls)
		
	def toolBarCalls(self):
	
		if self.sender() == self.HotelAct :
			self.hotel.freeRoomTester()
			self.dockWidget.setWidget(self.hotel)
	
		if self.sender() == self.HomeAct :
			self.dockWidget.setWidget(self.home)
	
		if self.sender() == self.CaisseAct :
			self.register.totalRegister()
			self.register.show()
	
		if self.sender() == self.TablesAct :
			self.TableEdit.show()
	
		if self.sender() == self.ComptaAct :
			self.Account.receiptBDD()
			self.Account.show()
	
		if self.sender() == self.EditAct :
			self.settings.selection()
			self.connecter.show()
		
		if self.sender() == self.connecter.connOk :
			self.connecter.login(self.settings)
		
		if self.sender() == self.ProviderAct :
			# self.settings.selection()
			self.dockWidget.setWidget(self.provider)
			
		if self.sender() == self.WorkerAct :
			self.dockWidget.setWidget(self.worker)
			
		if self.sender() == self.StatAct :
			self.dockWidget.setWidget(self.stat)
			
		if self.sender() == self.StockAct :
			self.stock.stockListing()
			self.stock.stockAlerter()
			self.stock.show()
	
	def closeEvent(self, event):
		event.ignore()
		query.execute("SELECT register_sum_init FROM Register WHERE register_date = '"+self.Date+"'")
		self.sominit = query.fetchone()
		self.sominit = str(self.sominit).strip("(',')")
		self.registerInit.sumInit.setText(self.sominit)
		self.registerInit.show()
		
	def closeWindow(self):
	
		if self.sender() == self.registerInit.INIT_OK :
			query.execute("UPDATE Register SET register_sum_init = '"+self.registerInit.sumInit.text()+"' WHERE register_date = '"+self.Date+"'")
			query.execute("SELECT register_sum_init FROM Register WHERE register_date = '"+self.Date+"'")
			self.sominit = query.fetchone()
			self.sominit = str(self.sominit).strip("(',')")
			self.registerInit.sumInit.setText(self.sominit)
			conn.commit()
			conn.close()
			self.registerInit.destroy()
			self.destroy()
		
		if self.sender() == self.registerInit.INIT_NO :
			self.registerInit.destroy()	
					
#================================TABLE PART	
		
	def tableData(self) :
		query.execute("SELECT table_name FROM Tables ORDER BY ID ASC")
		self.tabNameList = []
		self.tabNameList = query.fetchall()
		
		query.execute("SELECT table_place FROM Tables ORDER BY ID ASC")
		self.tabplaceList = []
		self.tabplaceList = query.fetchall()
		
		query.execute("SELECT table_state FROM Tables ORDER BY ID ASC")
		self.tabstateList = []
		self.tabstateList = query.fetchall()
		self.i = 0
		for tab in self.tabNameList :
			self.i+=1
		self.TableEdit.tableNbr.setText(str(self.i))
		
	def tableSignals(self) :
		self.home.tablesListWidget.itemClicked.connect(self.tableSlot)
		
	def tableSlot(self, item) :	
		
		query.execute("SELECT table_state FROM Tables ORDER BY ID ASC")
		self.tabstateList = []
		self.tabstateList = query.fetchall()
		self.i=0
		for tab in self.tabNameList :
			if item.text() == str(tab).strip("(',')"):
				self.home.viewTabNum.setText(item.text())
				
		self.home.calculTotal()
		self.home.prodListView()
		
	def tableEditSignals(self) :
		self.TableEdit.tableNbr.setText(str(self.i))
		self.TableEdit.tabOK.clicked.connect(self.tableEditSlot)
		self.TableEdit.tabAnnuler.clicked.connect(self.tableEditSlot)
		
		self.TableEdit.tabAdd.clicked.connect(self.tableEditSlot)
		self.TableEdit.tabModif.clicked.connect(self.tableEditSlot)
		self.TableEdit.tabDel.clicked.connect(self.tableEditSlot)
		
		self.TableEdit.modTabComboBox.currentTextChanged.connect(self.tabSubCat)
		self.TableEdit.delTabComboBox.currentTextChanged.connect(self.tabSubCat)
				
		self.TableEdit.numModifComboBox.currentTextChanged.connect(self.tabSubCat)	
		self.TableEdit.numDelComboBox.currentTextChanged.connect(self.tabSubCat)	

	def tableAdder(self):
		try:
			if self.TableEdit.newTabName.text() != "" and self.TableEdit.addTabComboBox.currentText() != "..." :
				query.execute("INSERT INTO Tables (table_name,table_place,table_total,table_state) VALUES ('Table "+str(self.TableEdit.newTabName.text())+\
				"','"+str(self.TableEdit.addTabComboBox.currentText())+"',0,'VIDE')")
				
				query.execute("SELECT table_name FROM Tables ORDER BY ID ASC")
				self.tabNameList = []
				self.tabNameList = query.fetchall()

				query.execute("SELECT table_place FROM Tables ORDER BY ID ASC")
				self.tabplaceList = []
				self.tabplaceList = query.fetchall()
						
				query.execute("SELECT table_state FROM Tables ORDER BY ID ASC")
				self.tabstateList = []
				self.tabstateList = query.fetchall()
				
				conn.commit()
				
				self.a=0
				self.home.tablesListWidget.clear()
				for tab in self.tabNameList :
					self.home.tablesListWidget.addItem(str(tab).strip("(',')"))
					self.a+=1
					
				self.TableEdit.tableNbr.setText(str(self.a))
				
				self.TableEdit.newTabName.setText("")
				self.TableEdit.addTabComboBox.clear()
				self.TableEdit.addTabComboBox.addItem('...')
				self.TableEdit.addTabComboBox.addItem('Interieur')
				self.TableEdit.addTabComboBox.addItem('Terasse')
				self.TableEdit.addTabComboBox.addItem('Etage')
				 
				self.messageFactory.raiseAdder("Table")
			
			else :
				self.messageFactory.raiseCaseExcept("toutes les cases")
				if  self.TableEdit.newTabName.text() == '':
					self.messageFactory.raiseCaseExcept("le Numéro de Table")
				if  self.TableEdit.addTabComboBox.currentText() == '...' :
					self.messageFactory.raiseCaseExcept("l'emplacment de Table")
		except:
			self.messageFactory.raiseCharExcept()	

	def tableModifer(self):
		try:
			if self.TableEdit.tabModifName.text() != "" and self.TableEdit.newModiTabComboBox.currentText() != "..." :
				query.execute("UPDATE Tables SET table_name ='"+self.TableEdit.tabModifName.text()+"',table_place ='"+self.TableEdit.newModiTabComboBox.currentText()+"'\
				WHERE table_name = '"+self.TableEdit.numModifComboBox.currentText()+"'")
				
				query.execute("SELECT table_name FROM Tables ORDER BY ID ASC")
				self.tabNameList = []
				self.tabNameList = query.fetchall()

				query.execute("SELECT table_place FROM Tables ORDER BY ID ASC")
				self.tabplaceList = []
				self.tabplaceList = query.fetchall()
						
				query.execute("SELECT table_state FROM Tables ORDER BY ID ASC")
				self.tabstateList = []
				self.tabstateList = query.fetchall()
				
				conn.commit()
				
				self.a=0
				self.home.tablesListWidget.clear()
				for tab in self.tabNameList :
					self.home.tablesListWidget.addItem(str(tab).strip("(',')"))
					self.a+=1
					
				self.TableEdit.tableNbr.setText(str(self.a))
				
				self.TableEdit.tabModifName.setText("")
				self.TableEdit.numModifComboBox.clear()
				
				self.messageFactory.raiseModifier("Table")
			
			else :
				self.messageFactory.raiseCaseExcept("toutes les cases")
				if  self.TableEdit.tabModifName.text() == '':
					self.messageFactory.raiseCaseExcept("le Nouveau Numéro de Table")
				if  self.TableEdit.numModifComboBox.currentText() == '...' :
					self.messageFactory.raiseCaseExcept("le Nouvel Emplacement de Table")
		except:
			self.messageFactory.raiseCharExcept()	

	def tableDeleter(self):
		try:
			if self.TableEdit.delTabName.text() != "" :
				query.execute("DELETE FROM 'Tables' WHERE table_name ='"+self.TableEdit.delTabName.text()+"'")
				
				query.execute("SELECT table_name FROM Tables ORDER BY ID ASC")
				self.tabNameList = []
				self.tabNameList = query.fetchall()

				query.execute("SELECT table_place FROM Tables ORDER BY ID ASC")
				self.tabplaceList = []
				self.tabplaceList = query.fetchall()
						
				query.execute("SELECT table_state FROM Tables ORDER BY ID ASC")
				self.tabstateList = []
				self.tabstateList = query.fetchall()
				
				conn.commit()
				
				self.a=0
				self.home.tablesListWidget.clear()
				for tab in self.tabNameList :
					self.home.tablesListWidget.addItem(str(tab).strip("(',')"))
					self.a+=1
					
				self.TableEdit.tableNbr.setText(str(self.a))
				
				self.TableEdit.delTabName.setText("")
				self.TableEdit.numDelComboBox.clear()
				self.messageFactory.raiseDeleter("Table")
			
			else :
				self.messageFactory.raiseCaseExcept("toutes les cases")
				if  self.TableEdit.delTabName.text() == '':
					self.messageFactory.raiseCaseExcept("le Numéro de Table")
		except:
			self.messageFactory.raiseCharExcept()

	def tableEditSlot(self) :	
		if self.sender() == self.TableEdit.tabAdd :
			self.tableAdder()
				
		if self.sender() == self.TableEdit.tabModif :
			self.tableModifer()
			
		if self.sender() == self.TableEdit.tabDel:
			self.tableDeleter()

		if self.sender() == self.TableEdit.tabOK or self.sender() == self.TableEdit.tabAnnuler :
			self.TableEdit.destroy()
			
	def tabSubCat(self):
		if self.sender() == self.TableEdit.modTabComboBox :
			#MODIF
			self.TableEdit.numModifComboBox.clear()

			query.execute("SELECT table_name FROM Tables WHERE table_place = '"+self.TableEdit.modTabComboBox.currentText()+"' ORDER BY id")
			self.tabList = []
			self.tabList = query.fetchall()

			for tab in self.tabList :
				self.TableEdit.numModifComboBox.addItem(str(tab).strip("(',')"))
			
		if self.sender() == self.TableEdit.numModifComboBox :
			self.TableEdit.tabModifName.setText(self.TableEdit.numModifComboBox.currentText())
			
		if self.sender() == self.TableEdit.delTabComboBox :
			#DEL
			self.TableEdit.numDelComboBox.clear()

			query.execute("SELECT table_name FROM Tables WHERE table_place = '"+self.TableEdit.delTabComboBox.currentText()+"' ORDER BY id")
			self.tabList = []
			self.tabList = query.fetchall()

			for tab in self.tabList :
				self.TableEdit.numDelComboBox.addItem(str(tab).strip("(',')"))
			
		if self.sender() == self.TableEdit.numDelComboBox :
			self.TableEdit.delTabName.setText(self.TableEdit.numDelComboBox.currentText())


class Opener(QDialog):
	def __init__(self):
		QDialog.__init__(self)
		
		self.declaredUi()
		
	def declaredUi(self):
		self.connecter = ConnectDialog()
		self.window = MainWindow()	
		
		self.connecter.connOk.clicked.connect(self.loginSlot)
		
	def loginSlot(self):
		if self.sender() == self.connecter.connOk :
			self.connecter.login(self.window)
				
#============================================================================================================END OPENER
	
if __name__ == "__main__":
	app = QApplication(sys.argv)
	opener = Opener()
	window = opener.connecter
	
	MAC=""
	MAC = uuid.UUID(int=uuid.getnode())
	MAC=str(MAC)
	
	query.execute("SELECT Key FROM MAC WHERE id=1")
	KEY = query.fetchone()
	tester = str(KEY).strip("(',')")

	if tester == '':
		query.execute("UPDATE MAC SET Key = '%s' WHERE id=1" %MAC)
		conn.commit()
		window.show()
		sys.exit(app.exec_())
	
	elif tester == str(MAC):
		window.show()
		sys.exit(app.exec_())
		
	elif tester != str(MAC):
		msg = QMessageBox()
		msg.setIcon(QMessageBox.Warning)

		msg.setText("Vous êtes en présence d'une version du produit obselete")
		msg.setInformativeText("La clé de sécurité est incompatible")
		msg.setWindowTitle("ALERT COPY OBSELETE !")
		msg.exec_()

#============================================================================================================END OF PROGRAMME

